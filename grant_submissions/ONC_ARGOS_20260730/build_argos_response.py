from __future__ import annotations

import hashlib
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
HERE = Path(__file__).resolve().parent
OUTPUT = HERE / "output"
MARKDOWN_PATH = HERE / "ARGOS_PARTNER_FIRST_CAPABILITY_RESPONSE_DRAFT.md"
DOCX_PATH = OUTPUT / "ARGOS_PARTNER_FIRST_CAPABILITY_RESPONSE_DRAFT.docx"
LOGO_PATH = ROOT / "assets" / "brand" / "lumaarc_arc_seal_v1.png"
LOGO_SHA256 = "1ed1c9b00e273aa9e781bd7fd0a4fcc3fc542257c6d294c8e8fbfada500701af"

NOTICE_ID = "ONC-ARGOS-SSN-2026-OS351107"
DEADLINE = "July 30, 2026 at 5:00 PM Eastern"
RESPONSE_TITLE = "Project Argos Capability Statement - Bounded Small-Business Response"

PRIVATE_FACT = "ACTION_TIME_PRIVATE_FACT_REQUIRED"
PRIVATE_FACT_DISPLAY = "Pending action-time fact"

INK = RGBColor(0x17, 0x22, 0x2B)
MUTED = RGBColor(0x52, 0x60, 0x6B)
ACCENT = RGBColor(0x00, 0x87, 0xB5)
LIGHT_FILL = "EAF5F8"
GRAY_FILL = "F1F3F5"
CAUTION_FILL = "FFF4CC"


EXECUTIVE_PARAGRAPHS = [
    (
        "Based on named first-party software artifacts, LumenCore can demonstrate "
        "source-artifact custody, versioned rule execution, evidence-package assembly, "
        "human approval gates, adverse-result retention, and reproducible reviewer "
        "handoff. Evidence packages are accompanied by SHA-256 integrity manifests. "
        "These are bounded component capabilities, not comparable federal-health past "
        "performance."
    ),
    (
        "For Project Argos, those patterns align most directly with requirements "
        "traceability, public-source ingestion provenance, deterministic validator "
        "orchestration, evidence-case generation, corrective-action retest records, "
        "and production handoff documentation. LumenCore is responding at component "
        "level and does not claim presently qualified full-scope health IT prime "
        "readiness, FHIR R4 conformance authority, HHS ATO leadership, or an evaluated "
        "health-domain deployment."
    ),
]

UNDERSTANDING_BULLETS = [
    (
        "Argos is a governed monitoring and evidence system, not a general-purpose "
        "chatbot. Standards-based and legal pass/fail decisions should remain "
        "deterministic wherever possible, with AI limited to bounded support such as "
        "summarization, classification, reconciliation, and drafting."
    ),
    (
        "The evidence chain begins with authoritative public sources and must preserve "
        "raw artifacts, request context, timestamps, response metadata, normalized "
        "records, lineage decisions, validation outputs, and hashes."
    ),
    (
        "Suspected non-conformities require human review, rule-level traceability, "
        "deduplication, plain-language narratives, controlled corrective-action support, "
        "and before/after retesting."
    ),
    (
        "Real-world endpoint observation must be safe and unauthenticated where "
        "appropriate, must avoid patient data and protected health information absent "
        "separate written authorization, and must operate within approved rate, scope, "
        "security, and hosting boundaries."
    ),
    (
        "Production readiness includes Government-aligned deployment, operations and "
        "rollback runbooks, cost and staffing visibility, and an HHS authorization path; "
        "a proof-of-concept result alone is not an Authority to Operate."
    ),
]

TASK_ROWS = [
    (
        "1",
        "Initiation, requirements, and regulatory traceability",
        "Partial capability",
        "Build a versioned requirements-to-test-to-evidence matrix with evidence IDs, "
        "severity, human checkpoints, schedules, and change control. No complete 45 CFR "
        "or CHPL mapping or authorized regulatory interpretation is claimed.",
    ),
    (
        "2",
        "Monthly administrative reporting",
        "Partial capability",
        "Produce source-backed status, blockers, milestones, event timelines, evidence "
        "IDs, decisions, and unresolved gates. Comparable HHS reporting performance is "
        "not claimed.",
    ),
    (
        "3",
        "Public-source discovery and ingestion",
        "Partial capability",
        "Implement source inventory, raw-response custody, timestamps, headers, hashes, "
        "lineage, precedence, dictionaries, and normalized projections. CHPL, Lantern, "
        "NPPES, and FHIR connectors are not demonstrated.",
    ),
    (
        "4",
        "Agentic architecture and governance",
        "Partial capability",
        "Separate deterministic validators from AI support; version policies, prompts, "
        "rules, retries, fallbacks, and audit logs; require human compliance decisions. "
        "No deployed health-domain architecture or secure HHS boundary is demonstrated.",
    ),
    (
        "5",
        "Publication and FHIR R4 conformance testing",
        "Not demonstrated",
        "Package supplied validator outputs as traceable evidence cases and regression "
        "receipts. No FHIR R4 test authority, quarterly operation, publication-failure "
        "detection, or conformance acceptance is claimed.",
    ),
    (
        "6",
        "Real-world endpoint observation",
        "Not demonstrated",
        "Contribute safe-observation logs, scope controls, drift records, hashes, and "
        "replay. Health-endpoint semantics, capability-sufficiency tests, "
        "published-versus-observed comparison, and approved access are not demonstrated.",
    ),
    (
        "7",
        "Evidence case generation and triage",
        "Strongest partial capability",
        "Assemble raw and normalized artifacts, logs, timestamps, hashes, severity, rule "
        "traceability, narratives, deduplication, screenshots, clustering, and "
        "human/machine exports. A designated HHS-system export remains unproven.",
    ),
    (
        "8",
        "Corrective-action workflow support",
        "Partial capability",
        "Provide controlled drafting, review, authorization, retest, before/after "
        "evidence, and recurrence fields. No ONC/ONC-ACB discretion or regulatory "
        "interpretation authority is claimed.",
    ),
    (
        "9",
        "HHS Authority to Operate",
        "Not demonstrated",
        "Supply component inventories, reproducibility and change receipts, evidence "
        "IDs, and control-test artifacts. No HHS ATO lead, FIPS 199/SSP delivery, 3PAO "
        "coordination, milestone reporting, or authorization package is claimed.",
    ),
    (
        "10",
        "Production-ready PPC release package",
        "Partial capability",
        "Package code, dependency locks, manifests, runbooks, rollback, limitations, "
        "tests, assumptions, costs, staffing, and reviewer receipts. No HHS-deployable "
        "package, hosting approval, or production acceptance is demonstrated.",
    ),
    (
        "11",
        "Public-artifact assessment and FY 2027 strategy",
        "Partial capability",
        "Apply traceability to approved public-document checks and locked prototype "
        "scoring. EHI/API criteria, controlling regulatory citations, and a validated "
        "FY 2027 strategy are not demonstrated.",
    ),
]

WORKSTREAM_STEPS = [
    (
        "Authorize sources and scope",
        "Record approved source, purpose, collection boundary, cadence, rate limits, "
        "and prohibited data before collection.",
    ),
    (
        "Preserve raw observations",
        "Store source payloads and request context with UTC timestamps, immutable "
        "identifiers, and SHA-256 manifests.",
    ),
    (
        "Normalize with lineage",
        "Create common records while retaining source precedence, conflicts, and "
        "reversible links to raw artifacts.",
    ),
    (
        "Run deterministic checks",
        "Apply versioned rules and validators; retain rule IDs, inputs, outputs, "
        "failures, and environment details.",
    ),
    (
        "Use AI only inside policy",
        "Permit bounded summaries, clustering, reconciliation, and draft narratives; "
        "block autonomous compliance determinations.",
    ),
    (
        "Assemble an evidence case",
        "Package artifacts, logs, hashes, timestamps, severity, traceability, narrative, "
        "and unresolved questions in human- and machine-readable forms.",
    ),
    (
        "Require human disposition",
        "Route suspected issues to authorized reviewers and record decisions, "
        "rationale, corrective-action drafts, and retest outcomes.",
    ),
]

EVIDENCE_ROWS = [
    (
        "First-party reproducibility package",
        "A named package and SHA-256 receipt record successful execution of its declared "
        "checks under an identified source state.",
        "First-party bounded reproducibility only; not external validation, agency "
        "certification, field performance, or health IT past performance.",
    ),
    (
        "Custody and validation controls",
        "Versioned manifests, SHA-256 receipts, schema checks, duplicate-action locks, "
        "and fail-closed gate records support an inspectable evidence workflow.",
        "Control artifacts do not establish health-domain correctness, production "
        "authorization, contract performance, or agency acceptance.",
    ),
    (
        "Adverse-result retention",
        "Named first-party records preserve failed promotion gates, negative findings, and "
        "unresolved authorities instead of converting them into favorable claims.",
        "Transparent failure handling is an engineering pattern, not proof of FHIR "
        "conformance, regulatory interpretation, or field performance.",
    ),
]

SIMILAR_SCOPE_ROWS = [
    (
        "Evidence custody, traceability, and deterministic validation",
        "Component pattern supported by named first-party code, tests, receipts, and "
        "reviewer artifacts.",
        "Available for authorized review after repository security screening.",
    ),
    (
        "FHIR R4, CHPL/Lantern/NPPES, and ONC Certification Program delivery",
        "No direct LumenCore prior-performance reference is claimed.",
        "Future acquisition staffing or separately authorized teaming would be required; "
        "no outside team is proposed in this response.",
    ),
    (
        "HHS ATO, FIPS 199, SSP/control implementation, and 3PAO coordination",
        "No direct LumenCore HHS authorization reference is claimed.",
        "Future acquisition staffing or separately authorized teaming would be required; "
        "no ATO lead or assessor is proposed in this response.",
    ),
    (
        "Full-scope federal health program integration",
        "No full-prime readiness or comparable federal health delivery is claimed.",
        "LumenCore should not be treated as a presently qualified full-scope prime.",
    ),
]

SECURITY_BULLETS = [
    (
        "Operate in a Government-approved environment with explicit authorization "
        "boundaries, least privilege, authenticated administration, dependency locking, "
        "secrets separation, immutable audit records, and rollback procedures."
    ),
    (
        "Treat HHS ATO as a managed authorization program: FIPS 199 categorization, "
        "boundary definition, SSP/control implementation, evidence collection, "
        "assessment coordination, POA&M handling, and Authorizing Official review."
    ),
    (
        "Do not collect, request, store, or use patient data or PHI unless separately "
        "authorized in writing and supported by approved architecture and controls."
    ),
    (
        "Keep AI components advisory and inspectable. Every model, prompt, policy, "
        "validator, retry, fallback, and rule version must be traceable to the evidence "
        "case it influenced."
    ),
]

CAPABILITY_GAP_ROWS = [
    (
        "Response posture",
        "LumenCore is the sole respondent. No teaming arrangement, joint venture, or "
        "subcontracting relationship is proposed.",
        "No external proposed team members exist to identify.",
    ),
    (
        "Health IT/FHIR and ONC program authority",
        "CHPL/Lantern/NPPES semantics, FHIR R4 Endpoint/Organization/Bundle testing, "
        "ONC regulatory mapping, corrective-action content review.",
        "Capability gap disclosed; no person or organization is represented.",
    ),
    (
        "Federal cybersecurity/ATO authority",
        "HHS authorization boundary, FIPS 199, SSP, control implementation, assessment "
        "coordination, POA&M and authorization package.",
        "Capability gap disclosed; no person or organization is represented.",
    ),
    (
        "Comparable full-program delivery",
        "Contract performance, staffing, Government coordination, hosting, delivery "
        "acceptance, health-domain prior performance, and integrated schedule.",
        "Not demonstrated and not claimed.",
    ),
]

MOBILIZATION_ROWS = [
    (
        "0-30 days",
        "Confirm scope, source authorization, regulatory ownership, hosting "
        "assumptions, traceability schema, security boundary, and pilot cohort.",
    ),
    (
        "31-90 days",
        "Stand up bounded ingestion, lineage, deterministic validator interfaces, "
        "evidence-case schema, human work queue, baseline test corpus, and audit logs.",
    ),
    (
        "91-180 days",
        "Expand approved endpoint observation, drift tracking, issue clustering, "
        "corrective-action draft/retest workflow, security documentation, and "
        "operational runbooks.",
    ),
    (
        "181 days and beyond",
        "Complete Government-directed hardening, assessment evidence, production "
        "handoff, PPC evaluation, cost/staffing analysis, and FY 2027 strategy.",
    ),
]

QUESTIONS = [
    "What proof-of-concept duration, target cohort size, and Government hosting environment should respondents assume for acquisition planning?",
    "Which organization owns final interpretation of 45 CFR and ONC Certification Program test outcomes, and what review service levels are expected?",
    "Will HHS provide an approved source list, rate-limit policy, synthetic test corpus, and expected evidence-case export schema?",
    "Which deliverables define the minimum viable first acquisition phase, and may HHS evaluate a bounded small-business contribution separately from full-scope prime readiness?",
    "What existing HHS authorization boundary, reusable controls, continuous-monitoring services, or 3PAO arrangements may be available to the PPC?",
]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_table_width(table, widths_in: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths_in) * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def set_page_number_start(section, start: int) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.0

    for style_name in ("Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.color.rgb = INK
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.0

    styles["Heading 1"].paragraph_format.space_before = Pt(12)
    styles["Heading 1"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].paragraph_format.space_before = Pt(9)
    styles["Heading 2"].paragraph_format.space_after = Pt(4)
    styles["Heading 3"].paragraph_format.space_before = Pt(6)
    styles["Heading 3"].paragraph_format.space_after = Pt(3)

    doc.core_properties.title = RESPONSE_TITLE
    doc.core_properties.subject = NOTICE_ID
    doc.core_properties.author = "LumenCore"
    doc.core_properties.keywords = "Project Argos, sources sought, evidence assurance"


def set_run(run, *, bold=None, italic=None, color=None, size=12) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc: Document, text: str, *, bold=False, italic=False, color=None, after=6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run(run, bold=bold, italic=italic, color=color)
    return paragraph


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.35)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run(run)


def add_numbered_item(doc: Document, label: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.35)
    paragraph.paragraph_format.first_line_indent = Inches(-0.35)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.line_spacing = 1.0
    first = paragraph.add_run(f"{label}. ")
    set_run(first, bold=True)
    rest = paragraph.add_run(text)
    set_run(rest)


def add_heading(doc: Document, text: str, level=1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run(run, bold=True, color=INK)


def fill_table(
    doc: Document,
    headers: list[str],
    rows: list[tuple[str, ...]],
    widths: list[float],
    *,
    header_fill=LIGHT_FILL,
    font_size=12,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    set_row_cant_split(header)
    for idx, label in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(label)
        set_run(run, bold=True, size=font_size)
    for row_data in rows:
        row = table.add_row()
        set_row_cant_split(row)
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(value)
            set_run(run, size=font_size)
        if len(table.rows) % 2 == 0:
            for cell in row.cells:
                set_cell_shading(cell, "F9FBFC")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_cover(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run()
    run.add_picture(str(LOGO_PATH), width=Inches(1.7))

    for text, bold, color, after in (
        ("LUMENCORE", True, INK, 2),
        ("LUMAARC BRAND MARK", False, MUTED, 4),
        ("MARKET RESEARCH CAPABILITY STATEMENT", True, INK, 4),
        ("Project Argos Agentic AI Proof of Concept", True, INK, 4),
        (f"Sources Sought Notice {NOTICE_ID}", False, MUTED, 18),
        ("DRAFT - HUMAN REVIEW AND ACTION-TIME FACTS REQUIRED", True, RGBColor(0x8A, 0x51, 0x00), 18),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(text)
        set_run(r, bold=bold, color=color)

    rows = [
        ("Company name", "LumenCore"),
        ("Responding legal entity", PRIVATE_FACT_DISPLAY),
        ("UEI / DUNS if applicable", PRIVATE_FACT_DISPLAY),
        ("Company address", PRIVATE_FACT_DISPLAY),
        ("Authorized point of contact", PRIVATE_FACT_DISPLAY),
        ("Telephone / email", PRIVATE_FACT_DISPLAY),
        ("Small-business designation(s)", PRIVATE_FACT_DISPLAY),
        ("Response deadline", DEADLINE),
    ]
    fill_table(doc, ["Required cover field", "Current response"], rows, [2.05, 4.45], header_fill=GRAY_FILL)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(8)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.line_spacing = 1.0
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "Market research response only. No proprietary, classified, confidential, "
        "CUI, patient, or sensitive information is included."
    )
    set_run(run, italic=True, color=MUTED)


def add_content_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run(f"LumenCore | Project Argos | {NOTICE_ID} | Draft")
    set_run(run, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_after = Pt(0)
    run = footer_p.add_run("Content Page ")
    set_run(run, color=MUTED)
    add_page_field(footer_p)


def build_docx() -> None:
    if not LOGO_PATH.exists():
        raise FileNotFoundError(LOGO_PATH)
    actual_logo_hash = sha256(LOGO_PATH)
    if actual_logo_hash != LOGO_SHA256:
        raise RuntimeError(f"Logo hash mismatch: {actual_logo_hash}")

    OUTPUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    content_section = doc.add_section(WD_SECTION.NEW_PAGE)
    content_section.page_width = Inches(8.5)
    content_section.page_height = Inches(11)
    content_section.top_margin = Inches(1)
    content_section.right_margin = Inches(1)
    content_section.bottom_margin = Inches(1)
    content_section.left_margin = Inches(1)
    content_section.header_distance = Inches(0.42)
    content_section.footer_distance = Inches(0.42)
    set_page_number_start(content_section, 1)
    add_content_header_footer(content_section)

    add_heading(doc, "1. Executive Fit and Recommended Role", level=1)
    for paragraph in EXECUTIVE_PARAGRAPHS:
        add_paragraph(doc, paragraph)

    callout = doc.add_table(rows=1, cols=1)
    callout.style = "Table Grid"
    set_table_width(callout, [6.5])
    cell = callout.cell(0, 0)
    set_cell_shading(cell, CAUTION_FILL)
    set_cell_margins(cell, top=120, bottom=120)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        "Response posture: bounded small-business component capability. LumenCore is "
        "the sole respondent; no outside team is proposed; full-prime readiness is not "
        "claimed."
    )
    set_run(r, bold=True)

    add_heading(doc, "2. Understanding of the Requirement", level=1)
    for item in UNDERSTANDING_BULLETS:
        add_bullet(doc, item)

    add_heading(doc, "3. Task-by-Task Capability and Gap Matrix", level=1)
    add_paragraph(
        doc,
        "The matrix distinguishes current first-party component capability from work "
        "that has not been demonstrated. It does not assign any task to an unnamed "
        "partner or represent another organization's authority.",
    )
    fill_table(
        doc,
        ["Task", "Requirement", "Proposed position", "Bounded contribution"],
        TASK_ROWS,
        [0.42, 1.43, 1.40, 3.25],
    )

    add_heading(doc, "4. Proposed Evidence-Assurance Workstream", level=1)
    add_paragraph(
        doc,
        "LumenCore proposes a fail-closed sequence in which missing authority, provenance, "
        "validation, or human disposition remains visible rather than becoming a "
        "promoted compliance claim.",
    )
    for idx, (label, text) in enumerate(WORKSTREAM_STEPS, start=1):
        add_numbered_item(doc, str(idx), f"{label}: {text}")

    add_heading(doc, "5. Component Evidence and Similar-Scope Status", level=1)
    add_paragraph(
        doc,
        "The following named first-party records support the engineering patterns "
        "offered here and are available for authorized review after repository security "
        "screening. They are not presented as healthcare, agency, field, or economic "
        "validation.",
    )
    fill_table(
        doc,
        ["Evidence record", "What the record supports", "What it does not support"],
        EVIDENCE_ROWS,
        [1.35, 2.75, 2.40],
    )
    add_paragraph(
        doc,
        "The notice requests experience of similar scope and complexity. The matrix "
        "below distinguishes LumenCore component evidence from material qualifications "
        "that are not demonstrated; it does not substitute adjacent technical work for "
        "federal health prior performance.",
    )
    fill_table(
        doc,
        ["Capability area", "Present support", "Acquisition implication"],
        SIMILAR_SCOPE_ROWS,
        [1.75, 2.30, 2.45],
    )
    add_paragraph(
        doc,
        "The internal evidence index states the present boundary directly: current "
        "strengths are artifact custody, deterministic replay, fail-closed claim "
        "governance, reviewer handoff packaging, and bounded pilot design. Independent "
        "scientific validation, field-validated savings, agency endorsement, certified "
        "safety, audited revenue, and customer adoption require separately identified "
        "external records.",
        italic=True,
    )

    add_heading(doc, "6. Delivery and Security Approach", level=1)
    for item in SECURITY_BULLETS:
        add_bullet(doc, item)
    add_paragraph(
        doc,
        "Current boundary: the cited first-party artifacts are not an HHS-authorized "
        "production system and do not establish an ATO, FHIR certification, 3PAO "
        "assessment, or federal health deployment. Those remain acquisition and "
        "authorization gaps.",
        bold=True,
    )

    add_heading(doc, "7. Standalone Response and Capability Gaps", level=1)
    fill_table(
        doc,
        ["Area", "Required or relevant capability", "Current status"],
        CAPABILITY_GAP_ROWS,
        [1.65, 3.25, 1.60],
    )
    add_paragraph(
        doc,
        "Teaming status as of submission: LumenCore has not formed or been authorized "
        "to represent a team for this response. No other organization or individual is "
        "proposed as a team member. The rows above identify capability gaps, not "
        "proposed or committed team members. LumenCore is open to future acquisition or "
        "teaming discussions, but this response does not represent another "
        "organization's commitment, qualifications, past performance, or authority.",
        italic=True,
    )

    add_heading(doc, "8. Illustrative Mobilization Plan", level=1)
    add_paragraph(
        doc,
        "The schedule below is an acquisition-planning outline, not a binding offer. It "
        "must be reconciled with Government scope, hosting, staffing, review cadence, "
        "security boundary, and the SOW's authorization milestones.",
    )
    fill_table(doc, ["Period", "Illustrative outcome"], MOBILIZATION_ROWS, [1.25, 5.25])

    add_heading(doc, "9. Questions and Requested Next Step", level=1)
    for idx, question in enumerate(QUESTIONS, start=1):
        add_numbered_item(doc, str(idx), question)
    add_paragraph(
        doc,
        "For acquisition planning, HHS may treat this response as evidence of bounded "
        "small-business capacity in evidence management and validation orchestration. "
        "If HHS conducts follow-up market research, LumenCore is available to "
        "demonstrate the cited artifacts and answer technical questions.",
        bold=True,
    )
    add_paragraph(
        doc,
        "This response is for market research only and is not an offer, proposal, "
        "certification, or representation that every draft SOW task is presently covered.",
        italic=True,
        color=MUTED,
    )

    doc.save(DOCX_PATH)


def table_markdown(headers: list[str], rows: list[tuple[str, ...]]) -> list[str]:
    output = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    for row in rows:
        cleaned = [value.replace("|", "\\|").replace("\n", " ") for value in row]
        output.append("| " + " | ".join(cleaned) + " |")
    return output


def build_markdown() -> None:
    lines = [
        f"# {RESPONSE_TITLE}",
        "",
        f"**Notice:** `{NOTICE_ID}`",
        f"**Deadline:** {DEADLINE}",
        "**Status:** `DRAFT - HUMAN REVIEW AND ACTION-TIME FACTS REQUIRED`",
        "",
        f"![Founder-selected LumaArc brand mark](../../assets/brand/lumaarc_arc_seal_v1.png)",
        "",
        "## Required Cover Facts",
        "",
    ]
    cover_rows = [
        ("Company name", "LumenCore"),
        ("Responding legal entity", PRIVATE_FACT),
        ("UEI / DUNS if applicable", PRIVATE_FACT),
        ("Company address", PRIVATE_FACT),
        ("Authorized point of contact", PRIVATE_FACT),
        ("Telephone / email", PRIVATE_FACT),
        ("Small-business designation(s)", PRIVATE_FACT),
    ]
    lines.extend(table_markdown(["Field", "Current response"], cover_rows))
    lines.extend(["", "## 1. Executive Fit and Recommended Role", ""])
    lines.extend(EXECUTIVE_PARAGRAPHS)
    lines.extend(
        [
            "",
            "**Response posture:** bounded small-business component capability. "
            "LumenCore is the sole respondent; no outside team is proposed; full-prime "
            "readiness is not claimed.",
            "",
            "## 2. Understanding of the Requirement",
            "",
        ]
    )
    lines.extend(f"- {item}" for item in UNDERSTANDING_BULLETS)
    lines.extend(
        [
            "",
            "## 3. Task-by-Task Capability and Gap Matrix",
            "",
            "This matrix distinguishes current first-party component capability from "
            "work that has not been demonstrated. It does not assign any task to an "
            "unnamed partner or represent another organization's authority.",
            "",
        ]
    )
    lines.extend(
        table_markdown(
            ["Task", "Requirement", "Proposed position", "Bounded contribution"],
            TASK_ROWS,
        )
    )
    lines.extend(["", "## 4. Proposed Evidence-Assurance Workstream", ""])
    lines.extend(
        f"{idx}. **{label}.** {text}"
        for idx, (label, text) in enumerate(WORKSTREAM_STEPS, start=1)
    )
    lines.extend(
        [
            "",
            "## 5. Component Evidence and Similar-Scope Status",
            "",
            "The named first-party records below are available for authorized review "
            "after repository security screening. They are not presented as healthcare, "
            "agency, field, or economic validation.",
            "",
        ]
    )
    lines.extend(
        table_markdown(
            ["Evidence record", "What it supports", "What it does not support"],
            EVIDENCE_ROWS,
        )
    )
    lines.extend(
        [
            "",
            "The notice requests experience of similar scope and complexity. The matrix "
            "below distinguishes LumenCore component evidence from material "
            "qualifications that are not demonstrated; adjacent technical work is not "
            "treated as federal health prior performance.",
            "",
        ]
    )
    lines.extend(
        table_markdown(
            ["Capability area", "Present support", "Acquisition implication"],
            SIMILAR_SCOPE_ROWS,
        )
    )
    lines.extend(["", "## 6. Delivery and Security Approach", ""])
    lines.extend(f"- {item}" for item in SECURITY_BULLETS)
    lines.extend(
        [
            "",
            "**Current boundary:** the cited first-party artifacts are not an "
            "HHS-authorized production system and do not establish an ATO, FHIR "
            "certification, 3PAO assessment, or federal health deployment.",
            "",
            "## 7. Standalone Response and Capability Gaps",
            "",
        ]
    )
    lines.extend(
        table_markdown(
            ["Area", "Required or relevant capability", "Current status"],
            CAPABILITY_GAP_ROWS,
        )
    )
    lines.extend(
        [
            "",
            "*Teaming status as of submission: LumenCore has not formed or been "
            "authorized to represent a team for this response. No other organization "
            "or individual is proposed as a team member. The rows above identify "
            "capability gaps, not proposed or committed team members. LumenCore is open "
            "to future acquisition or teaming discussions, but this response does not "
            "represent another organization's commitment, qualifications, past "
            "performance, or authority.*",
            "",
        ]
    )
    lines.extend(["", "## 8. Illustrative Mobilization Plan", ""])
    lines.extend(table_markdown(["Period", "Illustrative outcome"], MOBILIZATION_ROWS))
    lines.extend(["", "## 9. Questions and Requested Next Step", ""])
    lines.extend(f"{idx}. {question}" for idx, question in enumerate(QUESTIONS, start=1))
    lines.extend(
        [
            "",
            "**Requested next step:** for acquisition planning, HHS may treat this "
            "response as evidence of bounded small-business capacity in evidence "
            "management and validation orchestration. If HHS conducts follow-up market "
            "research, LumenCore is available to demonstrate the cited artifacts and "
            "answer technical questions.",
            "",
            "> This response is for market research only and is not an offer, proposal, "
            "certification, or representation that every draft SOW task is presently covered.",
            "",
        ]
    )
    markdown = "\n\n".join(lines).rstrip() + "\n"
    MARKDOWN_PATH.write_bytes(markdown.encode("utf-8"))


def main() -> int:
    build_markdown()
    build_docx()
    receipt = {
        "schema": "lumencore.argos_response_build_receipt.v1",
        "notice_id": NOTICE_ID,
        "status": "DRAFT_REVIEW_READY_NOT_SEND_READY",
        "logo_sha256": sha256(LOGO_PATH),
        "markdown": {
            "path": str(MARKDOWN_PATH.relative_to(ROOT)).replace("\\", "/"),
            "sha256": sha256(MARKDOWN_PATH),
        },
        "docx": {
            "path": str(DOCX_PATH.relative_to(ROOT)).replace("\\", "/"),
            "sha256": sha256(DOCX_PATH),
        },
        "submission_authorized": False,
        "exact_action_time_human_approval_required": True,
    }
    (OUTPUT / "build_receipt.json").write_text(
        json.dumps(receipt, indent=2) + "\n",
        encoding="utf-8",
        newline="\n",
    )
    print(json.dumps(receipt, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
