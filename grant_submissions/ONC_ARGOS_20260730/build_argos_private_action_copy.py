from __future__ import annotations

import argparse
import hashlib
import importlib.util
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ARGOS_DIR = Path(__file__).resolve().parent
ROOT = ARGOS_DIR.parents[1]
PUBLIC_MARKDOWN = ARGOS_DIR / "ARGOS_PARTNER_FIRST_CAPABILITY_RESPONSE_DRAFT.md"
PUBLIC_DOCX = (
    ARGOS_DIR / "output" / "ARGOS_PARTNER_FIRST_CAPABILITY_RESPONSE_DRAFT.docx"
)
SUBMISSION_GATE = ARGOS_DIR / "ARGOS_SUBMISSION_GATE_2026-07-26.json"
TEAM_REGISTER = ARGOS_DIR / "ARGOS_TEAMING_CANDIDATE_REGISTER_2026-07-27.json"
SECURITY_GATE = ARGOS_DIR / "ARGOS_PUBLIC_REPOSITORY_SECURITY_GATE_2026-07-28.json"
SECURITY_VERIFIER = (
    ROOT / "code" / "ops" / "VERIFY_PUBLIC_REPO_CREDENTIAL_HYGIENE.py"
)

NOTICE_ID = "ONC-ARGOS-SSN-2026-OS351107"
FACT_SCHEMA = "lumencore.argos_private_facts.v1"
RECEIPT_SCHEMA = "lumencore.argos_private_action_copy_receipt.v2"
RECEIPT_DECISION = "PRIVATE_COVER_READY_STANDALONE_DISPATCH_BLOCKED"
PRIVATE_MARKER = "ACTION_TIME_PRIVATE_FACT_REQUIRED"
PRIVATE_DISPLAY_MARKER = "Pending action-time fact"
PRIVATE_STATUS = "PRIVATE ACTION COPY - HUMAN REVIEW REQUIRED"
PUBLIC_REPOSITORY_URL = (
    "https://github.com/robertashworth1986-debug/lumen-core-public"
)
PUBLIC_SITE_URL = "https://lumen-core.ai/"
FORBIDDEN_EXTERNAL_ROUTE_TOKENS = (
    PUBLIC_REPOSITORY_URL,
    PUBLIC_SITE_URL,
    "github.com",
    "lumen-core-public",
    "robertashworth1986-debug",
)
TEXT_OUTPUT_NAME = "ARGOS_PRIVATE_ACTION_COPY.md"
DOCX_OUTPUT_NAME = "ARGOS_PRIVATE_ACTION_COPY.docx"
RECEIPT_OUTPUT_NAME = "ARGOS_PRIVATE_ACTION_COPY_RECEIPT.json"
PUBLIC_VAULT_ROOTS = (Path(r"E:\LumaProofVault"),)
CLAIM_BOUNDARY = (
    "This receipt proves only that a private action copy was generated from a "
    "schema-valid, user-attested facts file without changing the public templates. "
    "It also proves that the generated Markdown and DOCX contain no known repository "
    "or live-site route. It does not prove final dispatch verification, submission, "
    "acceptance, selection, award, compliance, certification, or authorization."
)
SAFEST_NEXT_ACTION = (
    "Keep the private copy outside Git and public mirrors. Review every inserted "
    "fact and the rendered cover with the user, preserve the link-free attachment "
    "boundary, then run the Government duplicate and final-dispatch gates before "
    "requesting single-use action-time approval. Credential rotation and public-history "
    "remediation remain required before any public repository promotion."
)

REQUIRED_FACT_KEYS = (
    "legal_company_name",
    "uei",
    "duns_if_notice_or_entity_record_requires_it",
    "company_address",
    "authorized_point_of_contact_name_and_title",
    "authorized_point_of_contact_phone",
    "authorized_point_of_contact_email",
    "small_business_designations",
    "sam_registration_status_and_expiration",
)
NOT_APPLICABLE_ALLOWED = {"duns_if_notice_or_entity_record_requires_it"}
SOURCE_KINDS = {
    "OFFICIAL_ENTITY_RECORD",
    "OFFICIAL_SAM_RECORD",
    "FOUNDER_VERIFIED_BUSINESS_RECORD",
    "OFFICIAL_CORRESPONDENCE",
}
PROHIBITED_KEYS = {
    "ein",
    "tin",
    "ssn",
    "bank_account",
    "routing_number",
    "password",
    "credential",
    "api_key",
    "otp",
}
MAX_VALUE_LENGTHS = {
    "legal_company_name": 120,
    "uei": 12,
    "duns_if_notice_or_entity_record_requires_it": 14,
    "company_address": 180,
    "authorized_point_of_contact_name_and_title": 120,
    "authorized_point_of_contact_phone": 40,
    "authorized_point_of_contact_email": 254,
    "small_business_designations": 160,
    "sam_registration_status_and_expiration": 120,
}


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def parse_utc(value: str) -> datetime:
    parsed = datetime.fromisoformat(value.replace("Z", "+00:00"))
    if parsed.tzinfo is None:
        raise ValueError("Timestamp must include a timezone")
    return parsed.astimezone(timezone.utc)


def is_within(path: Path, parent: Path) -> bool:
    try:
        path.resolve().relative_to(parent.resolve())
        return True
    except ValueError:
        return False


def public_storage_roots() -> tuple[Path, ...]:
    configured = tuple(root for root in PUBLIC_VAULT_ROOTS if root.is_absolute())
    return (ROOT, *configured)


def validate_private_path(path: Path, label: str) -> Path:
    resolved = path.expanduser().resolve()
    if any(is_within(resolved, public_root) for public_root in public_storage_roots()):
        raise ValueError(
            f"{label} must be outside the public repository and public mirror vaults"
        )
    return resolved


def _reject_duplicate_keys(pairs: list[tuple[str, object]]) -> dict:
    result = {}
    for key, value in pairs:
        if key in result:
            raise ValueError(f"duplicate JSON key: {key}")
        result[key] = value
    return result


def read_json(path: Path) -> dict:
    payload = json.loads(
        path.read_text(encoding="utf-8"),
        object_pairs_hook=_reject_duplicate_keys,
    )
    if not isinstance(payload, dict):
        raise ValueError(f"{path.name} must contain a JSON object")
    return payload


def validate_security_gate(payload: dict) -> dict[str, bool]:
    spec = importlib.util.spec_from_file_location(
        "argos_public_credential_hygiene",
        SECURITY_VERIFIER,
    )
    if spec is None or spec.loader is None:
        raise ValueError("public repository security verifier cannot be loaded")
    verifier = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(verifier)
    current_payload = verifier.build_payload()
    if verifier.canonical_json_bytes(payload) != verifier.canonical_json_bytes(
        current_payload
    ):
        raise ValueError(
            "public repository security gate is stale or was edited outside "
            "the canonical verifier"
        )
    link_allowed = current_payload["public_repository_link_allowed"]
    sanitized_allowed = current_payload["sanitized_external_response_allowed"]
    send_allowed = current_payload[
        "final_argos_send_allowed_by_security_gate"
    ]
    if send_allowed is not (link_allowed or sanitized_allowed):
        raise ValueError("public repository security decision is inconsistent")
    if link_allowed:
        provider_rotations_clear = all(
            row["confirmed"]
            for row in current_payload["provider_rotation"].values()
        )
        history = current_payload["history"]
        if not (
            current_payload["decision"]
            == "PASS_TARGETED_CREDENTIAL_AND_REMOTE_HISTORY_GATE"
            and provider_rotations_clear
            and history["remediation_confirmed"] is True
            and history["remote_public_history_verification_confirmed"] is True
            and history["historical_exposure_detected"] is False
            and history["scan_complete"] is True
            and history["scan_failure_count"] == 0
        ):
            raise ValueError(
                "public repository security gate cannot clear the repository link"
            )
    elif sanitized_allowed:
        history = current_payload["history"]
        if not (
            current_payload["decision"]
            == "ALLOW_SANITIZED_EXTERNAL_RESPONSE_BLOCK_PUBLIC_REPO_LINK"
            and current_payload["current_file"]["placeholder_only"] is True
            and current_payload["current_file"][
                "required_environment_references_present"
            ]
            is True
            and history["scan_complete"] is True
            and history["scan_failure_count"] == 0
        ):
            raise ValueError(
                "public repository security gate cannot clear the sanitized path"
            )
    if not send_allowed:
        raise ValueError("repository security preconditions block external response")
    return {
        "public_repository_link_allowed": link_allowed,
        "sanitized_external_response_allowed": sanitized_allowed,
        "final_argos_send_allowed_by_security_gate": send_allowed,
    }


def assert_no_external_repository_route(text: str, label: str) -> None:
    lowered = text.lower()
    found = [
        token
        for token in FORBIDDEN_EXTERNAL_ROUTE_TOKENS
        if token.lower() in lowered
    ]
    if found:
        raise ValueError(f"{label} contains a blocked external route")


def assert_docx_repo_isolated(path: Path) -> None:
    with ZipFile(path) as archive:
        for member in archive.infolist():
            data = archive.read(member)
            lowered = data.lower()
            if any(
                token.lower().encode("utf-8") in lowered
                for token in FORBIDDEN_EXTERNAL_ROUTE_TOKENS
            ):
                raise ValueError(
                    f"private DOCX contains a blocked external route in {member.filename}"
                )
            if member.filename.endswith(".rels") and b'targetmode="external"' in lowered:
                raise ValueError(
                    f"private DOCX contains an external relationship in {member.filename}"
                )


def validate_fact_value(key: str, fact: dict, evaluated: datetime) -> None:
    if set(fact) != {"value", "status", "source_kind", "verified_utc"}:
        raise ValueError(f"{key} must contain only value, status, source_kind, verified_utc")
    status = fact["status"]
    value = fact["value"]
    if status == "NOT_APPLICABLE":
        if key not in NOT_APPLICABLE_ALLOWED or value != "NOT_APPLICABLE":
            raise ValueError(f"{key} cannot be marked NOT_APPLICABLE")
    elif status != "VERIFIED":
        raise ValueError(f"{key} must be VERIFIED")
    if not isinstance(value, str) or not value.strip():
        raise ValueError(f"{key} must have a non-empty value")
    if len(value) > MAX_VALUE_LENGTHS[key]:
        raise ValueError(f"{key} exceeds the private-cover length limit")
    if any(ord(character) < 32 and character not in "\r\n\t" for character in value):
        raise ValueError(f"{key} contains a prohibited control character")
    if PRIVATE_MARKER in value or PRIVATE_DISPLAY_MARKER in value:
        raise ValueError(f"{key} still contains a placeholder")
    if fact["source_kind"] not in SOURCE_KINDS:
        raise ValueError(f"{key} has an unsupported source_kind")
    verified = parse_utc(fact["verified_utc"])
    if verified > evaluated:
        raise ValueError(f"{key} has a future verification timestamp")

    if key == "uei" and not re.fullmatch(r"[A-Za-z0-9]{12}", value):
        raise ValueError("uei must contain exactly 12 alphanumeric characters")
    if (
        key == "duns_if_notice_or_entity_record_requires_it"
        and status == "VERIFIED"
        and not re.fullmatch(r"\d{9}", value)
    ):
        raise ValueError("a verified DUNS must contain exactly 9 digits")
    if key == "authorized_point_of_contact_email" and not re.fullmatch(
        r"[^@\s]+@[^@\s]+\.[^@\s]+", value
    ):
        raise ValueError("authorized point-of-contact email is malformed")
    if key == "authorized_point_of_contact_phone":
        digits = re.sub(r"\D", "", value)
        if not 10 <= len(digits) <= 15:
            raise ValueError("authorized point-of-contact phone is malformed")


def validate_facts(payload: dict, evaluated: datetime) -> dict:
    expected_top_level = {
        "schema",
        "notice_id",
        "facts",
        "assertions",
    }
    if set(payload) != expected_top_level:
        raise ValueError("private facts file has unexpected top-level fields")
    if payload["schema"] != FACT_SCHEMA:
        raise ValueError("private facts schema is not supported")
    if payload["notice_id"] != NOTICE_ID:
        raise ValueError("private facts notice_id does not match Project Argos")

    facts = payload["facts"]
    if not isinstance(facts, dict):
        raise ValueError("facts must be an object")
    unexpected = set(facts) - set(REQUIRED_FACT_KEYS)
    prohibited = {key for key in unexpected if key.lower() in PROHIBITED_KEYS}
    if prohibited:
        raise ValueError("private facts file contains prohibited sensitive fields")
    if unexpected:
        raise ValueError("private facts file contains fields not required by this response")
    missing = set(REQUIRED_FACT_KEYS) - set(facts)
    if missing:
        raise ValueError("private facts file is missing required fields")
    for key in REQUIRED_FACT_KEYS:
        if not isinstance(facts[key], dict):
            raise ValueError(f"{key} must be an object")
        validate_fact_value(key, facts[key], evaluated)

    assertions = payload["assertions"]
    if set(assertions) != {
        "facts_current_and_accurate",
        "authorized_for_this_response",
        "minimum_necessary_business_information_only",
    }:
        raise ValueError("private facts assertions are incomplete")
    if not all(value is True for value in assertions.values()):
        raise ValueError("every private facts assertion must be explicitly true")
    return facts


def display_values(facts: dict) -> dict[str, str]:
    value = lambda key: facts[key]["value"].strip()
    duns = value("duns_if_notice_or_entity_record_requires_it")
    if facts["duns_if_notice_or_entity_record_requires_it"]["status"] == "NOT_APPLICABLE":
        duns = "Not applicable"
    return {
        "Responding legal entity": value("legal_company_name"),
        "UEI / DUNS if applicable": f"UEI: {value('uei')}; DUNS: {duns}",
        "Company address": value("company_address"),
        "Authorized point of contact": value(
            "authorized_point_of_contact_name_and_title"
        ),
        "Telephone / email": (
            f"{value('authorized_point_of_contact_phone')} / "
            f"{value('authorized_point_of_contact_email')}"
        ),
        "Small-business designation(s)": (
            f"{value('small_business_designations')}; "
            f"SAM: {value('sam_registration_status_and_expiration')}"
        ),
    }


def markdown_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace("|", "\\|").replace("\r", " ").replace(
        "\n", " "
    )


def render_private_markdown(
    template: str,
    values: dict[str, str],
) -> str:
    output = template.replace(
        "**Status:** `DRAFT - HUMAN REVIEW AND ACTION-TIME FACTS REQUIRED`",
        f"**Status:** `{PRIVATE_STATUS}`",
    )
    lines = []
    replaced = set()
    for line in output.splitlines():
        if line.startswith("|") and line.endswith("|"):
            parts = [part.strip() for part in line[1:-1].split("|")]
            if parts and parts[0] in values:
                label = parts[0]
                line = f"| {label} | {markdown_escape(values[label])} |"
                replaced.add(label)
        lines.append(line)
    if replaced != set(values):
        raise ValueError("public Markdown cover fields do not match the finalizer map")
    output = "\n".join(lines).rstrip() + "\n"
    if PRIVATE_MARKER in output or PRIVATE_DISPLAY_MARKER in output:
        raise ValueError("private Markdown still contains a private-fact placeholder")
    assert_no_external_repository_route(output, "private Markdown")
    handling = (
        "\n**Handling:** Minimum necessary verified business registration and contact "
        "facts only. No tax, banking, credential, patient, CUI, classified, or "
        "patent-sensitive information.\n"
    )
    status_line = f"**Status:** `{PRIVATE_STATUS}`\n"
    return output.replace(status_line, status_line + handling, 1)


def style_private_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(8.2)


def replace_paragraph_text(
    paragraph,
    text: str,
    *,
    bold: bool,
    italic: bool = False,
    size: float = 10,
) -> None:
    paragraph.text = text
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic


def render_private_docx(
    template_path: Path,
    output_path: Path,
    values: dict[str, str],
) -> None:
    doc = Document(template_path)
    replaced = set()
    for table in doc.tables:
        for row in table.rows:
            label = row.cells[0].text.strip()
            if label in values:
                row.cells[1].text = values[label]
                style_private_cell(row.cells[1])
                replaced.add(label)
    if replaced != set(values):
        raise ValueError("public DOCX cover fields do not match the finalizer map")

    status_replaced = False
    note_replaced = False
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == (
            "DRAFT - HUMAN REVIEW AND ACTION-TIME FACTS REQUIRED"
        ):
            replace_paragraph_text(paragraph, PRIVATE_STATUS, bold=True)
            status_replaced = True
        if paragraph.text.strip() == (
            "Market research response only. No proprietary, classified, confidential, "
            "CUI, patient, or sensitive information is included."
        ):
            replace_paragraph_text(
                paragraph,
                "Market research response only. Contains the minimum necessary verified "
                "business contact and registration facts; no tax, banking, credential, "
                "classified, CUI, patient, or patent-sensitive information is included.",
                bold=False,
                italic=True,
                size=8.2,
            )
            note_replaced = True
    if not status_replaced or not note_replaced:
        raise ValueError("public DOCX status or handling notice could not be replaced")

    doc.core_properties.title = "Project Argos Private Action Copy"
    doc.core_properties.subject = (
        "Private human-review copy; not evidence of submission or authorization"
    )
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                if "Draft" in run.text:
                    run.text = run.text.replace("Draft", "Private action copy")
    doc.save(output_path)
    assert_docx_repo_isolated(output_path)
    with ZipFile(output_path) as archive:
        document_xml = archive.read("word/document.xml")
    if (
        PRIVATE_MARKER.encode() in document_xml
        or PRIVATE_DISPLAY_MARKER.encode() in document_xml
    ):
        raise ValueError("private DOCX still contains a private-fact placeholder")


def output_paths(output_dir: Path) -> tuple[Path, Path, Path]:
    return (
        output_dir / TEXT_OUTPUT_NAME,
        output_dir / DOCX_OUTPUT_NAME,
        output_dir / RECEIPT_OUTPUT_NAME,
    )


def require_exact_keys(value: object, expected: set[str], label: str) -> dict:
    if not isinstance(value, dict):
        raise ValueError(f"{label} must be an object")
    actual = set(value)
    if actual != expected:
        missing = sorted(expected - actual)
        unexpected = sorted(actual - expected)
        raise ValueError(
            f"{label} keys are invalid; missing={missing}; unexpected={unexpected}"
        )
    return value


def expected_field_states(facts: dict) -> dict:
    return {
        key: {
            "status": facts[key]["status"],
            "source_kind": facts[key]["source_kind"],
            "verified_utc": parse_utc(facts[key]["verified_utc"]).strftime(
                "%Y-%m-%dT%H:%M:%SZ"
            ),
        }
        for key in REQUIRED_FACT_KEYS
    }


def reject_private_receipt_leaks(
    receipt: dict,
    facts: dict,
    private_facts_path: Path,
    private_output_dir: Path,
) -> None:
    receipt_text = json.dumps(receipt, sort_keys=True)
    raw_values = {
        fact["value"].strip()
        for fact in facts.values()
        if fact["value"].strip() != "NOT_APPLICABLE"
    }
    private_fragments = raw_values | set(display_values(facts).values())
    if any(fragment and fragment in receipt_text for fragment in private_fragments):
        raise ValueError("private value leaked into the redacted receipt")
    if any(
        str(private_path) in receipt_text
        for private_path in (private_facts_path, private_output_dir)
    ):
        raise ValueError("private path leaked into the redacted receipt")


def validate_private_receipt(
    receipt: dict,
    facts: dict,
    private_facts_path: Path,
    private_output_dir: Path,
    markdown_path: Path,
    docx_path: Path,
    evaluated: datetime,
) -> None:
    require_exact_keys(
        receipt,
        {
            "schema",
            "evaluated_utc",
            "notice_id",
            "decision",
            "required_fact_count",
            "private_value_count",
            "placeholder_count",
            "field_states",
            "facts_file",
            "public_templates",
            "public_repository_security",
            "outputs",
            "response_mode",
            "team_disclosure_resolved",
            "candidate_name_authorization_count",
            "government_send_ready",
            "submission_authorized",
            "external_action_performed",
            "private_values_logged",
            "private_output_mirrored_to_public_vault",
            "claim_boundary",
            "safest_next_action",
        },
        "private receipt",
    )
    reject_private_receipt_leaks(
        receipt,
        facts,
        private_facts_path,
        private_output_dir,
    )

    if receipt["schema"] != RECEIPT_SCHEMA:
        raise ValueError("private receipt schema is not supported")
    if receipt["notice_id"] != NOTICE_ID:
        raise ValueError("private receipt notice_id does not match Project Argos")
    if receipt["decision"] != RECEIPT_DECISION:
        raise ValueError("private receipt decision is not fail-closed")
    if parse_utc(receipt["evaluated_utc"]) > evaluated:
        raise ValueError("private receipt has a future evaluation timestamp")
    if receipt["required_fact_count"] != len(REQUIRED_FACT_KEYS):
        raise ValueError("private receipt required-fact count is stale")
    if receipt["private_value_count"] != len(facts):
        raise ValueError("private receipt value count is stale")
    if receipt["placeholder_count"] != 0:
        raise ValueError("private receipt reports unresolved placeholders")
    if receipt["field_states"] != expected_field_states(facts):
        raise ValueError("private receipt field states are stale")

    facts_file = require_exact_keys(
        receipt["facts_file"],
        {"bytes", "sha256", "path_logged"},
        "private receipt facts_file",
    )
    if facts_file["bytes"] != private_facts_path.stat().st_size:
        raise ValueError("private facts byte count is stale")
    if facts_file["sha256"] != sha256(private_facts_path):
        raise ValueError("private facts custody hash is stale")
    if facts_file["path_logged"] is not False:
        raise ValueError("private receipt claims a private path was logged")

    public_templates = require_exact_keys(
        receipt["public_templates"],
        {"markdown_sha256", "docx_sha256", "unchanged"},
        "private receipt public_templates",
    )
    if public_templates["markdown_sha256"] != sha256(PUBLIC_MARKDOWN):
        raise ValueError("public Markdown template changed")
    if public_templates["docx_sha256"] != sha256(PUBLIC_DOCX):
        raise ValueError("public DOCX template changed")
    if public_templates["unchanged"] is not True:
        raise ValueError("private receipt does not prove public template stability")

    security_gate = read_json(SECURITY_GATE)
    security_decision = validate_security_gate(security_gate)
    link_allowed = security_decision["public_repository_link_allowed"]
    sanitized_allowed = security_decision["sanitized_external_response_allowed"]
    security = require_exact_keys(
        receipt["public_repository_security"],
        {
            "gate_sha256",
            "decision",
            "public_repository_link_allowed",
            "sanitized_external_response_allowed",
            "government_send_security_precondition_allowed",
            "public_repository_link_included",
            "attachment_repo_isolated",
        },
        "private receipt public_repository_security",
    )
    if security["gate_sha256"] != sha256(SECURITY_GATE):
        raise ValueError("private receipt public repository security hash is stale")
    if security["decision"] != security_gate["decision"]:
        raise ValueError("private receipt public repository security decision is stale")
    if security["public_repository_link_allowed"] is not link_allowed:
        raise ValueError("private receipt repository-link permission is stale")
    if security["sanitized_external_response_allowed"] is not sanitized_allowed:
        raise ValueError("private receipt sanitized-response permission is stale")
    if (
        security["government_send_security_precondition_allowed"]
        is not security_decision["final_argos_send_allowed_by_security_gate"]
    ):
        raise ValueError("private receipt security precondition is stale")
    if security["public_repository_link_included"] is not False:
        raise ValueError("private receipt repository-link inclusion is unsafe")
    if security["attachment_repo_isolated"] is not True:
        raise ValueError("private receipt does not prove repository isolation")

    outputs = require_exact_keys(
        receipt["outputs"],
        {"markdown", "docx"},
        "private receipt outputs",
    )
    markdown = require_exact_keys(
        outputs["markdown"],
        {"name", "bytes", "sha256"},
        "private receipt Markdown output",
    )
    docx = require_exact_keys(
        outputs["docx"],
        {"name", "bytes", "sha256"},
        "private receipt DOCX output",
    )
    if markdown["name"] != TEXT_OUTPUT_NAME or docx["name"] != DOCX_OUTPUT_NAME:
        raise ValueError("private receipt output names are stale")
    if markdown["bytes"] != markdown_path.stat().st_size:
        raise ValueError("private Markdown byte count is stale")
    if docx["bytes"] != docx_path.stat().st_size:
        raise ValueError("private DOCX byte count is stale")
    if markdown["sha256"] != sha256(markdown_path):
        raise ValueError("private Markdown custody hash is stale")
    if docx["sha256"] != sha256(docx_path):
        raise ValueError("private DOCX custody hash is stale")
    markdown_text = markdown_path.read_text(encoding="utf-8")
    with ZipFile(docx_path) as archive:
        document_xml = archive.read("word/document.xml")
    assert_no_external_repository_route(markdown_text, "private Markdown")
    assert_docx_repo_isolated(docx_path)

    gate = read_json(SUBMISSION_GATE)
    team = read_json(TEAM_REGISTER)
    expected_team_disclosure = bool(gate["send_gate"]["team_disclosure_resolved"])
    expected_name_authorizations = sum(
        int(candidate["verification"]["authorization_to_name_in_response"])
        for candidate in team["candidates"]
    )
    if receipt["response_mode"] != gate["response"]["response_mode"]:
        raise ValueError("private receipt response mode is stale")
    if receipt["team_disclosure_resolved"] is not expected_team_disclosure:
        raise ValueError("private receipt team-disclosure state is stale")
    if (
        receipt["candidate_name_authorization_count"]
        != expected_name_authorizations
    ):
        raise ValueError("private receipt candidate-authorization count is stale")

    required_false_controls = (
        "government_send_ready",
        "submission_authorized",
        "external_action_performed",
        "private_values_logged",
        "private_output_mirrored_to_public_vault",
    )
    if any(receipt[key] is not False for key in required_false_controls):
        raise ValueError("private receipt external-action controls are not fail-closed")
    if receipt["claim_boundary"] != CLAIM_BOUNDARY:
        raise ValueError("private receipt claim boundary is stale")
    if receipt["safest_next_action"] != SAFEST_NEXT_ACTION:
        raise ValueError("private receipt safest-next-action text is stale")


def public_summary(status: str, receipt: dict) -> dict:
    return {
        "status": status,
        "decision": receipt["decision"],
        "required_fact_count": receipt["required_fact_count"],
        "private_value_count": receipt["private_value_count"],
        "placeholder_count": receipt["placeholder_count"],
        "government_send_ready": receipt["government_send_ready"],
        "submission_authorized": receipt["submission_authorized"],
        "external_action_performed": receipt["external_action_performed"],
        "private_values_logged": receipt["private_values_logged"],
        "public_repository_link_included": receipt[
            "public_repository_security"
        ]["public_repository_link_included"],
        "attachment_repo_isolated": receipt[
            "public_repository_security"
        ]["attachment_repo_isolated"],
    }


def build_private_copy(
    facts_path: Path,
    output_dir: Path,
    as_of_utc: str,
) -> dict:
    evaluated = parse_utc(as_of_utc)
    private_facts_path = validate_private_path(facts_path, "facts path")
    private_output_dir = validate_private_path(output_dir, "output directory")
    payload = read_json(private_facts_path)
    facts = validate_facts(payload, evaluated)
    values = display_values(facts)
    gate = read_json(SUBMISSION_GATE)
    team = read_json(TEAM_REGISTER)
    security_gate = read_json(SECURITY_GATE)
    security_decision = validate_security_gate(security_gate)
    public_repository_link_allowed = security_decision[
        "public_repository_link_allowed"
    ]

    public_hashes_before = {
        "markdown_sha256": sha256(PUBLIC_MARKDOWN),
        "docx_sha256": sha256(PUBLIC_DOCX),
    }
    private_output_dir.mkdir(parents=True, exist_ok=True)
    markdown_path, docx_path, receipt_path = output_paths(private_output_dir)
    existing = [path.name for path in (markdown_path, docx_path, receipt_path) if path.exists()]
    if existing:
        raise FileExistsError("private output files already exist; use a new output directory")
    temporary_paths = {
        markdown_path: private_output_dir / f".{markdown_path.name}.tmp",
        docx_path: private_output_dir / f".{docx_path.name}.tmp",
        receipt_path: private_output_dir / f".{receipt_path.name}.tmp",
    }
    if any(path.exists() for path in temporary_paths.values()):
        raise FileExistsError("private staging files already exist; use a new output directory")

    private_markdown = render_private_markdown(
        PUBLIC_MARKDOWN.read_text(encoding="utf-8"),
        values,
    )
    staged_markdown = temporary_paths[markdown_path]
    staged_docx = temporary_paths[docx_path]
    staged_receipt = temporary_paths[receipt_path]
    final_paths = (markdown_path, docx_path, receipt_path)
    try:
        staged_markdown.write_text(private_markdown, encoding="utf-8", newline="\n")
        render_private_docx(
            PUBLIC_DOCX,
            staged_docx,
            values,
        )

        placeholder_count = (
            private_markdown.count(PRIVATE_MARKER)
            + private_markdown.count(PRIVATE_DISPLAY_MARKER)
        )
        public_hashes_after = {
            "markdown_sha256": sha256(PUBLIC_MARKDOWN),
            "docx_sha256": sha256(PUBLIC_DOCX),
        }
        field_states = expected_field_states(facts)
        team_disclosure_resolved = bool(
            gate["send_gate"]["team_disclosure_resolved"]
        )
        receipt = {
            "schema": RECEIPT_SCHEMA,
            "evaluated_utc": evaluated.strftime("%Y-%m-%dT%H:%M:%SZ"),
            "notice_id": NOTICE_ID,
            "decision": RECEIPT_DECISION,
            "required_fact_count": len(REQUIRED_FACT_KEYS),
            "private_value_count": len(facts),
            "placeholder_count": placeholder_count,
            "field_states": field_states,
            "facts_file": {
                "bytes": private_facts_path.stat().st_size,
                "sha256": sha256(private_facts_path),
                "path_logged": False,
            },
            "public_templates": {
                **public_hashes_after,
                "unchanged": public_hashes_before == public_hashes_after,
            },
            "public_repository_security": {
                "gate_sha256": sha256(SECURITY_GATE),
                "decision": security_gate["decision"],
                "public_repository_link_allowed": public_repository_link_allowed,
                "sanitized_external_response_allowed": security_decision[
                    "sanitized_external_response_allowed"
                ],
                "government_send_security_precondition_allowed": security_decision[
                    "final_argos_send_allowed_by_security_gate"
                ],
                "public_repository_link_included": False,
                "attachment_repo_isolated": True,
            },
            "outputs": {
                "markdown": {
                    "name": markdown_path.name,
                    "bytes": staged_markdown.stat().st_size,
                    "sha256": sha256(staged_markdown),
                },
                "docx": {
                    "name": docx_path.name,
                    "bytes": staged_docx.stat().st_size,
                    "sha256": sha256(staged_docx),
                },
            },
            "response_mode": gate["response"]["response_mode"],
            "team_disclosure_resolved": team_disclosure_resolved,
            "candidate_name_authorization_count": sum(
                int(candidate["verification"]["authorization_to_name_in_response"])
                for candidate in team["candidates"]
            ),
            "government_send_ready": False,
            "submission_authorized": False,
            "external_action_performed": False,
            "private_values_logged": False,
            "private_output_mirrored_to_public_vault": False,
            "claim_boundary": CLAIM_BOUNDARY,
            "safest_next_action": SAFEST_NEXT_ACTION,
        }
        validate_private_receipt(
            receipt,
            facts,
            private_facts_path,
            private_output_dir,
            staged_markdown,
            staged_docx,
            evaluated,
        )
        staged_receipt.write_text(
            json.dumps(receipt, indent=2, sort_keys=True) + "\n",
            encoding="utf-8",
            newline="\n",
        )
        staged_markdown.replace(markdown_path)
        staged_docx.replace(docx_path)
        staged_receipt.replace(receipt_path)
    except Exception:
        for generated_path in (*temporary_paths.values(), *final_paths):
            generated_path.unlink(missing_ok=True)
        raise
    return receipt


def check_private_copy(facts_path: Path, output_dir: Path, as_of_utc: str) -> dict:
    evaluated = parse_utc(as_of_utc)
    private_facts_path = validate_private_path(facts_path, "facts path")
    private_output_dir = validate_private_path(output_dir, "output directory")
    facts = validate_facts(read_json(private_facts_path), evaluated)
    markdown_path, docx_path, receipt_path = output_paths(private_output_dir)
    if not all(path.is_file() for path in (markdown_path, docx_path, receipt_path)):
        raise FileNotFoundError("private action-copy output set is incomplete")
    receipt = read_json(receipt_path)
    validate_private_receipt(
        receipt,
        facts,
        private_facts_path,
        private_output_dir,
        markdown_path,
        docx_path,
        evaluated,
    )
    markdown_text = markdown_path.read_text(encoding="utf-8")
    if PRIVATE_MARKER in markdown_text or PRIVATE_DISPLAY_MARKER in markdown_text:
        raise ValueError("private Markdown contains a placeholder")
    with ZipFile(docx_path) as archive:
        document_xml = archive.read("word/document.xml")
    if (
        PRIVATE_MARKER.encode() in document_xml
        or PRIVATE_DISPLAY_MARKER.encode() in document_xml
    ):
        raise ValueError("private DOCX contains a placeholder")
    return receipt


def main() -> int:
    parser = argparse.ArgumentParser(
        description=(
            "Build or verify a private Project Argos action copy without logging "
            "private values or writing inside the public repository."
        )
    )
    parser.add_argument("--facts", required=True, type=Path)
    parser.add_argument("--output-dir", required=True, type=Path)
    parser.add_argument("--as-of-utc")
    parser.add_argument("--check", action="store_true")
    args = parser.parse_args()
    as_of = args.as_of_utc or datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
    try:
        receipt = (
            check_private_copy(args.facts, args.output_dir, as_of)
            if args.check
            else build_private_copy(args.facts, args.output_dir, as_of)
        )
    except (ValueError, FileNotFoundError, FileExistsError, KeyError, json.JSONDecodeError) as exc:
        print(
            json.dumps(
                {
                    "status": "BLOCKED",
                    "error_type": type(exc).__name__,
                    "private_values_logged": False,
                    "external_action_performed": False,
                },
                indent=2,
            )
        )
        return 1
    status = "CURRENT" if args.check else "PRIVATE_ACTION_COPY_WRITTEN"
    print(json.dumps(public_summary(status, receipt), indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
