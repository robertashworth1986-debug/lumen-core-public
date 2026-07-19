from __future__ import annotations

import argparse
import re
from collections.abc import Iterator, Sequence
from pathlib import Path
from uuid import uuid4
from zipfile import ZIP_DEFLATED, ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


HERE = Path(__file__).resolve().parent
DEFAULT_SOURCE = HERE / "NV063_DSIP_VOLUME2_FINAL_CANDIDATE_2026-07-16.md"
DEFAULT_TEMPLATE = (
    HERE
    / "source_attachments"
    / "DON_PHASE_I_CONVENTIONAL_VOLUME2_TEMPLATE_2026-03-04.docx"
)
DEFAULT_OUTPUT = HERE / "NV063_DSIP_VOLUME2_FINAL_CANDIDATE_2026-07-16.docx"

FONT = "Arial"
TOPIC_NUMBER = "DON26BZ03-NV063"
DSIP_TOPIC_SEGMENT = "NV063"
REVIEW_PROPOSAL_LABEL = "Proposal No. [assigned in DSIP]"
# Mirrors the assigned DSIP shape already recorded in this repository without
# guessing HarborSentinel's agency prefix or assigned serial.
DSIP_PROPOSAL_NUMBER = re.compile(
    rf"[A-Z0-9]{{3,16}}-{DSIP_TOPIC_SEGMENT}-[A-Z0-9]{{4,16}}"
)
PROPOSAL_NUMBER_PLACEHOLDER_TERMS = (
    "ASSIGN",
    "INSERT",
    "PENDING",
    "PLACEHOLDER",
    "UNKNOWN",
    "TBD",
)
REQUIRED_SECTION_HEADINGS = (
    "1.0 Description of Proposed Phase I Technical Effort",
    "1.1 Phase I Technical Objectives",
    "1.2 Phase I Base and Option Statement of Work",
    "1.3 Related Work",
    "2.0 Key Personnel",
    "3.0 Commercialization and Transition Plan Summary",
    "4.0 Facilities and Equipment",
    "5.0 Letters of Support",
)
REQUIRED_CLAIM_BOUNDARIES = (
    "The system is advisory. It does not autonomously determine hostile intent or authorize an operational action.",
    "Phase I does not claim access to Navy radar, classified sensor data, tactical SSDS software, operational watch-floor data, or Government-furnished interfaces.",
    "Existing evidence is internal feasibility work, not field validation.",
    "Controlled injections are not real adversary labels. The natural queue is not a false-positive rate.",
    "It does not mean SSDS integration, classified sensor validation, operational threat classification, field readiness, or CMMC or clearance completion.",
    "No Navy endorsement, operational access, classified work, Government-furnished data, SSDS integration, or field result is claimed.",
    "No current Navy customer, pilot, revenue, field performance, or production deployment is claimed.",
    "The current ordinary software-development facility is not represented as a CUI enclave, classified facility, accredited system, or cleared facility.",
    "The absence of a letter is not replaced with an unsupported partner, customer, Navy sponsor, transition commitment, or validation claim.",
)
RELEASE_BODY_REPLACEMENTS = {
    "No letter of support is included in this review candidate.": (
        "No letter of support is included in this Volume 2."
    ),
}
RELEASE_FORBIDDEN_MARKERS = (
    "REVIEW CANDIDATE",
    "NOT CERTIFIED",
    "[ASSIGNED IN DSIP]",
    "[INSERT PROPOSAL NUMBER",
    "PROPOSAL NO. ASSIGNED IN DSIP",
    "PROPOSAL NUMBER ASSIGNED IN DSIP",
    "REMOVE THE DRAFT CONTROL",
    "COMPLETE LIVE DSIP",
)
GENERIC_PLACEHOLDER = re.compile(
    r"\b(?:TBD|TODO|PLACEHOLDER)\b|\[(?:ASSIGNED|INSERT|FIRM NAME|REMOVE)\b",
    re.IGNORECASE,
)
FIRST_PAGE_LEGEND = (
    "This proposal includes data that must not be disclosed outside the Government "
    "and must not be duplicated, used, or disclosed-in whole or in part-for any "
    "purpose other than to evaluate this proposal. If, however, a contract is "
    "awarded to this proposing SBC as a result of-or in connection with-the "
    "submission of this data, the Government has the right to duplicate, use, or "
    "disclose the data to the extent provided in the resulting contract. This "
    "restriction does not limit the Government's right to use information contained "
    "in this data if it is obtained from another source without restriction. The "
    "data subject to this restriction are contained in all pages of this Volume 2."
)
PAGE_LEGEND = (
    "Use or disclosure of data contained on this page is subject to the restriction "
    "on the first page of this volume."
)


def scrub_package_artifacts(docx_path: Path) -> None:
    tmp_path = docx_path.with_name(
        f".{docx_path.name}.{uuid4().hex}.scrubbed.docx"
    )
    try:
        with ZipFile(docx_path, "r") as src, ZipFile(
            tmp_path, "w", ZIP_DEFLATED
        ) as dst:
            for item in src.infolist():
                name = item.filename
                if name == "docProps/custom.xml" or name.startswith("customXml/"):
                    continue
                data = src.read(name)
                if name == "[Content_Types].xml":
                    root = ET.fromstring(data)
                    for node in list(root):
                        part_name = node.attrib.get("PartName", "")
                        if (
                            part_name.startswith("/customXml/")
                            or part_name == "/docProps/custom.xml"
                        ):
                            root.remove(node)
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                elif name.endswith(".rels"):
                    root = ET.fromstring(data)
                    for node in list(root):
                        rel_type = node.attrib.get("Type", "")
                        target = node.attrib.get("Target", "").replace("\\", "/")
                        if (
                            rel_type.endswith("/customXml")
                            or rel_type.endswith("/custom-properties")
                            or target.startswith("customXml/")
                            or target.startswith("../customXml/")
                            or target == "docProps/custom.xml"
                        ):
                            root.remove(node)
                    data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                dst.writestr(item, data)
        tmp_path.replace(docx_path)
    finally:
        tmp_path.unlink(missing_ok=True)


def clear_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_margins(cell, top=72, start=90, bottom=72, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_simple_field(paragraph, instruction: str, cached_text: str) -> None:
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = cached_text
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def add_page_number(paragraph) -> None:
    paragraph.add_run("Page ")
    add_simple_field(paragraph, "PAGE", "1")
    paragraph.add_run(" of ")
    add_simple_field(paragraph, "NUMPAGES", "1")


def clear_header_or_footer(part) -> None:
    element = part._element
    for child in list(element):
        if child.tag == qn("w:p"):
            element.remove(child)


def create_bullet_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "\u2022")
    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "left")
    paragraph_properties = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "360")
    indent.set(qn("w:hanging"), "180")
    paragraph_properties.append(indent)
    level.extend([start, num_fmt, level_text, level_justification, paragraph_properties])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id: int) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    num_properties = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_properties.extend([level, num])
    paragraph_properties.append(num_properties)


def configure_document(
    doc: Document, released: bool, proposal_number: str | None = None
) -> int:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = False
    doc.settings.odd_and_even_pages_header_footer = False

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1
    normal.paragraph_format.left_indent = Inches(0)
    normal.paragraph_format.right_indent = Inches(0)
    normal.paragraph_format.first_line_indent = Inches(0)

    for style_name, size, before, after in (
        ("Heading 1", 12, 7, 3),
        ("Heading 2", 11, 5, 2),
        ("Heading 3", 10, 4, 1),
    ):
        try:
            style = doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.left_indent = Inches(0)
        style.paragraph_format.right_indent = Inches(0)
        style.paragraph_format.first_line_indent = Inches(0)

    try:
        list_bullet = doc.styles["List Bullet"]
    except KeyError:
        list_bullet = doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
    list_bullet.font.name = FONT
    list_bullet.font.size = Pt(10)
    list_bullet.paragraph_format.space_before = Pt(0)
    list_bullet.paragraph_format.space_after = Pt(1)
    list_bullet.paragraph_format.line_spacing = 1

    clear_header_or_footer(section.header)
    header = section.header.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    proposal_label = (
        f"Proposal No. {proposal_number}" if released else REVIEW_PROPOSAL_LABEL
    )
    header_run = header.add_run(
        f"Robert Ashworth d/b/a LumenCore | {TOPIC_NUMBER} | {proposal_label}"
    )
    header_run.font.name = FONT
    header_run.font.size = Pt(7.5)

    clear_header_or_footer(section.footer)
    footer_legend = section.footer.add_paragraph()
    footer_legend.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_legend.paragraph_format.space_after = Pt(0)
    legend = footer_legend.add_run(PAGE_LEGEND)
    legend.font.name = FONT
    legend.font.size = Pt(6.5)

    footer_page = section.footer.add_paragraph()
    footer_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_page.paragraph_format.space_after = Pt(0)
    add_page_number(footer_page)
    if not released:
        draft = footer_page.add_run(" | REVIEW CANDIDATE - NOT CERTIFIED")
        draft.font.name = FONT
        draft.font.size = Pt(7.5)
        draft.font.bold = True

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    return create_bullet_numbering(doc)


def add_title_block(doc: Document, released: bool) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Volume 2: Technical Volume")
    run.font.name = FONT
    run.font.size = Pt(14)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(
        "HarborSentinel: Explainable Low-Storage Pattern-of-Life Analysis for Congested Maritime Environments"
    )
    run.font.name = FONT
    run.font.size = Pt(11)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("DON26BZ03-NV063 | Navy SBIR 2026 Release 3 Phase I")
    run.font.name = FONT
    run.font.size = Pt(9)

    if not released:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run("REVIEW CANDIDATE - COMPLETE LIVE DSIP AND HUMAN CERTIFICATION GATES BEFORE SUBMISSION")
        run.font.name = FONT
        run.font.size = Pt(8)
        run.font.bold = True

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(FIRST_PAGE_LEGEND)
    run.font.name = FONT
    run.font.size = Pt(8)
    run.font.italic = True


def markdown_blocks(text: str) -> Iterator[tuple[str, str]]:
    paragraph: list[str] = []

    def flush():
        if paragraph:
            yield "paragraph", " ".join(paragraph).strip()
            paragraph.clear()

    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            yield from flush()
            continue
        if line.startswith("# "):
            yield from flush()
            yield "title", line[2:].strip()
        elif line.startswith("## "):
            yield from flush()
            yield "h1", line[3:].strip()
        elif line.startswith("### "):
            yield from flush()
            yield "h2", line[4:].strip()
        elif line.startswith("- "):
            yield from flush()
            yield "bullet", line[2:].strip()
        else:
            paragraph.append(line)
    yield from flush()


def document_blocks(
    source_text: str, released: bool = False
) -> Iterator[tuple[str, str]]:
    skip_prefixes = (
        "Topic:",
        "Program:",
        "Proposal title:",
        "Status:",
    )
    for kind, text in markdown_blocks(source_text):
        if kind == "title" or text.startswith(skip_prefixes):
            continue
        if released:
            for review_text, release_text in RELEASE_BODY_REPLACEMENTS.items():
                text = text.replace(review_text, release_text)
        yield kind, text


def add_markdown(
    doc: Document,
    source_text: str,
    bullet_num_id: int,
    released: bool = False,
) -> None:
    for kind, text in document_blocks(source_text, released=released):
        if kind == "h1":
            doc.add_paragraph(text, style="Heading 1")
        elif kind == "h2":
            doc.add_paragraph(text, style="Heading 2")
        elif kind == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1
            apply_bullet(p, bullet_num_id)
            run = p.add_run(text)
            run.font.name = FONT
            run.font.size = Pt(10)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.right_indent = Inches(0)
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run(text)
            run.font.name = FONT
            run.font.size = Pt(10)


def remove_empty_template_tables(doc: Document) -> None:
    for table in list(doc.tables):
        table._element.getparent().remove(table._element)


def validate_release_proposal_number(value: str | None) -> str:
    if not isinstance(value, str) or not value:
        raise ValueError("RELEASE_REQUIRES_EXPLICIT_DSIP_PROPOSAL_NUMBER")
    if value != value.strip():
        raise ValueError("INVALID_DSIP_PROPOSAL_NUMBER_FORMAT: surrounding whitespace")
    if DSIP_PROPOSAL_NUMBER.fullmatch(value) is None:
        raise ValueError(
            "INVALID_DSIP_PROPOSAL_NUMBER_FORMAT: expected an uppercase DSIP identifier "
            f"with the {DSIP_TOPIC_SEGMENT} topic segment"
        )
    if any(term in value for term in PROPOSAL_NUMBER_PLACEHOLDER_TERMS):
        raise ValueError("INVALID_DSIP_PROPOSAL_NUMBER_FORMAT: placeholder value")
    return value


def release_text_violations(text: str) -> list[str]:
    upper_text = text.upper()
    violations = [
        marker for marker in RELEASE_FORBIDDEN_MARKERS if marker in upper_text
    ]
    if GENERIC_PLACEHOLDER.search(text):
        violations.append("GENERIC_PLACEHOLDER")
    return sorted(set(violations))


def validate_release_source(source_text: str) -> None:
    errors: list[str] = []
    heading_positions: list[int] = []
    for heading in REQUIRED_SECTION_HEADINGS:
        token = f"## {heading}"
        position = source_text.find(token)
        if position < 0:
            errors.append(f"missing section: {heading}")
        else:
            heading_positions.append(position)
    if len(heading_positions) == len(REQUIRED_SECTION_HEADINGS):
        if heading_positions != sorted(heading_positions):
            errors.append("required sections are out of order")

    for boundary in REQUIRED_CLAIM_BOUNDARIES:
        if boundary not in source_text:
            errors.append(f"missing claim boundary: {boundary}")

    for required_fact in (
        "Phase I Base, months 1-6, not to exceed $200,000",
        "Phase I Option, months 7-12, not to exceed $115,000",
    ):
        if required_fact not in source_text:
            errors.append(f"missing Base/Option boundary: {required_fact}")

    rendered_source = "\n".join(
        text for _, text in document_blocks(source_text, released=True)
    )
    for violation in release_text_violations(rendered_source):
        errors.append(f"rendered source contains {violation}")

    if errors:
        raise ValueError("RELEASE_SOURCE_VALIDATION_FAILED: " + "; ".join(errors))


def part_text(part) -> str:
    paragraphs = [paragraph.text for paragraph in part.paragraphs]
    for table in part.tables:
        for row in table.rows:
            paragraphs.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs).strip()


def package_story_text(docx_path: Path) -> tuple[str, list[str]]:
    story_texts: list[str] = []
    custom_parts: list[str] = []
    with ZipFile(docx_path, "r") as package:
        corrupt_part = package.testzip()
        if corrupt_part is not None:
            raise ValueError(f"DOCX_PACKAGE_CORRUPT: {corrupt_part}")
        for name in package.namelist():
            if name == "docProps/custom.xml" or name.startswith("customXml/"):
                custom_parts.append(name)
            if not (
                name == "word/document.xml"
                or name == "docProps/core.xml"
                or re.fullmatch(r"word/(?:header|footer)\d+\.xml", name)
            ):
                continue
            root = ET.fromstring(package.read(name))
            story_texts.append("".join(root.itertext()))
    return "\n".join(story_texts), custom_parts


def validate_release_document(
    docx_path: Path, source_text: str, proposal_number: str
) -> None:
    errors: list[str] = []
    doc = Document(docx_path)

    if len(doc.sections) != 1:
        errors.append(f"expected one section, found {len(doc.sections)}")
    if doc.tables:
        errors.append(f"expected no template tables, found {len(doc.tables)}")

    section = doc.sections[0]
    expected_dimensions = (
        ("page width", section.page_width, 8.5),
        ("page height", section.page_height, 11.0),
        ("top margin", section.top_margin, 1.0),
        ("right margin", section.right_margin, 1.0),
        ("bottom margin", section.bottom_margin, 1.0),
        ("left margin", section.left_margin, 1.0),
    )
    for label, measurement, expected_inches in expected_dimensions:
        if measurement is None or abs(measurement.inches - expected_inches) > 0.001:
            errors.append(f"invalid {label}")

    if section.different_first_page_header_footer:
        errors.append("different first-page header/footer must be disabled")
    if doc.settings.odd_and_even_pages_header_footer:
        errors.append("odd/even header/footer mode must be disabled")

    required_headers = [item.header for item in doc.sections]
    required_footers = [item.footer for item in doc.sections]
    expected_header = (
        f"Robert Ashworth d/b/a LumenCore | {TOPIC_NUMBER} | "
        f"Proposal No. {proposal_number}"
    )
    for index, header in enumerate(required_headers, start=1):
        header_text = part_text(header)
        if header_text != expected_header:
            errors.append(f"section {index} header is not exact")
        if header_text.count(proposal_number) != 1:
            errors.append(
                f"section {index} header does not contain the proposal number exactly once"
            )

    for index, footer in enumerate(required_footers, start=1):
        footer_text = part_text(footer)
        if PAGE_LEGEND not in footer_text:
            errors.append(f"section {index} footer is missing the proprietary legend")
        field_instructions = {
            (node.get(qn("w:instr")) or "").strip()
            for node in footer._element.iter(qn("w:fldSimple"))
        }
        if not {"PAGE", "NUMPAGES"}.issubset(field_instructions):
            errors.append(f"section {index} footer is missing page fields")

    update_fields = doc.settings._element.find(qn("w:updateFields"))
    if update_fields is None or update_fields.get(qn("w:val")) != "true":
        errors.append("Word field updates are not enabled")

    for style_name, minimum_size in (
        ("Normal", 10),
        ("Heading 1", 10),
        ("Heading 2", 10),
        ("Heading 3", 10),
        ("List Bullet", 10),
    ):
        try:
            style = doc.styles[style_name]
        except KeyError:
            errors.append(f"missing style: {style_name}")
            continue
        if style.font.name != FONT:
            errors.append(f"{style_name} does not use {FONT}")
        if style.font.size is None or style.font.size.pt < minimum_size:
            errors.append(f"{style_name} is smaller than {minimum_size} point")

    body_texts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
    expected_prefix = [
        "Volume 2: Technical Volume",
        "HarborSentinel: Explainable Low-Storage Pattern-of-Life Analysis for Congested Maritime Environments",
        f"{TOPIC_NUMBER} | Navy SBIR 2026 Release 3 Phase I",
        FIRST_PAGE_LEGEND,
    ]
    if body_texts[: len(expected_prefix)] != expected_prefix:
        errors.append("title block or first-page proprietary legend is incomplete")

    expected_source_blocks = [
        text for _, text in document_blocks(source_text, released=True)
    ]
    if body_texts[len(expected_prefix) :] != expected_source_blocks:
        errors.append("rendered source blocks differ from the bounded release source")

    body_text = "\n".join(body_texts)
    heading_positions: list[int] = []
    for heading in REQUIRED_SECTION_HEADINGS:
        position = body_text.find(heading)
        if position < 0:
            errors.append(f"rendered document is missing section: {heading}")
        else:
            heading_positions.append(position)
    if len(heading_positions) == len(REQUIRED_SECTION_HEADINGS):
        if heading_positions != sorted(heading_positions):
            errors.append("rendered sections are out of order")
    for boundary in REQUIRED_CLAIM_BOUNDARIES:
        if boundary not in body_text:
            errors.append(f"rendered document lost claim boundary: {boundary}")

    package_text, custom_parts = package_story_text(docx_path)
    if custom_parts:
        errors.append("custom package artifacts remain: " + ", ".join(custom_parts))
    combined_text = "\n".join(
        [
            body_text,
            *(part_text(header) for header in required_headers),
            *(part_text(footer) for footer in required_footers),
            package_text,
        ]
    )
    for violation in release_text_violations(combined_text):
        errors.append(f"release package contains {violation}")
    embedded_numbers = set(DSIP_PROPOSAL_NUMBER.findall(combined_text))
    if embedded_numbers != {proposal_number}:
        errors.append(
            "release package contains a missing or mismatched DSIP proposal number"
        )

    if errors:
        raise ValueError("RELEASE_DOCUMENT_VALIDATION_FAILED: " + "; ".join(errors))


def build(
    source: Path,
    template: Path,
    output: Path,
    released: bool,
    proposal_number: str | None = None,
    overwrite: bool = False,
) -> Path:
    if released:
        proposal_number = validate_release_proposal_number(proposal_number)
    elif proposal_number is not None:
        raise ValueError("PROPOSAL_NUMBER_IS_RELEASE_ONLY")

    if not source.is_file():
        raise FileNotFoundError(f"SOURCE_NOT_FOUND: {source}")
    if not template.is_file():
        raise FileNotFoundError(f"TEMPLATE_NOT_FOUND: {template}")
    if released and output.suffix.lower() != ".docx":
        raise ValueError("RELEASE_OUTPUT_MUST_BE_DOCX")
    if output.resolve() in {source.resolve(), template.resolve()}:
        raise ValueError("OUTPUT_MUST_NOT_OVERWRITE_AN_INPUT")
    if released and output.exists() and not overwrite:
        raise FileExistsError(
            f"RELEASE_OUTPUT_EXISTS: {output}; pass --overwrite only after verifying the target"
        )

    source_text = source.read_text(encoding="utf-8")
    if released:
        validate_release_source(source_text)

    doc = Document(template)
    clear_body(doc)
    remove_empty_template_tables(doc)
    bullet_num_id = configure_document(
        doc, released=released, proposal_number=proposal_number
    )
    add_title_block(doc, released=released)
    add_markdown(doc, source_text, bullet_num_id, released=released)

    doc.core_properties.title = "HarborSentinel Navy SBIR Phase I Volume 2"
    doc.core_properties.subject = TOPIC_NUMBER
    doc.core_properties.author = "Robert Ashworth"
    doc.core_properties.comments = (
        "Release build; explicit DSIP proposal number structurally verified"
        if released
        else "Review candidate; not certified"
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    temporary_output = output.with_name(
        f".{output.stem}.{uuid4().hex}.tmp{output.suffix or '.docx'}"
    )
    try:
        doc.save(temporary_output)
        scrub_package_artifacts(temporary_output)
        if released:
            validate_release_document(
                temporary_output, source_text, proposal_number
            )
            if output.exists() and not overwrite:
                raise FileExistsError(
                    f"RELEASE_OUTPUT_EXISTS: {output}; target appeared during build"
                )
        temporary_output.replace(output)
    finally:
        temporary_output.unlink(missing_ok=True)

    print(output)
    return output


def main(argv: Sequence[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Build the HarborSentinel DSIP Volume 2 review or guarded release DOCX."
    )
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument(
        "--release",
        action="store_true",
        help="Remove visible review-candidate controls only after all live DSIP gates clear.",
    )
    parser.add_argument(
        "--proposal-number",
        help="Exact proposal number assigned by DSIP; required with --release.",
    )
    parser.add_argument(
        "--overwrite",
        "--allow-overwrite",
        action="store_true",
        help="Allow a verified release build to replace an existing output.",
    )
    args = parser.parse_args(argv)

    if not args.release and args.proposal_number is not None:
        parser.error("--proposal-number may only be used with --release")
    if not args.release and args.overwrite:
        parser.error("--overwrite may only be used with --release")
    try:
        build(
            args.source,
            args.template,
            args.output,
            released=args.release,
            proposal_number=args.proposal_number,
            overwrite=args.overwrite,
        )
    except (FileExistsError, FileNotFoundError, ValueError) as exc:
        parser.error(str(exc))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
