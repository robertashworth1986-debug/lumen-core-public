from __future__ import annotations

from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
SOURCE = HERE / "NV063_VOLUME2_TECHNICAL_DRAFT_2026-06-19.md"
OUTPUT = HERE / "NV063_VOLUME2_TECHNICAL_DRAFT_2026-06-19.docx"

FONT = "Calibri"
INK = RGBColor(0x00, 0x00, 0x00)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
GRAY = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0x9B, 0x1C, 0x1C)


def _remove_package_relationships(xml_bytes: bytes) -> bytes:
    root = ET.fromstring(xml_bytes)
    changed = False
    for rel in list(root):
        rel_type = rel.attrib.get("Type", "")
        target = rel.attrib.get("Target", "").replace("\\", "/")
        if (
            rel_type.endswith("/customXml")
            or rel_type.endswith("/custom-properties")
            or target.startswith("customXml/")
            or target.startswith("../customXml/")
            or target == "docProps/custom.xml"
        ):
            root.remove(rel)
            changed = True
    return ET.tostring(root, encoding="utf-8", xml_declaration=True) if changed else xml_bytes


def _remove_content_type_overrides(xml_bytes: bytes) -> bytes:
    root = ET.fromstring(xml_bytes)
    changed = False
    for override in list(root):
        part_name = override.attrib.get("PartName", "")
        if part_name.startswith("/customXml/") or part_name == "/docProps/custom.xml":
            root.remove(override)
            changed = True
    return ET.tostring(root, encoding="utf-8", xml_declaration=True) if changed else xml_bytes


def scrub_package_artifacts(docx_path: Path) -> None:
    """Remove default DOCX customXml sidecars that are not part of the proposal."""
    tmp_path = docx_path.with_suffix(".scrubbed.tmp.docx")
    with ZipFile(docx_path, "r") as src, ZipFile(tmp_path, "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            name = item.filename
            if name == "docProps/custom.xml" or name.startswith("customXml/"):
                continue
            data = src.read(name)
            if name == "[Content_Types].xml":
                data = _remove_content_type_overrides(data)
            elif name.endswith(".rels"):
                data = _remove_package_relationships(data)
            dst.writestr(item, data)
    tmp_path.replace(docx_path)


def set_run(run, *, size: float = 11, bold: bool = False, italic: bool = False, color=INK) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def configure_document(doc: Document) -> int:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, before, after, color in (
        ("Heading 1", 16, 16, 8, BLUE),
        ("Heading 2", 13, 12, 6, BLUE),
        ("Heading 3", 12, 8, 4, DARK_BLUE),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(footer.add_run("WORKING DRAFT - NOT APPROVED FOR SUBMISSION"), size=8.5, color=GRAY)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(header.add_run("NV063 HarborSentinel | Volume 2 Draft"), size=8.5, color=GRAY)

    return create_bullet_numbering(doc)


def create_bullet_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "\u2022")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "280")
    p_pr.append(tabs)
    p_pr.append(ind)
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(r_fonts)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    set_run(p.add_run("HarborSentinel"), size=22, bold=True, color=DARK_BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(4)
    set_run(
        subtitle.add_run(
            "Anomalous Behavior Detection and Alerting for Congested Maritime Environments"
        ),
        size=12,
        bold=True,
        color=GRAY,
    )

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(10)
    set_run(
        meta.add_run(
            "DON26BZ03-NV063 | Navy SBIR 2026 Release 3 Phase I | Volume 2 Technical Draft"
        ),
        size=10,
        color=GRAY,
    )

    warning = doc.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning.paragraph_format.space_after = Pt(10)
    set_run(warning.add_run("WORKING DRAFT - NOT APPROVED FOR SUBMISSION"), size=10, bold=True, color=RED)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles[f"Heading {min(level, 3)}"]
    set_run(paragraph.add_run(text), size=16 if level == 1 else 13 if level == 2 else 12, bold=True, color=BLUE if level < 3 else DARK_BLUE)


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    set_run(paragraph.add_run(text), size=11)


def add_bullet(doc: Document, text: str, num_id: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.208
    apply_bullet(paragraph, num_id)
    set_run(paragraph.add_run(text), size=11)


def iter_markdown_blocks(text: str):
    pending_kind: str | None = None
    pending_text: str = ""

    def flush():
        nonlocal pending_kind, pending_text
        if pending_kind and pending_text:
            yield pending_kind, pending_text.strip()
        pending_kind = None
        pending_text = ""

    for raw_line in text.splitlines():
        stripped = raw_line.strip()
        if not stripped:
            yield from flush()
            continue
        if stripped.startswith("#"):
            yield from flush()
            if stripped.startswith("### "):
                yield "h3", stripped[4:].strip()
            elif stripped.startswith("## "):
                yield "h2", stripped[3:].strip()
            else:
                yield "h1", stripped[2:].strip()
            continue
        if stripped.startswith("- "):
            yield from flush()
            pending_kind = "bullet"
            pending_text = stripped[2:].strip()
            continue
        if pending_kind in {"paragraph", "bullet"}:
            pending_text += " " + stripped
        else:
            pending_kind = "paragraph"
            pending_text = stripped

    yield from flush()


def build() -> None:
    doc = Document()
    bullet_num_id = configure_document(doc)
    add_title_block(doc)

    for kind, text in iter_markdown_blocks(SOURCE.read_text(encoding="utf-8")):
        if kind == "h1" and text == "HarborSentinel Volume 2 Technical Draft":
            continue
        if kind == "h3":
            add_heading(doc, text, 3)
        elif kind == "h2":
            add_heading(doc, text, 2)
        elif kind == "h1":
            add_heading(doc, text, 1)
        elif kind == "bullet":
            add_bullet(doc, text, bullet_num_id)
        else:
            add_body(doc, text)

    doc.core_properties.title = "HarborSentinel Volume 2 Technical Draft"
    doc.core_properties.subject = "DON26BZ03-NV063 Navy SBIR Phase I"
    doc.core_properties.author = "Robert Ashworth"
    doc.core_properties.comments = (
        "Working draft for DSIP/Navy Volume 2 conversion. Not approved for submission."
    )
    doc.save(OUTPUT)
    scrub_package_artifacts(OUTPUT)
    # Normalize the OOXML package after relationship cleanup so LibreOffice
    # accepts the file while keeping the non-proposal customXml sidecars removed.
    Document(OUTPUT).save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
