#!/usr/bin/env python3
"""Build the MissionWeave DSIP Volume 2 DOCX from its bounded Markdown source."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
DEFAULT_SOURCE = HERE / "MISSIONWEAVE_DSIP_VOLUME2_FINAL_CANDIDATE_2026-07-16.md"
DEFAULT_OUTPUT = HERE / "MISSIONWEAVE_DSIP_VOLUME2_FINAL_CANDIDATE_2026-07-16.docx"
DEFAULT_METADATA = HERE / "MISSIONWEAVE_DSIP_VOLUME2_BUILD_METADATA_2026-07-16.json"
NEUTRAL_PROPOSAL_HEADER = "Proposal No. assigned in DSIP"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}

FONT = "Arial"
BLACK = RGBColor(0x00, 0x00, 0x00)
NAVY = RGBColor(0x1F, 0x4D, 0x78)
BLUE = RGBColor(0x2E, 0x62, 0x8A)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = "F2F4F7"
BORDER_GRAY = "AEB7C2"

PRESET = {
    "preset": "grant_proposal",
    "named_override": "DLA single-spaced technical volume",
    "page": {
        "size": "US Letter portrait",
        "margins_inches": 1.0,
        "header_footer_inches": 0.5,
        "content_width_dxa": CONTENT_WIDTH_DXA,
    },
    "body": {
        "font": FONT,
        "size_pt": 10.5,
        "line_spacing": "single",
        "space_after_pt": 4.5,
    },
    "headings": {
        "title": {"size_pt": 18, "color": "000000", "after_pt": 8},
        "h1": {"size_pt": 12.5, "color": "2E628A", "before_pt": 10, "after_pt": 4},
        "h2": {"size_pt": 11.5, "color": "1F4D78", "before_pt": 7, "after_pt": 3},
        "h3": {"size_pt": 10.5, "color": "1F4D78", "before_pt": 5, "after_pt": 2},
    },
    "lists": {
        "marker_aligned_dxa": 260,
        "text_indent_dxa": 540,
        "hanging_dxa": 280,
        "space_after_pt": 2,
    },
    "tables": {
        "width_dxa": CONTENT_WIDTH_DXA,
        "indent_dxa": TABLE_INDENT_DXA,
        "cell_margins_dxa": CELL_MARGIN_DXA,
        "font_size_pt": 10,
        "header_fill": LIGHT_GRAY,
    },
}


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def set_run_font(run, *, size: float = 10.5, color=BLACK, bold=None, italic=None) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style_font(style, *, size: float, color=BLACK, bold=None, italic=None) -> None:
    style.font.name = FONT
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(size)
    style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    set_style_font(normal, size=10.5, color=BLACK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.widow_control = True

    if "Proposal Title" not in doc.styles:
        title = doc.styles.add_style("Proposal Title", WD_STYLE_TYPE.PARAGRAPH)
        title.base_style = normal
    else:
        title = doc.styles["Proposal Title"]
    set_style_font(title, size=18, color=BLACK, bold=True)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    title.paragraph_format.keep_with_next = True
    title_p_pr = title._element.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)

    heading_specs = [
        ("Heading 1", 12.5, BLUE, 10, 4),
        ("Heading 2", 11.5, NAVY, 7, 3),
        ("Heading 3", 10.5, NAVY, 5, 2),
    ]
    for name, size, color, before, after in heading_specs:
        style = doc.styles[name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Proposal Metadata" not in doc.styles:
        metadata = doc.styles.add_style("Proposal Metadata", WD_STYLE_TYPE.PARAGRAPH)
    else:
        metadata = doc.styles["Proposal Metadata"]
    set_style_font(metadata, size=10.5, color=BLACK)
    metadata.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_before = Pt(0)
    metadata.paragraph_format.space_after = Pt(2)
    metadata.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    metadata.paragraph_format.keep_with_next = True

    if "Table Text" not in doc.styles:
        table_text = doc.styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
    else:
        table_text = doc.styles["Table Text"]
    set_style_font(table_text, size=10, color=BLACK)
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    if "Reference" not in doc.styles:
        reference = doc.styles.add_style("Reference", WD_STYLE_TYPE.PARAGRAPH)
    else:
        reference = doc.styles["Reference"]
    set_style_font(reference, size=10, color=BLACK)
    reference.paragraph_format.left_indent = Inches(0.25)
    reference.paragraph_format.first_line_indent = Inches(-0.25)
    reference.paragraph_format.space_after = Pt(2)
    reference.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)


def set_cell_margins(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGIN_DXA.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER_GRAY)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def apply_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        set_row_cant_split(row)
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    set_table_borders(table)
    set_repeat_table_header(table.rows[0])


def table_widths(headers: list[str]) -> list[int]:
    key = tuple(h.strip().lower() for h in headers)
    if key == ("objective", "phase i acceptance evidence"):
        return [2700, 6660]
    if key == ("task", "schedule", "work and outputs"):
        return [2400, 1200, 5760]
    if key == (
        "condition",
        "mean on-time delta vs cross-trained fifo",
        "paired 95% bootstrap interval",
        "better / tied / worse seeds",
    ):
        return [1900, 2400, 2650, 2410]
    if key == ("deliverable", "planned timing"):
        return [6900, 2460]
    if key == ("risk", "control and decision gate"):
        return [2450, 6910]
    if key == ("milestone", "evidence required for advancement"):
        return [2600, 6760]
    if key == (
        "effort and authoritative identifiers",
        "submission, award, and pi facts",
        "relationship to missionweave and cost separation",
    ):
        return [3300, 2500, 3560]
    if key == (
        "technical data or computer software to be furnished with restrictions",
        "basis for assertion",
        "asserted rights category",
        "person or organization asserting restrictions",
    ):
        return [3200, 2300, 1850, 2010]
    if len(headers) == 2:
        return [2800, 6560]
    if len(headers) == 3:
        return [1800, 2500, 5060]
    if len(headers) == 4:
        return [1800, 2500, 2500, 2560]
    base = CONTENT_WIDTH_DXA // len(headers)
    widths = [base] * len(headers)
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def create_numbering(doc: Document) -> tuple[int, int]:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abstract = max(abstract_ids, default=0) + 1
    next_num = max(num_ids, default=0) + 1

    def add_definition(abstract_id: int, num_id: int, fmt: str, text: str, font: str | None = None) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        level.append(num_fmt)
        level_text = OxmlElement("w:lvlText")
        level_text.set(qn("w:val"), text)
        level.append(level_text)
        justify = OxmlElement("w:lvlJc")
        justify.set(qn("w:val"), "left")
        level.append(justify)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "280")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "40")
        spacing.set(qn("w:line"), "240")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        level.append(p_pr)
        if font:
            r_pr = OxmlElement("w:rPr")
            fonts = OxmlElement("w:rFonts")
            fonts.set(qn("w:ascii"), font)
            fonts.set(qn("w:hAnsi"), font)
            r_pr.append(fonts)
            level.append(r_pr)
        abstract.append(level)
        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        ref = OxmlElement("w:abstractNumId")
        ref.set(qn("w:val"), str(abstract_id))
        num.append(ref)
        numbering.append(num)

    add_definition(next_abstract, next_num, "bullet", "•", "Arial")
    add_definition(next_abstract + 1, next_num + 1, "decimal", "%1.")
    return next_num, next_num + 1


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline(paragraph, text: str, *, size: float = 10.5, color=BLACK, base_bold=False) -> None:
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, color=color, bold=base_bold)
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
        else:
            run = paragraph.add_run(token)
            set_run_font(run, size=size, color=color, bold=base_bold)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=10, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, placeholder, end])


def configure_header_footer(doc: Document, proposal_number: str) -> None:
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    text = f"Robert Ashworth d/b/a LumenCore | DLA26BZ03-NV011 | {proposal_number}"
    run = p.add_run(text)
    set_run_font(run, size=10, color=GRAY)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_page_number(p)


def parse_table(lines: list[str], start: int) -> tuple[list[str], list[list[str]], int]:
    headers = [cell.strip() for cell in lines[start].strip().strip("|").split("|")]
    index = start + 2
    rows: list[list[str]] = []
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        row = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if len(row) != len(headers):
            raise ValueError(f"Malformed Markdown table row {index + 1}: {lines[index]}")
        rows.append(row)
        index += 1
    return headers, rows, index


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for col, text in enumerate(headers):
        cell = table.rows[0].cells[col]
        cell._tc.get_or_add_tcPr().append(_cell_shading(LIGHT_GRAY))
        paragraph = cell.paragraphs[0]
        paragraph.style = doc.styles["Table Text"]
        add_inline(paragraph, text, size=10, base_bold=True)
    for row_data in rows:
        cells = table.add_row().cells
        for col, text in enumerate(row_data):
            paragraph = cells[col].paragraphs[0]
            paragraph.style = doc.styles["Table Text"]
            add_inline(paragraph, text, size=10)
    apply_table_geometry(table, table_widths(headers))
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    after.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _cell_shading(fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    return shading


def build_document(source: Path, output: Path, metadata_output: Path, proposal_number: str) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    bullet_num_id, decimal_num_id = create_numbering(doc)
    configure_header_footer(doc, proposal_number)

    title_seen = False
    body_started = False
    in_references = False
    index = 0
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[index + 1]):
            headers, rows, index = parse_table(lines, index)
            add_table(doc, headers, rows)
            continue

        if stripped.startswith("# "):
            if title_seen:
                raise ValueError("Only one level-1 title is supported")
            paragraph = doc.add_paragraph(style="Proposal Title")
            add_inline(paragraph, stripped[2:].strip(), size=18, base_bold=True)
            title_seen = True
            index += 1
            continue

        if stripped.startswith("## "):
            heading_text = stripped[3:].strip()
            paragraph = doc.add_paragraph(style="Heading 1")
            add_inline(paragraph, heading_text, size=12.5, color=BLUE, base_bold=True)
            body_started = True
            in_references = heading_text == "References"
            index += 1
            continue

        if stripped.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 2")
            add_inline(paragraph, stripped[4:].strip(), size=11.5, color=NAVY, base_bold=True)
            index += 1
            continue

        if stripped.startswith("#### "):
            paragraph = doc.add_paragraph(style="Heading 3")
            add_inline(paragraph, stripped[5:].strip(), size=10.5, color=NAVY, base_bold=True)
            index += 1
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            apply_numbering(paragraph, bullet_num_id)
            add_inline(paragraph, stripped[2:].strip())
            index += 1
            continue

        if re.match(r"^\d+\.\s+", stripped) and not in_references:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            apply_numbering(paragraph, decimal_num_id)
            add_inline(paragraph, re.sub(r"^\d+\.\s+", "", stripped))
            index += 1
            continue

        if in_references and re.match(r"^\d+\.\s+", stripped):
            paragraph = doc.add_paragraph(style="Reference")
            add_inline(paragraph, stripped, size=10)
            index += 1
            continue

        style = "Proposal Metadata" if title_seen and not body_started else "Normal"
        paragraph = doc.add_paragraph(style=style)
        add_inline(paragraph, stripped)
        index += 1

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.core_properties.title = "MissionWeave - DLA26BZ03-NV011 Phase I Technical Volume"
    doc.core_properties.subject = "DoW SBIR 2026 Release 3 Phase I"
    doc.core_properties.author = "Robert Ashworth"
    doc.core_properties.last_modified_by = "LumenCore proposal build"
    doc.core_properties.keywords = "MissionWeave, DLA, SBIR, digital twin, mission readiness"

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)

    proposal_number_is_neutral = proposal_number == NEUTRAL_PROPOSAL_HEADER
    metadata = {
        "schema": "missionweave_dsip_volume2_build_metadata.v2",
        "source": source.name,
        "source_sha256": sha256(source),
        "output": output.name,
        "output_sha256": sha256(output),
        "proposal_number_header_state": (
            "NEUTRAL_PLACEHOLDER" if proposal_number_is_neutral else "ASSIGNED_PRIVATE_VALUE"
        ),
        "proposal_number_value_exposed": False,
        "proposal_number_sha256": (
            None
            if proposal_number_is_neutral
            else hashlib.sha256(proposal_number.encode("utf-8")).hexdigest().upper()
        ),
        "design_tokens": PRESET,
    }
    metadata_output.write_text(json.dumps(metadata, indent=2) + "\n", encoding="utf-8")


def validate_cli_proposal_number(value: str) -> str:
    if value != NEUTRAL_PROPOSAL_HEADER:
        raise ValueError("ASSIGNED_PROPOSAL_NUMBER_REQUIRES_PRIVATE_FINALIZER")
    return value


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--metadata-output", type=Path, default=DEFAULT_METADATA)
    parser.add_argument(
        "--proposal-number",
        default=NEUTRAL_PROPOSAL_HEADER,
        help=(
            "Neutral placeholder only. Assigned proposal numbers must be read from the "
            "ignored private record by the guarded finalizer."
        ),
    )
    args = parser.parse_args()
    try:
        proposal_number = validate_cli_proposal_number(args.proposal_number)
    except ValueError as exc:
        parser.error(str(exc))
    build_document(args.source, args.output, args.metadata_output, proposal_number)
    print(args.output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
