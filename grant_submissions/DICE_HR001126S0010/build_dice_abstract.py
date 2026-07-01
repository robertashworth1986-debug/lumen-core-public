from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


HERE = Path(__file__).resolve().parent
TEMPLATE = HERE / "A1_DICE_Abstract_Template_OFFICIAL.docx"
OUTPUT = HERE / "LumenCore_DICE_Abstract_WORKING_DRAFT.docx"

TITLE = (
    "Coherence-Bounded Peer Mesh: Sparse Task Markets and Local Inference "
    "Control for Resilient Heterogeneous AI Collectives"
)

INK = RGBColor(0x00, 0x00, 0x00)
BLUE = RGBColor(0x1F, 0x4E, 0x78)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_BLUE = "D9EAF7"

CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CORE_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"
DCTERMS_NS = "http://purl.org/dc/terms/"
XSI_NS = "http://www.w3.org/2001/XMLSchema-instance"

ET.register_namespace("cp", CORE_NS)
ET.register_namespace("dc", DC_NS)
ET.register_namespace("dcterms", DCTERMS_NS)
ET.register_namespace("xsi", XSI_NS)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is not None:
        return
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _remove_package_relationships(xml_bytes: bytes) -> bytes:
    root = ET.fromstring(xml_bytes)
    changed = False
    for rel in list(root):
        rel_type = rel.attrib.get("Type", "")
        target = rel.attrib.get("Target", "")
        target_norm = target.replace("\\", "/")
        remove = (
            rel_type.endswith("/comments")
            or rel_type.endswith("/commentsExtended")
            or rel_type.endswith("/commentsExtensible")
            or rel_type.endswith("/commentsIds")
            or rel_type.endswith("/people")
            or rel_type.endswith("/classificationlabels")
            or rel_type.endswith("/customXml")
            or rel_type.endswith("/custom-properties")
            or target_norm in {
                "comments.xml",
                "commentsExtended.xml",
                "commentsExtensible.xml",
                "commentsIds.xml",
                "people.xml",
                "docMetadata/LabelInfo.xml",
            }
            or target_norm.startswith("customXml/")
            or target_norm.startswith("../customXml/")
            or target_norm == "docProps/custom.xml"
        )
        if remove:
            root.remove(rel)
            changed = True
    return ET.tostring(root, encoding="utf-8", xml_declaration=True) if changed else xml_bytes


def _remove_content_type_overrides(xml_bytes: bytes) -> bytes:
    root = ET.fromstring(xml_bytes)
    changed = False
    for override in list(root):
        part_name = override.attrib.get("PartName", "")
        if (
            part_name
            in {
                "/word/comments.xml",
                "/word/commentsExtended.xml",
                "/word/commentsExtensible.xml",
                "/word/commentsIds.xml",
                "/word/people.xml",
                "/docMetadata/LabelInfo.xml",
                "/docProps/custom.xml",
            }
            or part_name.startswith("/customXml/")
        ):
            root.remove(override)
            changed = True
    return ET.tostring(root, encoding="utf-8", xml_declaration=True) if changed else xml_bytes


def _set_text(root: ET.Element, tag: str, value: str) -> None:
    node = root.find(tag)
    if node is None:
        node = ET.SubElement(root, tag)
    node.text = value


def _clean_core_properties(xml_bytes: bytes) -> bytes:
    root = ET.fromstring(xml_bytes)
    _set_text(root, f"{{{DC_NS}}}creator", "Robert Ashworth")
    _set_text(root, f"{{{CORE_NS}}}lastModifiedBy", "LumenCore")
    _set_text(root, f"{{{CORE_NS}}}revision", "1")
    last_printed = root.find(f"{{{CORE_NS}}}lastPrinted")
    if last_printed is not None:
        root.remove(last_printed)
    _set_text(root, f"{{{DCTERMS_NS}}}created", "2026-06-19T00:00:00Z")
    created = root.find(f"{{{DCTERMS_NS}}}created")
    if created is not None:
        created.set(f"{{{XSI_NS}}}type", "dcterms:W3CDTF")
    _set_text(root, f"{{{DCTERMS_NS}}}modified", "2026-06-19T00:00:00Z")
    modified = root.find(f"{{{DCTERMS_NS}}}modified")
    if modified is not None:
        modified.set(f"{{{XSI_NS}}}type", "dcterms:W3CDTF")
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def scrub_template_artifacts(docx_path: Path) -> None:
    """Remove hidden template review/SharePoint artifacts after generation."""
    remove_names = {
        "word/comments.xml",
        "word/commentsExtended.xml",
        "word/commentsExtensible.xml",
        "word/commentsIds.xml",
        "word/people.xml",
        "docMetadata/LabelInfo.xml",
        "docProps/custom.xml",
    }
    tmp_path = docx_path.with_suffix(".scrubbed.tmp.docx")
    with ZipFile(docx_path, "r") as src, ZipFile(tmp_path, "w", ZIP_DEFLATED) as dst:
        for item in src.infolist():
            name = item.filename
            if name in remove_names or name.startswith("customXml/"):
                continue
            data = src.read(name)
            if name == "[Content_Types].xml":
                data = _remove_content_type_overrides(data)
            elif name.endswith(".rels"):
                data = _remove_package_relationships(data)
            elif name == "docProps/core.xml":
                data = _clean_core_properties(data)
            dst.writestr(item, data)
    tmp_path.replace(docx_path)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.0

    for style_name, size, before, after in (
        ("Heading 1", 14, 12, 6),
        ("Heading 2", 12, 8, 4),
        ("Heading 3", 12, 6, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        p_pr = style.element.get_or_add_pPr()
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is not None:
            p_pr.remove(num_pr)


def set_run(run, *, bold=False, italic=False, size=12, color=INK) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")


def add_body_paragraph(doc: Document, text: str, *, bold_lead: str | None = None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run(lead, bold=True)
        remainder = p.add_run(text[len(bold_lead) :])
        set_run(remainder)
    else:
        run = p.add_run(text)
        set_run(run)
    return p


def ensure_bullet_numbering(doc: Document) -> int:
    existing = getattr(doc, "_dice_bullet_num_id", None)
    if existing is not None:
        return existing

    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract_num.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "\u2022")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "280")
    p_pr.extend([tabs, indent])
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Symbol")
    fonts.set(qn("w:hAnsi"), "Symbol")
    r_pr.append(fonts)
    level.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    abstract_num.append(level)
    numbering.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    doc._dice_bullet_num_id = num_id
    return num_id


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), str(ensure_bullet_numbering(doc)))
    num_pr.extend([ilvl, num_id])
    p._p.get_or_add_pPr().append(num_pr)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    set_run(p.add_run(text))
    return p


def add_heading(doc: Document, text: str, level=1):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    size = 14 if level == 1 else 12
    before = 12 if level == 1 else 8
    after = 6 if level == 1 else 4
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    set_run(p.add_run(text), bold=True, size=size, color=BLUE)
    return p


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])
    set_run(run, size=9, color=GRAY)


def clear_document_body(doc: Document) -> None:
    body = doc._element.body
    section_properties = body.sectPr
    for child in list(body):
        if child is not section_properties:
            body.remove(child)


def set_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_run(p.add_run("COVER SHEET"), bold=True, size=16, color=BLUE)

    warning = doc.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning.paragraph_format.space_after = Pt(10)
    set_run(
        warning.add_run("WORKING DRAFT - NOT APPROVED FOR SUBMISSION"),
        bold=True,
        size=10,
        color=RGBColor(0x9B, 0x1C, 0x1C),
    )

    rows = [
        ("Abstract Title", TITLE),
        ("Technical Area", "TA1 & TA2"),
        ("Proposer Organization", "Robert Ashworth d/b/a LumenCore"),
        (
            "Technical Point of Contact (POC)",
            "Robert Ashworth\n2613 Paddle Wheel Dr, Nashville, TN 37214\n"
            "615-438-2502\nrobertashworth4444@gmail.com",
        ),
        (
            "Administrative POC",
            "Robert Ashworth\n2613 Paddle Wheel Dr, Nashville, TN 37214\n"
            "615-438-2502\nrobertashworth4444@gmail.com",
        ),
        (
            "Other Team Members (subcontractors and consultants), if known/applicable",
            "Not finalized at abstract stage. Planned recruitment/teaming targets: "
            "distributed-systems research lead; inference-control/mechanistic-"
            "interpretability lead; cloud/HPC engineering support; independent "
            "red-team and evaluation advisor.",
        ),
        ("Award Instrument Requested", "OT for Research"),
        (
            "Estimated Total Cost (Base + Options)",
            "$4,920,000 abstract-stage ROM planning estimate; validate before full proposal",
        ),
        ("Estimated Period of Performance", "36 months; planning start February 1, 2027"),
        (
            "Identify any other solicitation(s) to which this concept has been proposed.",
            "None. Related anomaly-detection work is being prepared for a distinct Navy "
            "SBIR topic and does not propose this decentralized-agent concept.",
        ),
        ("SAM.gov Unique Entity ID (UEI)", "SQY2XW71ZM51"),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.05)
    table.columns[1].width = Inches(4.45)
    for idx, (label, value) in enumerate(rows):
        left, right = table.rows[idx].cells
        left.width = Inches(2.05)
        right.width = Inches(4.45)
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(left, LIGHT_BLUE)
        set_cell_margins(left)
        set_cell_margins(right)
        lp = left.paragraphs[0]
        rp = right.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        rp.paragraph_format.space_after = Pt(0)
        set_run(lp.add_run(label), bold=True, size=10)
        set_run(rp.add_run(value), size=10)

    doc.add_page_break()


def add_toc(doc: Document) -> None:
    add_heading(doc, "Table of Contents", 1)
    entries = [
        "1. Goals and Impact",
        "2. Technical Approach",
        "3. Capabilities/Management Plan",
        "4. Cost and Schedule",
        "5. Publications",
        "6. Bibliography",
    ]
    for entry in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        set_run(p.add_run(entry))
    doc.add_page_break()


def add_goals(doc: Document) -> None:
    add_heading(doc, "1. Goals and Impact", 1)
    add_body_paragraph(
        doc,
        "LumenCore proposes Coherence-Bounded Peer Mesh (CBPM), a local adaptor "
        "architecture for heterogeneous open-weight and black-box AI agents. CBPM "
        "replaces a persistent central orchestrator with sparse peer task markets, "
        "reputation-weighted context fusion, and a local inference-control loop that "
        "exposes measurable role-coherence bounds to the coordination layer. "
        "The research addresses TA1 and TA2 together because planning depth is limited "
        "by the period over which each agent can remain aligned to its assigned role.",
    )
    add_body_paragraph(
        doc,
        "The central hypothesis is that agents can coordinate through bounded local "
        "capability advertisements and task auctions while each adaptor continuously "
        "measures role drift, mission alignment, uncertainty, and behavioral diversity. "
        "When the controller detects drift or compromised behavior, it can repair "
        "context, constrain tools, challenge a commitment, reduce reputation, re-auction "
        "the task, or isolate the agent without globally replanning the mission.",
    )
    add_body_paragraph(
        doc,
        "If successful, CBPM will improve the number of agents and mission interactions "
        "that can be coordinated, reduce messages required to recover from node failure "
        "or deception, sustain role coherence over longer interaction horizons, and "
        "preserve multiple viable courses of action. The target progression is the DICE "
        "scale of 500 agents and 5,000 messages in Phase 1, 5,000 agents and 50,000 "
        "messages in Phase 2, and 100,000 agents and 1 million messages in Phase 3.",
    )
    add_body_paragraph(
        doc,
        "The innovation is not a new orchestration wrapper or a static system prompt. "
        "It is a coupled coordination-control mechanism: TA1 consumes explicit "
        "coherence-horizon and confidence bounds from TA2; TA2 derives control targets "
        "from the current task contract, local dependencies, and mission constraints. "
        "Every local decision produces an auditable event record suitable for failure "
        "analysis without requiring a central authority to approve normal coordination.",
    )


def add_technical_approach(doc: Document) -> None:
    add_heading(doc, "2. Technical Approach", 1)
    add_heading(doc, "2.1 Local Adaptor and Interfaces", 2)
    add_body_paragraph(
        doc,
        "Each agent receives a CBPM adaptor with four modules: (1) a capability and "
        "constraint descriptor; (2) a sparse peer task-market client; (3) a "
        "reputation-weighted context and commitment store; and (4) a local inference "
        "controller. The common adaptor interface accepts mission fragments, proposed "
        "task contracts, evidence packets, role constraints, and tool policies, and "
        "returns bids, commitments, confidence, coherence bounds, execution evidence, "
        "and challenge responses.",
    )
    add_body_paragraph(
        doc,
        "For black-box agents, the controller operates through context engineering, "
        "structured memory, tool permissions, response contracts, counterfactual "
        "self-checks, and local critique. For open-weight agents, the same interface "
        "adds optional representation and activation monitors. The adaptor will support "
        "A2A- and MCP-compatible transports while keeping the coordination algorithm "
        "independent of any single agent framework.",
    )

    add_heading(doc, "2.2 TA1: Sparse Task Markets and Resilient Context Fusion", 2)
    for text in [
        "Mission propagation and decomposition. Agents advertise compact, signed "
        "capability summaries to bounded neighborhoods. A receiving agent may accept a "
        "mission fragment, decompose it, or forward an unclaimed remainder. Task "
        "contracts carry prerequisites, evidence requirements, deadlines, expected "
        "message cost, and the minimum TA2 coherence horizon.",
        "Local auction and consensus. Candidate agents bid using capability fit, current "
        "load, evidence quality, expected cost, and locally maintained reputation. "
        "Selection uses a quorum of independent neighborhood observations rather than a "
        "global leader. The research will compare deterministic, probabilistic, and "
        "game-theoretic selection rules.",
        "Failure and compromise response. Failed commitments trigger local challenge and "
        "re-auction. Contradictory evidence is retained as competing hypotheses with "
        "source-specific confidence. Agents that repeatedly provide unverifiable, "
        "strategically timed, or mutually inconsistent commitments are discounted or "
        "isolated. Limits will be measured as a function of compromised fraction, "
        "collusion topology, communication delay, and neighborhood degree.",
        "Sparse adaptation. Capability caches and dependency-local replanning target "
        "communication growth from an initial bounded neighborhood regime toward "
        "O(n log n) and O(n) message behavior as scale increases. The proposal will "
        "measure mission success, coordination stability, message count, and recovery "
        "messages rather than claiming asymptotic performance from implementation form.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "2.3 TA2: Coherence-Bounded Local Inference Control", 2)
    for text in [
        "Role contract. Each task establishes an explicit role vector containing allowed "
        "objectives, prohibited actions, information boundaries, required evidence, "
        "dependencies, and expiration conditions.",
        "Runtime monitor. The adaptor estimates role drift, goal drift, context conflict, "
        "uncertainty, and behavioral-collapse risk at each inference step. Black-box "
        "agents are measured through observable commitments and counterfactual probes; "
        "open-weight agents may additionally expose activation-space signals.",
        "Controller. Interventions escalate from context repair and memory pruning to "
        "tool restriction, commitment challenge, role reassignment, or local isolation. "
        "The controller reports a finite coherence horizon and confidence to TA1 instead "
        "of asserting permanent alignment.",
        "Cognitive agility. The evaluation will preserve multiple admissible courses of "
        "action and quantify diversity, novelty, and mission success. A controller that "
        "maximizes role conformity by collapsing every agent to one answer will fail the "
        "agility requirement.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "2.4 Preliminary Synthetic Feasibility Evidence", 2)
    add_body_paragraph(
        doc,
        "A deterministic discrete-event harness was frozen on June 13, 2026 to test the "
        "measurement plan before proposal drafting. Across 200 paired seeds, each trial "
        "used 500 stochastic task executors, 1,200 tasks, 12 roles, 10% failed agents, "
        "and 5% compromised agents. Compared with a centralized assignment baseline, a "
        "cached-capability peer auction with local reputation and role repair maintained "
        "essentially identical mission completion while reducing mean coordination "
        "messages by 31.5% (95% bootstrap CI 31.3%-31.7%), reducing recovery messages by "
        "40.1% (39.7%-40.5%), and improving observed role coherence by 6.45 percentage "
        "points (6.24-6.69).",
    )
    add_body_paragraph(
        doc,
        "Scale stress checks then executed 50 paired trials with 5,000 agents and "
        "15,000 tasks, and five paired trials with 100,000 agents and 250,000 tasks. "
        "At 5,000 agents, the peer protocol reduced messages by 31.8% and recovery "
        "messages by 33.4%; at 100,000 agents, it reduced messages by 30.1% and "
        "recovery messages by 38.3%. Mean mission-success differences were near zero "
        "and their bootstrap intervals crossed zero at both scales. The 100,000-agent "
        "peer run averaged approximately 734,102 messages, below the DICE notional "
        "one-million-message scale, but with simplified stochastic executors.",
    )
    add_body_paragraph(
        doc,
        "A disjoint constraint-contract ablation then compared peer reputation with "
        "bids carrying explicit role, horizon, evidence-lineage, expiration, and risk "
        "fields. Across 30 validation scenarios per condition with 500 agents and 1,200 "
        "tasks, the contract method improved safe completion by 4.16 percentage points "
        "in benign conditions (95% paired bootstrap CI 3.84-4.46), 3.52 points under "
        "10% independent compromise (3.04-3.98), 3.50 points under 10% collusion "
        "(3.05-3.96), 2.10 points under monitor shift (1.69-2.49), and 1.35 points under "
        "25% high collusion (0.95-1.75). Modeled constraint violations fell by 5.66-7.24 "
        "points and messages per safe completion fell by 1.07-1.81. The tradeoff was a "
        "2.5-4.5 point reduction in raw completion from deterministic rejection, with "
        "false rejection reaching 11.38% under monitor shift. Under high collusion, "
        "compromised assignments increased by 0.96 points because locally consistent "
        "forgeries passed field checks. Contract fields were piggybacked on existing "
        "messages; byte, latency, and cryptographic costs were not measured.",
    )
    add_body_paragraph(
        doc,
        "Evidence boundary: these agents are stochastic task executors, not language "
        "models. The result demonstrates only that the proposed metrics and a narrow "
        "protocol hypothesis can be evaluated reproducibly. It does not establish DICE "
        "program metric attainment, foundation-model inference at these scales, "
        "operational performance, or adversarial security. Inputs, source commit, trial "
        "CSVs, scorecards, and SHA-256 manifests are retained for audit.",
        bold_lead="Evidence boundary:",
    )

    add_heading(doc, "2.5 Experimental Plan and Risks", 2)
    for text in [
        "Phase 1: implement the adaptor, establish centralized and decentralized SOTA "
        "baselines, test 500 heterogeneous agents, and demonstrate integration with the "
        "TA3 interface. Preregister mission, message, recovery, coherence, alignment, "
        "agility, latency, and cost metrics.",
        "Phase 2: add Byzantine behavior, delayed and deceptive context, collusion, prompt "
        "infection, strategic failure timing, and adaptive reputation attacks. Estimate "
        "breakdown points and false-isolation rates rather than assuming perfect threat "
        "detection.",
        "Phase 3: move to event-driven inference, hierarchical sparse neighborhoods, and "
        "100,000-agent scale. Demonstrate recovery and mission continuity under the "
        "TA3 team-versus-team evaluation.",
        "Primary risks are communication cascades, colluding agents manipulating local "
        "reputation, controllers suppressing useful novelty, non-transfer of open-weight "
        "signals to black-box models, and compute cost. Mitigations include bounded "
        "neighborhoods, independent evidence challenges, diversity floors, observable-"
        "behavior fallbacks, event-driven inference, small reasoning models, and staged "
        "scale gates.",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "2.6 Existing Capabilities, Repositories, and Compute", 2)
    add_body_paragraph(
        doc,
        "Existing software capabilities include a Python/FastAPI human-approval hub, "
        "centralized multi-engine orchestration, fail-closed runtime controls, streaming "
        "anomaly detection, walk-forward benchmarking, and SHA-256 evidence manifests. "
        "A public demonstration and evidence portal is available at "
        "https://lumen-core.ai and the public repository is available at "
        "https://github.com/robertashworth1986-debug/lumen-core-public with MIT-licensed "
        "public code; proposal-specific adaptor research and data-rights assertions "
        "will be identified in the full proposal.",
    )
    add_body_paragraph(
        doc,
        "Current local compute is an HP OMEN laptop with an Intel Core Ultra 9 185H "
        "(16 cores/22 logical processors), approximately 32 GB RAM, and an NVIDIA RTX "
        "4070 Laptop GPU. This is sufficient for development, discrete-event simulation, "
        "small-model experiments, and interface testing but not for program-scale "
        "foundation-model inference. The cost plan therefore includes elastic cloud GPU "
        "capacity, object storage, experiment tracking, and checkpointed event-driven "
        "simulation. The team will benchmark vLLM or equivalent efficient inference "
        "stacks and use smaller reasoning models where fidelity permits.",
    )


def add_management(doc: Document) -> None:
    add_heading(doc, "3. Capabilities/Management Plan", 1)
    add_body_paragraph(
        doc,
        "Robert Ashworth will serve as Principal Investigator and system architect. He "
        "has built the existing orchestration, evidence, validation, deployment, and "
        "human-approval components and will own adaptor architecture, experimental "
        "reproducibility, and integration. The current organization is a one-person "
        "small business; this is a material execution risk, not a hidden assumption.",
    )
    add_body_paragraph(
        doc,
        "Before a full proposal, LumenCore will seek named commitments for three "
        "complementary roles: distributed systems and Byzantine consensus; inference-"
        "time control and mechanistic interpretability; and scalable cloud/HPC "
        "simulation. An independent red-team/evaluation advisor will review adversarial "
        "design and claim boundaries. Teaming arrangements are not asserted as final in "
        "this abstract.",
    )
    add_body_paragraph(
        doc,
        "The program will use phase-gated technical reviews, frozen experiment manifests, "
        "paired baselines, independent reruns, and a risk register tied to measurable exit "
        "criteria. External submissions, release of proprietary artifacts, and changes to "
        "security posture remain human-approved. TA1 and TA2 work is proposed at the "
        "Unclassified level. The full proposal will select an award vehicle and "
        "information boundary consistent with actual FCI/CUI handling and verified CMMC "
        "status.",
    )


def add_cost(doc: Document) -> None:
    add_heading(doc, "4. Cost and Schedule", 1)
    add_body_paragraph(
        doc,
        "The following is an abstract-stage ROM planning estimate, not a certified cost "
        "proposal or reviewed cost volume. The provisional basis is 21,060 direct labor "
        "hours at a $100/hour blended direct rate, 25% fringe, 28% combined overhead/G&A "
        "on labor plus fringe, $350,000 of specialized subaward/consultant effort, "
        "$1,040,000 of cloud/HPC, $80,000 of travel, and $80,000 of software/data/"
        "equipment. Rates, indirect costs, subaward scopes, resource-sharing terms, and "
        "vendor quotations must be validated before any full-proposal cost submission.",
    )

    rows = [
        ("Phase", "Months", "Program focus", "Total", "Compute included"),
        ("1", "1-9", "Adaptor, baselines, 500-agent evaluation", "$1,050,000", "$180,000"),
        ("2", "10-24", "Adversarial robustness and team playoffs", "$2,050,000", "$420,000"),
        ("3", "25-36", "100,000-agent scaling and transition", "$1,820,000", "$440,000"),
        ("Total", "36", "Base plus options", "$4,920,000", "$1,040,000"),
    ]
    table = doc.add_table(rows=len(rows), cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.65, 0.65, 3.0, 1.1, 1.1]
    for row_idx, values in enumerate(rows):
        for col_idx, value in enumerate(values):
            cell = table.rows[row_idx].cells[col_idx]
            cell.width = Inches(widths[col_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=60, bottom=60, start=70, end=70)
            if row_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            p = cell.paragraphs[0]
            p.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if col_idx == 2 else WD_ALIGN_PARAGRAPH.CENTER
            )
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(value), bold=row_idx in (0, len(rows) - 1), size=9)
    set_repeat_table_header(table.rows[0])

    add_body_paragraph(
        doc,
        "Major milestones: M3 adaptor prototype; M4 self-evaluation; M6 baseline "
        "integration; M8 Phase 1 comparison; M13/16/19 adversarial evaluations; M23 "
        "down-select; M30 scaling evaluation; M33 self-play; M35 transition demo.",
    )


def add_publications(doc: Document) -> None:
    add_heading(doc, "5. Publications", 1)
    reports = [
        "Ashworth, Robert. \"DICE Preliminary Synthetic Benchmark V1: Sparse Peer Auction "
        "with Local Role-Coherence Control.\" Technical report and SHA-256 manifest, 2026.",
        "Ashworth, Robert. \"DICE Constraint-Carrying Commitment Benchmark: Role, "
        "Evidence-Lineage, Expiration, and Coherence-Horizon Ablation.\" Generated "
        "discrete-event technical report and SHA-256 manifest, 2026.",
        "Ashworth, Robert. \"HarborSentinel Synthetic Benchmark: Explainable Streaming "
        "Maritime Anomaly Detection.\" Technical report and reproducibility package, 2026.",
        "Ashworth, Robert. \"Harmonic Validation Protocol: Leakage-Resistant Walk-Forward "
        "Evidence and Frozen Delta Requirements.\" Technical report, 2026.",
        "Ashworth, Robert. \"LumenCore Platform, Proof, and Commercialization Map.\" "
        "Technical report, 2026.",
        "Ashworth, Robert. \"Luma Production Context: Evidence-Calibrated Multi-Horizon "
        "Decision Platform.\" Technical report, 2026.",
    ]
    for report in reports:
        add_bullet(doc, report)


def add_bibliography(doc: Document) -> None:
    add_heading(doc, "6. Bibliography", 1)
    refs = [
        "Defense Advanced Research Projects Agency. (2026). HR001126S0010, "
        "Decentralized Artificial Intelligence through Controlled Emergence (DICE). "
        "https://www.darpa.mil/research/programs/"
        "decentralized-artificial-intelligence-through-controlled-emergence",
        "Turner, A. M., et al. (2023). Steering Language Models with Activation "
        "Engineering. arXiv:2308.10248. https://arxiv.org/abs/2308.10248",
        "Shapira, N., et al. (2024). Prompt Infection: LLM-to-LLM Prompt Injection "
        "within Multi-Agent Systems. arXiv:2410.07283. "
        "https://arxiv.org/abs/2410.07283",
        "Friston, K. J., et al. (2024). Designing Ecosystems of Intelligence from First "
        "Principles. Collective Intelligence, 3(1). "
        "https://doi.org/10.1177/26339137231222481",
        "Zhou, H., et al. (2025). ReSo: A Reward-Driven Self-Organizing LLM-Based "
        "Multi-Agent System for Reasoning Tasks. Proceedings of EMNLP 2025. "
        "https://aclanthology.org/2025.emnlp-main.808/",
        "Lee, H., et al. (2026). Robust Multi-Agent LLMs under Byzantine Faults. "
        "arXiv:2605.09076. https://arxiv.org/abs/2605.09076",
        "Model Context Protocol project. (2026). Open protocol repository. "
        "https://github.com/modelcontextprotocol",
        "Agent2Agent Protocol project. (2026). Open protocol repository. "
        "https://github.com/a2aproject",
        "vLLM project. (2026). vLLM documentation. https://docs.vllm.ai/",
        "Fujimoto, R. M. (1990). Parallel Discrete Event Simulation. Communications "
        "of the ACM, 33(10), 30-53. https://doi.org/10.1145/84537.84545",
    ]
    for ref in refs:
        add_bullet(doc, ref)


def main() -> None:
    doc = Document(TEMPLATE)
    clear_document_body(doc)
    configure_styles(doc)
    for section in doc.sections:
        set_section(section)

    add_cover(doc)
    add_toc(doc)
    add_goals(doc)
    add_technical_approach(doc)
    add_management(doc)
    add_cost(doc)
    add_publications(doc)
    add_bibliography(doc)

    for section in doc.sections:
        set_section(section)
        footer = section.footer
        footer.is_linked_to_previous = True
        p = footer.paragraphs[0]
        p.clear()
        left = p.add_run("UNCLASSIFIED | WORKING DRAFT - NOT APPROVED FOR SUBMISSION")
        set_run(left, size=8, color=GRAY)
        p.add_run(" " * 8)
        add_page_number(p)

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "DARPA DICE HR001126S0010 TA1/TA2 Abstract"
    doc.core_properties.author = "Robert Ashworth"
    doc.core_properties.comments = (
        "Working draft generated from the official DICE abstract template. "
        "Requires human review, cost validation, teaming, and compliance verification."
    )
    doc.save(OUTPUT)
    scrub_template_artifacts(OUTPUT)
    # Normalize the package after template scrubbing so LibreOffice can render it
    # without retaining stale relationship state from the official DOCX template.
    Document(OUTPUT).save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
