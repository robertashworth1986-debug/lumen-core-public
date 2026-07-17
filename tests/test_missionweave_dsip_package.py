from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
PACKAGE = ROOT / "grant_submissions" / "DLA26BZ03_NV011_MissionWeave"
VOLUME1 = PACKAGE / "MISSIONWEAVE_DSIP_VOLUME1_PUBLIC_TEXT_2026-07-16.md"
VOLUME2_MD = PACKAGE / "MISSIONWEAVE_DSIP_VOLUME2_FINAL_CANDIDATE_2026-07-16.md"
VOLUME2_DOCX = PACKAGE / "MISSIONWEAVE_DSIP_VOLUME2_FINAL_CANDIDATE_2026-07-16.docx"
COST = PACKAGE / "MISSIONWEAVE_DSIP_VOLUME3_COST_INPUTS_2026-07-16.md"
MANIFEST = PACKAGE / "MISSIONWEAVE_DSIP_PACKAGE_MANIFEST_2026-07-16.json"


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    digest.update(path.read_bytes())
    return digest.hexdigest().upper()


def _volume1_sections() -> tuple[str, str]:
    text = VOLUME1.read_text(encoding="utf-8")
    abstract = re.search(
        r"## Technical Abstract\s+(.*?)\s+## Anticipated Benefits", text, re.DOTALL
    ).group(1).strip()
    benefits = re.search(
        r"## Anticipated Benefits\s+(.*?)\s+## DSIP Entry Checks", text, re.DOTALL
    ).group(1).strip()
    return abstract, benefits


def test_official_source_hashes_are_locked() -> None:
    expected = {
        "DLA_26BZ_RELEASE_3_COMPONENT_INSTRUCTIONS.pdf":
            "17B06B6FE3DBE6F2035F8287A0DD01FDBFE180858F3C490F1E91D121819239A2",
        "DLA26BZ03_NV011_OFFICIAL_TOPIC_DETAILS.json":
            "4D01354A456BF91BAFF5AB638311BD5822176B49AD7911F6DEC1C379B3B217B7",
        "DoW_2026_SBIR_BAA_RELEASE_3_AMENDMENT_2.pdf":
            "FB809186F3B43313A1877A5C36E2383A45B84F62566DDA3F6BCDB93779CE08AE",
    }
    source_dir = PACKAGE / "source_attachments"
    for name, digest in expected.items():
        assert _sha256(source_dir / name) == digest


def test_volume1_public_entries_fit_dsip_limits_and_stay_nonproprietary() -> None:
    abstract, benefits = _volume1_sections()
    assert 1000 <= len(abstract) <= 3000
    assert 700 <= len(benefits) <= 3000
    combined = f"{abstract}\n{benefits}".lower()
    forbidden = [
        "dla validated",
        "field proven",
        "production ready",
        "10x improvement achieved",
        "cmmc compliant",
        "itar compliant",
        "has realized savings",
        "current dla customer",
    ]
    assert not any(phrase in combined for phrase in forbidden)
    assert "No customer commitment, realized savings, or production deployment is claimed" in benefits


def test_volume2_uses_required_section_order() -> None:
    text = VOLUME2_MD.read_text(encoding="utf-8")
    required = [
        "## 1. Identification and Significance of the Problem or Opportunity",
        "## 2. Phase I Technical Objectives",
        "## 3. Phase I Statement of Work",
        "## 4. Related Work",
        "## 5. Relationship with Future Research or Research and Development",
        "## 6. Commercialization Strategy",
        "## 7. Key Personnel",
        "## 8. Foreign Citizens",
        "## 9. Facilities and Equipment",
        "## 10. Subcontractors and Consultants",
        "## 11. Prior, Current, or Pending Support for Similar Proposals or Awards",
        "## 12. Technical Data and Software Rights Assertions",
    ]
    positions = [text.index(heading) for heading in required]
    assert positions == sorted(positions)


def test_volume2_preserves_exact_benchmark_boundary_and_negative_evidence() -> None:
    text = VOLUME2_MD.read_text(encoding="utf-8")
    required_fragments = [
        "+0.0578",
        "+0.1156",
        "+0.1175",
        "+0.1266",
        "+0.0302",
        "25 / 4 / 1",
        "24 / 1 / 5",
        "25 / 3 / 2",
        "28 / 0 / 2",
        "23 / 0 / 7",
        "0.240 for cross-trained FIFO and 0.270 for MissionWeave",
        "no 10x claim",
        "BD5FB806A6F524DE2E60D48E4D091D916F86B35B2FD73E3889667B2D8B2385DB",
    ]
    for fragment in required_fragments:
        assert fragment in text


def test_volume2_has_no_submission_placeholders_or_unsupported_escalation() -> None:
    text = VOLUME2_MD.read_text(encoding="utf-8")
    assert "TBD" not in text
    assert "TODO" not in text
    assert "[confirm" not in text.lower()
    forbidden = [
        "DLA-validated",
        "DLA validated",
        "field-proven",
        "production-ready",
        "CMMC compliant",
        "ITAR compliant",
        "guaranteed ROI",
        "achieved 10x",
        "current DLA customer",
    ]
    for phrase in forbidden:
        assert phrase not in text


def test_cost_model_reconciles_to_ceiling_and_all_prime_work_share() -> None:
    text = COST.read_text(encoding="utf-8")
    amounts = [60800, 12160, 18240, 4000, 2000, 2800, 0, 0]
    assert sum(amounts) == 100000
    for amount in amounts[:-2]:
        assert f"${amount:,}" in text
    assert "$100,000" in text
    assert "91.2%" in text
    assert "640 hours" in text
    assert "None proposed" in text
    assert "Not requested" in text


def test_docx_page_geometry_header_fonts_and_table_geometry() -> None:
    doc = Document(VOLUME2_DOCX)
    section = doc.sections[0]
    assert round(section.page_width.inches, 3) == 8.5
    assert round(section.page_height.inches, 3) == 11.0
    for margin in (section.top_margin, section.right_margin, section.bottom_margin, section.left_margin):
        assert round(margin.inches, 3) == 1.0
    header_text = " ".join(p.text for p in section.header.paragraphs)
    assert "Robert Ashworth d/b/a LumenCore" in header_text
    assert "DLA26BZ03-NV011" in header_text
    assert "Proposal No. assigned in DSIP" in header_text

    for style_name, minimum in (("Normal", 10), ("Heading 1", 10), ("Heading 2", 10), ("Table Text", 10)):
        assert doc.styles[style_name].font.name == "Arial"
        assert doc.styles[style_name].font.size.pt >= minimum

    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        assert tbl_pr.find(qn("w:tblW")).get(qn("w:w")) == "9360"
        assert tbl_pr.find(qn("w:tblInd")).get(qn("w:w")) == "120"
        grid_widths = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
        assert sum(grid_widths) == 9360
        for row in table.rows:
            cell_widths = [int(cell._tc.get_or_add_tcPr().find(qn("w:tcW")).get(qn("w:w"))) for cell in row.cells]
            assert cell_widths == grid_widths


def test_package_manifest_verifies_when_present() -> None:
    if not MANIFEST.exists():
        return
    manifest = json.loads(MANIFEST.read_text(encoding="utf-8"))
    assert manifest["topic"] == "DLA26BZ03-NV011"
    assert manifest["file_count"] == len(manifest["files"])
    for item in manifest["files"]:
        path = PACKAGE / item["path"]
        assert path.stat().st_size == item["bytes"]
        assert _sha256(path) == item["sha256"]
