from __future__ import annotations

import importlib.util
import json
import sys
from pathlib import Path

import pytest
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
BUILDER_PATH = ROOT / "code" / "ops" / "BUILD_PRIVATE_MASTER_WHITEPAPER.py"
CORPUS_PATH = (
    ROOT / "code" / "ops" / "BUILD_PRIVATE_MASTER_WHITEPAPER_CORPUS.py"
)


def load_module(name: str, path: Path):
    spec = importlib.util.spec_from_file_location(name, path)
    assert spec and spec.loader
    module = importlib.util.module_from_spec(spec)
    sys.modules[name] = module
    spec.loader.exec_module(module)
    return module


builder = load_module("private_master_whitepaper", BUILDER_PATH)
corpus = load_module("private_master_whitepaper_corpus", CORPUS_PATH)


def test_corpus_covers_current_research_and_presentation_governance() -> None:
    source_ids = {source.source_id for source in corpus.SOURCES}

    assert "whitehole_whiteholelab_audit" in source_ids
    assert "pitch_deck_governance" in source_ids
    assert builder.DISPOSITIONS["whitehole_whiteholelab_audit"] == {
        "disposition": "HISTORICAL_PROVENANCE_CURRENT_CLAIM_BOUNDARY",
        "claim_class": "implemented_governance",
        "use": "Body and source ledger: preserve custody history while excluding heuristic ranks and the legacy site from current evidence.",
    }
    assert builder.DISPOSITIONS["pitch_deck_governance"] == {
        "disposition": "CURRENT_PRESENTATION_RELEASE_CONTROL",
        "claim_class": "implemented_governance",
        "use": "Body and source ledger: bind the single current review deck and quarantine stale or high-risk presentation artifacts.",
    }


def test_privacy_scan_does_not_treat_long_count_sentence_as_street_address() -> None:
    sentence = (
        "126 executed direct candidate-source-baseline comparisons. No individual "
        "comparison is positive after the current global Holm correction and no "
        "candidate passes the complete promotion gate. A distinct EIA hourly lane"
    )

    builder.audit_private_output(sentence)
    assert builder.PRIVACY_PATTERNS["street_address"].search(sentence) is None
    assert "street_address" not in corpus.flag_text(sentence)
    short_sentence = "126 predictions remain in this lane"
    builder.audit_private_output(short_sentence)
    assert "street_address" not in corpus.flag_text(short_sentence)


def test_privacy_scan_still_detects_a_plausible_street_address() -> None:
    address = "Reviewer correspondence should be sent to 123 Main Street."

    with pytest.raises(ValueError, match="street_address"):
        builder.audit_private_output(address)
    assert "street_address" in corpus.flag_text(address)


def configure_output_fixture(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> dict:
    paths = {
        "OUTPUT_MD": tmp_path / "master.md",
        "OUTPUT_DOCX": tmp_path / "master.docx",
        "OUTPUT_PDF": tmp_path / "master.pdf",
        "SOURCE_LEDGER_JSON": tmp_path / "ledger.json",
        "SOURCE_LEDGER_MD": tmp_path / "ledger.md",
        "CROSSWALK_JSON": tmp_path / "crosswalk.json",
        "CROSSWALK_MD": tmp_path / "crosswalk.md",
        "ARCHITECTURE_FIGURE": tmp_path / "architecture.png",
        "MATURITY_FIGURE": tmp_path / "maturity.png",
        "OUTPUT_MANIFEST": tmp_path / "manifest.json",
    }
    for name, path in paths.items():
        monkeypatch.setattr(builder, name, path)
    generated = tuple(
        paths[name]
        for name in (
            "OUTPUT_MD",
            "OUTPUT_DOCX",
            "OUTPUT_PDF",
            "SOURCE_LEDGER_JSON",
            "SOURCE_LEDGER_MD",
            "CROSSWALK_JSON",
            "CROSSWALK_MD",
            "ARCHITECTURE_FIGURE",
            "MATURITY_FIGURE",
        )
    )
    monkeypatch.setattr(builder, "GENERATED_ARTIFACTS", generated)

    crosswalk = {
        "schema": "lumencore_eia_hourly_supersession_crosswalk.v1",
        "status": builder.STATUS,
        "generated_utc": "2026-08-02T13:00:00+00:00",
        "publication_authorized": False,
        "version_succession_is_evidence_promotion": False,
        "automatic_promotion_allowed": False,
        "versions": [
            {
                "version": "v2",
                "protocol_role": "all_authority_direct_parent",
                "disposition": "ACTIVE_PRESERVED_PARENT",
                "performance_claim_ready": False,
            },
            {
                "version": "v3",
                "protocol_role": "frozen_future_only_hybrid_confirmation",
                "disposition": "ACTIVE_FROZEN_PROSPECTIVE_CONFIRMATION",
                "performance_claim_ready": False,
            },
            {
                "version": "v4",
                "protocol_role": "disjoint_temporal_replication",
                "disposition": "DEFERRED_UNTIL_DISJOINT_TEMPORAL_REPLICATION",
                "performance_claim_ready": False,
            },
            {
                "version": "v5",
                "protocol_role": "named_independent_evaluator_replication",
                "disposition": "DEFERRED_UNTIL_NAMED_INDEPENDENT_EVALUATOR",
                "performance_claim_ready": False,
            },
        ],
        "claim_boundary": "Version succession is not evidence promotion.",
    }
    crosswalk["crosswalk_sha256"] = builder.stable_hash(crosswalk)
    source = tmp_path / "source.json"
    source.write_text("{}\n", encoding="utf-8")
    payload = {
        "generated_utc": "2026-08-02T13:00:00+00:00",
        "paper_payload_sha256": "a" * 64,
        "source_ledger": {"ledger_sha256": "b" * 64, "records": []},
        "canonical_source_receipts": [builder.file_receipt(source)],
        "supersession_crosswalk": crosswalk,
    }
    current_markdown = (
        "Paper payload SHA-256: `" + payload["paper_payload_sha256"] + "`\n"
        "1537 sealed predictions.\n"
    )
    current_ledger_markdown = "# Current source ledger\n"
    monkeypatch.setattr(builder, "render_master_markdown", lambda _: current_markdown)
    monkeypatch.setattr(
        builder,
        "render_source_ledger_markdown",
        lambda _: current_ledger_markdown,
    )

    paths["OUTPUT_MD"].write_text(current_markdown, encoding="utf-8")
    paths["SOURCE_LEDGER_JSON"].write_text(
        json.dumps(payload["source_ledger"]), encoding="utf-8"
    )
    paths["SOURCE_LEDGER_MD"].write_text(
        current_ledger_markdown, encoding="utf-8"
    )
    paths["CROSSWALK_JSON"].write_text(
        json.dumps(crosswalk), encoding="utf-8"
    )
    paths["CROSSWALK_MD"].write_text(
        builder.render_crosswalk_markdown(crosswalk), encoding="utf-8"
    )
    document = Document()
    document.add_paragraph(
        "Paper payload SHA-256: " + payload["paper_payload_sha256"]
    )
    document.save(paths["OUTPUT_DOCX"])
    paths["ARCHITECTURE_FIGURE"].write_bytes(b"architecture")
    paths["MATURITY_FIGURE"].write_bytes(b"maturity")
    paths["OUTPUT_PDF"].write_bytes(b"private-review-pdf")
    pdf_text = (
        builder.STATUS
        + "\nPaper payload SHA-256: "
        + payload["paper_payload_sha256"]
        + "\n"
        + ("searchable private review content " * 300)
    )
    monkeypatch.setattr(
        builder, "extract_pdf_text", lambda _: (pdf_text, 3, False)
    )
    return {"paths": paths, "payload": payload, "generated": generated}


def test_manifest_contract_contains_exactly_builder_generated_outputs(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    fixture = configure_output_fixture(monkeypatch, tmp_path)

    manifest = builder.write_manifest(fixture["payload"])

    assert [row["path"] for row in manifest["generated_artifacts"]] == [
        str(path) for path in fixture["generated"]
    ]
    assert any(
        Path(row["path"]).suffix.casefold() == ".pdf"
        for row in manifest["generated_artifacts"]
    )
    assert manifest["canonical_pdf"]["inspection"]["all_checks_pass"] is True
    assert manifest["supersession_crosswalk_sha256"] == fixture["payload"][
        "supersession_crosswalk"
    ]["crosswalk_sha256"]
    assert builder.verify_manifest() == manifest


def test_manifest_rejects_tampered_canonical_pdf(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    fixture = configure_output_fixture(monkeypatch, tmp_path)
    builder.write_manifest(fixture["payload"])
    fixture["paths"]["OUTPUT_PDF"].write_bytes(b"tampered-private-review-pdf")

    with pytest.raises(ValueError, match="Artifact receipt mismatch"):
        builder.verify_manifest()


def test_patent_archive_is_hash_inventoried_without_opening_zip(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    archive = tmp_path / "sealed-patent.zip"
    archive.write_bytes(b"PK\x03\x04opaque-patent-package")
    output_root = tmp_path / "private-review"
    monkeypatch.setattr(corpus, "ROOT", tmp_path)
    monkeypatch.setattr(corpus, "OUTPUT_ROOT", output_root)
    monkeypatch.setattr(corpus, "EXTRACTED_ROOT", output_root / "extracted")
    monkeypatch.setattr(corpus, "INVENTORY_PATH", output_root / "inventory.json")
    monkeypatch.setattr(
        corpus,
        "SOURCES",
        (
            corpus.SourceSpec(
                "sealed_patent",
                "repo",
                archive.name,
                "existence_and_lineage_only",
                "patent_package_do_not_expand",
                False,
            ),
        ),
    )

    class ForbiddenZipFile:
        def __init__(self, *args, **kwargs):
            raise AssertionError("Patent archive must not be opened")

    monkeypatch.setattr(corpus.zipfile, "ZipFile", ForbiddenZipFile)
    inventory = corpus.build_inventory()
    record = inventory["records"][0]

    assert record["status"] == "INVENTORIED_NOT_EXPANDED"
    assert record["kind"] == "opaque_archive"
    assert record["archive_contents_inspected"] is False
    assert record["archive_member_count"] is None
    assert "extracted_path" not in record
    corpus.verify_inventory(inventory)


def test_markdown_table_rows_are_kept_together() -> None:
    document = Document()

    builder.add_markdown_table(
        document,
        [
            ["Source", "Disposition"],
            ["long_source_name", "A multi-line disposition that must remain intact."],
        ],
    )

    table = document.tables[0]
    assert all("w:cantSplit" in row._tr.xml for row in table.rows)


def test_supersession_crosswalk_preserves_readiness_without_promotion() -> None:
    v1 = {
        "prediction_count": 10,
        "settlement_count": 8,
        "common_settled_hour_count": 4,
        "sample_gates": {
            "preliminary_ready": False,
            "confirmatory_ready": False,
            "durability_ready": False,
        },
        "promotion_evaluation_complete": False,
    }
    v2 = {
        "prediction_panel_count": 208,
        "settlement_panel_count": 204,
        "sealed_authority_prediction_count": 1664,
        "settled_authority_prediction_count": 1632,
        "common_settled_hour_count": 204,
        "sample_gates": {
            "preliminary_ready": True,
            "confirmatory_ready": False,
            "durability_ready": False,
        },
        "promotion_evaluation_complete": False,
        "performance_claim_ready": False,
    }
    v3 = {
        "v3_prediction_panel_count": 8,
        "v3_settlement_panel_count": 4,
        "v3_sealed_authority_prediction_count": 64,
        "v3_settled_authority_prediction_count": 32,
        "complete_utc_day_count": 0,
        "sample_gates": {
            "operational_shakeout_ready": False,
            "preliminary_sample_ready": False,
            "confirmatory_sample_ready": False,
        },
        "performance": {
            "scores_suppressed": True,
            "promotion_evaluation_complete": False,
        },
        "protocol_commit": None,
    }

    crosswalk = builder.build_supersession_crosswalk(
        "2026-08-02T13:00:00+00:00", v1, v2, v3
    )
    by_version = {row["version"]: row for row in crosswalk["versions"]}

    assert by_version["v2"]["preliminary_sample_ready_only"] is True
    assert by_version["v2"]["performance_claim_ready"] is False
    assert by_version["v3"]["protocol_commit_bound"] is False
    assert by_version["v3"]["scores_suppressed"] is True
    assert by_version["v4"]["disposition"].startswith("DEFERRED")
    assert by_version["v5"]["disposition"].startswith("DEFERRED")
    assert crosswalk["version_succession_is_evidence_promotion"] is False


def test_seal_rejects_stale_quantitative_markdown(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    fixture = configure_output_fixture(monkeypatch, tmp_path)
    fixture["paths"]["OUTPUT_MD"].write_text(
        "Paper payload SHA-256: `" + "c" * 64 + "`\n1483 sealed predictions.\n",
        encoding="utf-8",
    )

    with pytest.raises(ValueError, match="stale relative to the current source payload"):
        builder.write_manifest(fixture["payload"])

    assert not fixture["paths"]["OUTPUT_MANIFEST"].exists()


def test_seal_rejects_stale_docx_payload_binding(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    fixture = configure_output_fixture(monkeypatch, tmp_path)
    document = Document()
    document.add_paragraph("Paper payload SHA-256: " + "c" * 64)
    document.save(fixture["paths"]["OUTPUT_DOCX"])

    with pytest.raises(ValueError, match="DOCX is stale"):
        builder.write_manifest(fixture["payload"])

    assert not fixture["paths"]["OUTPUT_MANIFEST"].exists()


def test_markdown_layout_markers_are_materialized() -> None:
    source = (
        "{PAGE_BREAK}\n{FIGURE:maturity}\n{FIGURE:architecture}\n"
        "{FIGURE:concept_overlay}\n"
    )

    rendered = builder.materialize_markdown_layout(source)

    builder.audit_layout_markers(rendered)
    assert "<!-- PAGE_BREAK -->" in rendered
    assert "figures/evidence_maturity_ladder.png" in rendered
    assert "figures/evidence_governed_architecture.png" in rendered
    assert "<!-- FIGURE:concept_overlay -->" in rendered


def test_docx_builder_consumes_materialized_layout(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    output_docx = tmp_path / "master.docx"
    monkeypatch.setattr(builder, "OUTPUT_DOCX", output_docx)

    def fake_picture(document, path, width, alt_text):
        document.add_paragraph("FIGURE:" + Path(path).name)

    monkeypatch.setattr(builder, "add_picture_with_alt", fake_picture)
    payload = {
        "repository_maturity_level": 3,
        "paper_payload_sha256": "a" * 64,
    }
    markdown = """<!-- BODY -->
<!-- PAGE_BREAK -->
![Evidence maturity ladder with the current repository at Level 3](figures/evidence_maturity_ladder.png)

*Figure 1. Claim-specific evidence maturity. Levels 4 and 5 remain closed.*

![Evidence-governed architecture from authorized sources through independent review](figures/evidence_governed_architecture.png)

*Figure 2. Evidence-governed architecture. Nature-inspired forms enter as hypotheses.*

<!-- FIGURE:concept_overlay -->

## References

- [R1] A compact reference entry.
"""

    builder.build_docx(payload, markdown)

    text = builder.extract_docx_text(output_docx)
    builder.audit_layout_markers(text)
    assert "FIGURE:evidence_maturity_ladder.png" in text
    assert "FIGURE:evidence_governed_architecture.png" in text
    assert "FIGURE:metatron_golden_ratio_overlay.png" in text
    assert "Figure 1. Claim-specific evidence maturity." in text
    assert "Figure 2. Evidence-governed architecture." in text
    document = Document(output_docx)
    reference = next(p for p in document.paragraphs if "[R1]" in p.text)
    assert reference.paragraph_format.space_after.pt == 2
    assert reference.paragraph_format.line_spacing == pytest.approx(1.08, abs=0.001)
    assert all(run.font.size.pt == 9.5 for run in reference.runs if run.text.strip())


def test_seal_rejects_raw_layout_marker_in_docx(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> None:
    fixture = configure_output_fixture(monkeypatch, tmp_path)
    document = Document(fixture["paths"]["OUTPUT_DOCX"])
    document.add_paragraph("{FIGURE:maturity}")
    document.save(fixture["paths"]["OUTPUT_DOCX"])

    with pytest.raises(ValueError, match="raw layout marker remains"):
        builder.write_manifest(fixture["payload"])

    assert not fixture["paths"]["OUTPUT_MANIFEST"].exists()
