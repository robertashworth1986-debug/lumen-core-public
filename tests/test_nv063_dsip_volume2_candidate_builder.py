from __future__ import annotations

import importlib.util
import subprocess
import sys
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
PACKAGE = ROOT / "grant_submissions" / "NV063_HarborSentinel"
BUILDER_PATH = PACKAGE / "build_nv063_dsip_volume2_candidate.py"
# Format-valid but deliberately synthetic; every release build in these tests stays
# under pytest's temporary directory.
SYNTHETIC_PROPOSAL_NUMBER = "N26BZ-NV063-0000"


def load_builder():
    spec = importlib.util.spec_from_file_location("nv063_volume2_builder", BUILDER_PATH)
    assert spec is not None and spec.loader is not None
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


BUILDER = load_builder()


def run_builder(*arguments: object) -> subprocess.CompletedProcess[str]:
    return subprocess.run(
        [sys.executable, str(BUILDER_PATH), *(str(argument) for argument in arguments)],
        cwd=ROOT,
        capture_output=True,
        text=True,
        check=False,
    )


def result_text(result: subprocess.CompletedProcess[str]) -> str:
    return f"{result.stdout}\n{result.stderr}"


def test_release_requires_explicit_proposal_number(tmp_path: Path) -> None:
    output = tmp_path / "missing-number.docx"

    result = run_builder("--release", "--output", output)

    assert result.returncode == 2
    assert "RELEASE_REQUIRES_EXPLICIT_DSIP_PROPOSAL_NUMBER" in result_text(result)
    assert not output.exists()


@pytest.mark.parametrize(
    "proposal_number",
    [
        "N26BZ-NV063-ASSIGNED",
        "N26BZ-NV999-1234",
        "n26bz-NV063-1234",
        "N26BZ-NV063-12 34",
    ],
)
def test_release_rejects_invalid_proposal_number(
    tmp_path: Path, proposal_number: str
) -> None:
    output = tmp_path / "invalid-number.docx"

    result = run_builder(
        "--release",
        "--proposal-number",
        proposal_number,
        "--output",
        output,
    )

    assert result.returncode == 2
    assert "INVALID_DSIP_PROPOSAL_NUMBER_FORMAT" in result_text(result)
    assert not output.exists()


def test_release_build_in_temp_has_exact_number_and_no_review_controls(
    tmp_path: Path,
) -> None:
    output = tmp_path / "harborsentinel-volume2-release.docx"

    result = run_builder(
        "--release",
        "--proposal-number",
        SYNTHETIC_PROPOSAL_NUMBER,
        "--output",
        output,
    )

    assert result.returncode == 0, result_text(result)
    assert output.is_file()

    doc = Document(output)
    assert len(doc.sections) == 1
    for section in doc.sections:
        header_text = BUILDER.part_text(section.header)
        assert header_text.endswith(f"Proposal No. {SYNTHETIC_PROPOSAL_NUMBER}")
        assert header_text.count(SYNTHETIC_PROPOSAL_NUMBER) == 1

    body_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "No letter of support is included in this Volume 2." in body_text
    for boundary in BUILDER.REQUIRED_CLAIM_BOUNDARIES:
        assert boundary in body_text

    package_text, custom_parts = BUILDER.package_story_text(output)
    assert custom_parts == []
    assert SYNTHETIC_PROPOSAL_NUMBER in package_text
    assert BUILDER.release_text_violations(package_text) == []

    with ZipFile(output) as package:
        assert package.testzip() is None


def test_release_refuses_existing_output_without_explicit_overwrite(
    tmp_path: Path,
) -> None:
    output = tmp_path / "existing.docx"
    original = b"do not replace"
    output.write_bytes(original)

    refused = run_builder(
        "--release",
        "--proposal-number",
        SYNTHETIC_PROPOSAL_NUMBER,
        "--output",
        output,
    )

    assert refused.returncode == 2
    assert "RELEASE_OUTPUT_EXISTS" in result_text(refused)
    assert output.read_bytes() == original

    allowed = run_builder(
        "--release",
        "--proposal-number",
        SYNTHETIC_PROPOSAL_NUMBER,
        "--output",
        output,
        "--overwrite",
    )
    assert allowed.returncode == 0, result_text(allowed)
    with ZipFile(output) as package:
        assert package.testzip() is None


def test_review_candidate_defaults_remain_backward_compatible(tmp_path: Path) -> None:
    output = tmp_path / "review-candidate.docx"

    result = run_builder("--output", output)

    assert result.returncode == 0, result_text(result)
    doc = Document(output)
    header_text = BUILDER.part_text(doc.sections[0].header)
    footer_text = BUILDER.part_text(doc.sections[0].footer)
    body_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert BUILDER.REVIEW_PROPOSAL_LABEL in header_text
    assert "REVIEW CANDIDATE - NOT CERTIFIED" in footer_text
    assert "REVIEW CANDIDATE - COMPLETE LIVE DSIP" in body_text
