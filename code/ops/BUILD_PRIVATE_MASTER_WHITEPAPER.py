"""Build the private LumenCore master whitepaper and its source ledger.

This builder is deliberately fail-closed. It reports current local evidence,
preserves negative results, treats legacy concept papers as hypotheses, omits
private identifiers and patent detail, and has no publication capability.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import subprocess
import tempfile
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_ROOT = ROOT / "out" / "private_review" / "master_whitepaper"
INVENTORY_PATH = OUTPUT_ROOT / "corpus_inventory.json"
SOURCE_LEDGER_JSON = OUTPUT_ROOT / "MASTER_WHITEPAPER_SOURCE_LEDGER_2026-08-02.json"
SOURCE_LEDGER_MD = OUTPUT_ROOT / "MASTER_WHITEPAPER_SOURCE_LEDGER_2026-08-02.md"
OUTPUT_MD = OUTPUT_ROOT / "LUMENCORE_MASTER_WHITEPAPER_PRIVATE_REVIEW_2026-08-02.md"
OUTPUT_DOCX = OUTPUT_ROOT / "LumenCore_Master_Whitepaper_PRIVATE_REVIEW_2026-08-02.docx"
OUTPUT_PDF = OUTPUT_ROOT / "LumenCore_Master_Whitepaper_PRIVATE_REVIEW_2026-08-02.pdf"
OUTPUT_MANIFEST = OUTPUT_ROOT / "MASTER_WHITEPAPER_MANIFEST_2026-08-02.json"
CROSSWALK_JSON = OUTPUT_ROOT / "MASTER_WHITEPAPER_SUPERSESSION_CROSSWALK_2026-08-02.json"
CROSSWALK_MD = OUTPUT_ROOT / "MASTER_WHITEPAPER_SUPERSESSION_CROSSWALK_2026-08-02.md"
FIGURE_ROOT = OUTPUT_ROOT / "figures"
ARCHITECTURE_FIGURE = FIGURE_ROOT / "evidence_governed_architecture.png"
MATURITY_FIGURE = FIGURE_ROOT / "evidence_maturity_ladder.png"
GENERATED_ARTIFACTS = (
    OUTPUT_MD,
    OUTPUT_DOCX,
    OUTPUT_PDF,
    SOURCE_LEDGER_JSON,
    SOURCE_LEDGER_MD,
    CROSSWALK_JSON,
    CROSSWALK_MD,
    ARCHITECTURE_FIGURE,
    MATURITY_FIGURE,
)

README_PATH = ROOT / "README.md"
CURRENT_WHITEPAPER = ROOT / "docs" / "LUMENCORE_SOURCE_NATIVE_BENCHMARK_WHITEPAPER_CURRENT.md"
SOURCE_NATIVE_LEDGER = ROOT / "out" / "ops" / "source_native_family_baseline_ledger_latest.json"
SOURCE_NATIVE_V3_STATUS = (
    ROOT
    / "out"
    / "time_series_source_native_prospective_v3"
    / "prospective_status_latest.json"
)
EIA_V1_STATUS = (
    ROOT / "out" / "eia_grid_prospective_hourly_router" / "prospective_status_latest.json"
)
EIA_V1_PROTOCOL = ROOT / "config" / "eia_grid_prospective_hourly_router_protocol_v1.json"
EIA_V1_RUNTIME = ROOT / "code" / "eia_grid_prospective_hourly_router.py"
EIA_V2_STATUS = (
    ROOT
    / "out"
    / "eia_grid_all_authority_direct_hourly_router"
    / "prospective_status_latest.json"
)
EIA_V2_PROTOCOL = ROOT / "config" / "eia_grid_all_authority_direct_hourly_protocol_v2.json"
EIA_V2_RUNTIME = ROOT / "code" / "eia_grid_all_authority_direct_hourly_router.py"
EIA_V3_STATUS = (
    ROOT
    / "out"
    / "eia_grid_hourly_hybrid_confirmation_v3"
    / "prospective_status_latest.json"
)
EIA_V3_PROTOCOL = ROOT / "config" / "eia_grid_hourly_hybrid_confirmation_protocol_v3.json"
EIA_V3_RUNTIME = ROOT / "code" / "eia_grid_hourly_hybrid_confirmation.py"
EIA_V3_WRAPPER = ROOT / "tools" / "Run-EiaHourlyHybridConfirmationCycle.ps1"
EIA_V3_TEST = ROOT / "tests" / "test_eia_grid_hourly_hybrid_confirmation.py"
EIA_V2_TO_V5_LADDER = ROOT / "docs" / "EIA_GRID_HOURLY_HYBRID_V2_TO_V5_LADDER_2026-08-02.md"
REVIEWER_START = ROOT / "docs" / "REVIEWER_START_HERE.md"
NOAHS_ARCHITECTURE = ROOT / "docs" / "NOAHS_REVIEWER_ARCHITECTURE_2026-07-25.md"
PUBLIC_EVIDENCE_LEDGER = (
    ROOT / "docs" / "PUBLIC_SAFE_MODEL_AND_GEOMETRY_EVIDENCE_LEDGER_2026-07-13.md"
)
PROOFLOCK_OFFER = (
    ROOT / "docs" / "LUMENCORE_EVIDENCE_PROTOCOL_REVIEW_FIXED_SCOPE_OFFER_2026-07-30.md"
)
PRODUCT_AUDIT = ROOT / "docs" / "PRODUCT_LANE_EVIDENCE_AUDIT_2026-07-29.md"
CONCEPT_OVERLAY = Path.home() / "iCloudDrive" / "metatron_golden_ratio_overlay.png"
LOGO_PATH = ROOT / "assets" / "brand" / "lumaarc_eclipse_corona_concept_v1.png"

STATUS = "PRIVATE_REVIEW_NOT_FOR_PUBLICATION"
BOUNDARY = (
    "This paper reports implemented local software, source-bound experiments, "
    "protocol states, and historical research hypotheses. It is not peer review, "
    "independent validation, field validation, a performance guarantee, a realized "
    "savings claim, a patent opinion, regulatory approval, or deployment authority."
)

NAVY = "17324D"
TEAL = "0A6B72"
GOLD = "B78A2E"
PALE_TEAL = "E8F3F3"
PALE_GOLD = "F7F1E3"
PALE_BLUE = "EAF0F5"
PALE_RED = "F7EAEA"
PALE_GRAY = "F2F4F6"
INK = "1B1F23"
MUTED = "5F6B76"
RED = "9B2C2C"
GREEN = "2E6B50"


DISPOSITIONS: dict[str, dict[str, str]] = {
    "current_source_native_benchmark": {
        "disposition": "PRIMARY_CURRENT_EVIDENCE",
        "claim_class": "implemented_measured_and_protocol",
        "use": "Body: methods, current results, limitations, and references.",
    },
    "master_context_registry": {
        "disposition": "CURRENT_INTERNAL_GOVERNANCE",
        "claim_class": "implemented_governance",
        "use": "Body: continuity, privacy, and canonical-source rules.",
    },
    "whitehole_whiteholelab_audit": {
        "disposition": "HISTORICAL_PROVENANCE_CURRENT_CLAIM_BOUNDARY",
        "claim_class": "implemented_governance",
        "use": "Body and source ledger: preserve custody history while excluding heuristic ranks and the legacy site from current evidence.",
    },
    "pitch_deck_governance": {
        "disposition": "CURRENT_PRESENTATION_RELEASE_CONTROL",
        "claim_class": "implemented_governance",
        "use": "Body and source ledger: bind the single current review deck and quarantine stale or high-risk presentation artifacts.",
    },
    "nature_inspired_climate_energy": {
        "disposition": "CONCEPT_HISTORY_REQUIRES_REDERIVATION",
        "claim_class": "conceptual",
        "use": "Body: architecture inspiration only; benefit claims and pitch placeholders excluded.",
    },
    "harmonic_backprop_onepager": {
        "disposition": "TESTABLE_ALGORITHM_HYPOTHESIS",
        "claim_class": "conceptual",
        "use": "Body: software and instrument-control research question; no physical validation claim.",
    },
    "nuclear_harmonization": {
        "disposition": "SPECULATIVE_QUARANTINE",
        "claim_class": "speculative",
        "use": "Appendix only: historical concept; requires licensed partner and conventional physics reformulation.",
    },
    "datacenter_whitepaper": {
        "disposition": "UNSUPPORTED_BENEFIT_CLAIMS_QUARANTINED",
        "claim_class": "speculative",
        "use": "Appendix only: application motivation; thermal, EMI, reliability, and cognition claims excluded.",
    },
    "lumenlogic_fresh_whitepaper": {
        "disposition": "UNSUPPORTED_BENEFIT_AND_MARKET_CLAIMS_QUARANTINED",
        "claim_class": "speculative",
        "use": "Appendix only: historical naming and application map.",
    },
    "flowform_complete_specs": {
        "disposition": "PATENT_SENSITIVE_HARDWARE_HYPOTHESIS",
        "claim_class": "conceptual",
        "use": "Body: bounded geometry variables and proposed matched tests; exact design detail omitted.",
    },
    "flowform_updated_specs": {
        "disposition": "VISUAL_ONLY_LOW_TEXT_SOURCE",
        "claim_class": "conceptual",
        "use": "Source ledger only; no evidentiary weight.",
    },
    "flowform_pdf_specs": {
        "disposition": "MISNAMED_DOCX_HARDWARE_HYPOTHESIS",
        "claim_class": "conceptual",
        "use": "Body: candidate thermal, EMI, and packaging tests; no claimed gain.",
    },
    "lumenshell_flowform_bundle": {
        "disposition": "PATENT_SENSITIVE_CONCEPT_BUNDLE",
        "claim_class": "conceptual",
        "use": "Appendix summary only; no design dimensions or patent-sensitive mechanisms reproduced.",
    },
    "lumenshell_spec_01": {
        "disposition": "PROTOTYPE_REQUIREMENTS_WITH_AUTHORSHIP_CONFLICT",
        "claim_class": "conceptual",
        "use": "Body: useful VR testbed requirements only; third-party boilerplate excluded.",
    },
    "aetherframe_genesis_logbook": {
        "disposition": "FOUNDER_CONCEPT_HISTORY",
        "claim_class": "conceptual",
        "use": "Appendix only: naming lineage; health and frequency claims excluded.",
    },
    "aetherframe_pitch_deck": {
        "disposition": "BYTE_DUPLICATE_MISLABELED_SOURCE",
        "claim_class": "duplicate",
        "use": "Source ledger only; byte-identical to FlowForm complete specifications.",
    },
    "novacore_concept_paper": {
        "disposition": "PRIVATE_DOSSIER_WITH_UNVERIFIED_CLAIMS",
        "claim_class": "speculative",
        "use": "Appendix only: subsystem naming; private identifiers and illustrative ROI excluded.",
    },
    "agentic_ai_blueprint": {
        "disposition": "ROADMAP_NOT_IMPLEMENTATION_RECEIPT",
        "claim_class": "conceptual",
        "use": "Body: orchestration roadmap context; not proof of deployed capabilities.",
    },
    "telecom_deck": {
        "disposition": "MARKETING_CONCEPT_SOURCE",
        "claim_class": "conceptual",
        "use": "Appendix only: domain hypothesis; no telecom performance claim.",
    },
    "wormhole_field_research": {
        "disposition": "SPECULATIVE_PHYSICS_EXCLUDED",
        "claim_class": "speculative",
        "use": "Historical source ledger only; no scientific or funding use.",
    },
    "wormhole_grant_proposal": {
        "disposition": "SPECULATIVE_PHYSICS_EXCLUDED",
        "claim_class": "speculative",
        "use": "Historical source ledger only; not suitable for agency submission.",
    },
    "logic_specs_placeholder": {
        "disposition": "INVALID_PLACEHOLDER",
        "claim_class": "unusable",
        "use": "Excluded; 41-byte file is not a readable DOCX.",
    },
    "lumencore_patent_bundle": {
        "disposition": "SEALED_PATENT_ARCHIVE",
        "claim_class": "patent_sensitive",
        "use": "Existence and hash only; archive not expanded.",
    },
    "lumenshell_patent_package": {
        "disposition": "SEALED_PATENT_ARCHIVE",
        "claim_class": "patent_sensitive",
        "use": "Existence and hash only; archive not expanded.",
    },
}


PRIVACY_PATTERNS = {
    "email": re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.IGNORECASE),
    "phone": re.compile(r"(?:\+?1[-. ]?)?\(?\d{3}\)?[-. ]\d{3}[-. ]\d{4}"),
    "federal_identifier": re.compile(r"\b(?:UEI|CAGE|EIN)\b", re.IGNORECASE),
    "patent_application_number": re.compile(
        r"\b(?:patent|application)\s*(?:number|#|no\.)?\s*\d{2}/\d{3}[, ]?\d{3}\b",
        re.IGNORECASE,
    ),
    "street_address": re.compile(
        r"\b\d{2,5}\s+(?:(?-i:[A-Z])[A-Z0-9.'-]{0,30}\s+){1,6}"
        r"(?:ST|STREET|AVE|AVENUE|RD|ROAD|DR|DRIVE|LN|LANE|BLVD|BOULEVARD)\b",
        re.IGNORECASE,
    ),
}


def read_json(path: Path) -> dict[str, Any]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError(f"Expected object: {path}")
    return payload


def receipt_from_bytes(path: Path, content: bytes) -> dict[str, Any]:
    try:
        display_path = path.relative_to(ROOT).as_posix()
    except ValueError:
        display_path = str(path)
    return {
        "path": display_path,
        "exists": True,
        "bytes": len(content),
        "sha256": hashlib.sha256(content).hexdigest(),
    }


def read_json_with_receipt(path: Path) -> tuple[dict[str, Any], dict[str, Any]]:
    require_file(path)
    content = path.read_bytes()
    payload = json.loads(content.decode("utf-8"))
    if not isinstance(payload, dict):
        raise ValueError(f"Expected object: {path}")
    return payload, receipt_from_bytes(path, content)


def stable_hash(payload: Any) -> str:
    encoded = json.dumps(
        payload, sort_keys=True, ensure_ascii=True, separators=(",", ":")
    ).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


def file_receipt(path: Path) -> dict[str, Any]:
    if not path.is_file():
        return {"path": str(path), "exists": False, "bytes": 0, "sha256": None}
    content = path.read_bytes()
    return receipt_from_bytes(path, content)


def require_file(path: Path) -> None:
    if not path.is_file() or path.stat().st_size == 0:
        raise ValueError(f"Required source is missing or empty: {path}")


def build_source_ledger(inventory: dict[str, Any], generated_utc: str) -> dict[str, Any]:
    records = inventory.get("records")
    if not isinstance(records, list) or len(records) != len(DISPOSITIONS):
        raise ValueError("Corpus inventory record count is unexpected")
    ledger_records: list[dict[str, Any]] = []
    seen: set[str] = set()
    for source in records:
        if not isinstance(source, dict):
            raise ValueError("Corpus inventory contains a non-object record")
        source_id = source.get("source_id")
        if not isinstance(source_id, str) or source_id not in DISPOSITIONS:
            raise ValueError(f"Unclassified corpus source: {source_id}")
        seen.add(source_id)
        resolution = DISPOSITIONS[source_id]
        resolved_path = source.get("resolved_path")
        filename = Path(resolved_path).name if isinstance(resolved_path, str) else None
        ledger_records.append(
            {
                "source_id": source_id,
                "filename": filename,
                "source_status": source.get("status"),
                "kind": source.get("kind"),
                "bytes": source.get("bytes"),
                "source_sha256": source.get("source_sha256"),
                "intended_role": source.get("intended_role"),
                "sensitivity": source.get("sensitivity"),
                "content_flags": source.get("content_flags", []),
                **resolution,
            }
        )
    if seen != set(DISPOSITIONS):
        raise ValueError("Source disposition coverage is incomplete")

    by_class: dict[str, int] = {}
    for record in ledger_records:
        claim_class = str(record["claim_class"])
        by_class[claim_class] = by_class.get(claim_class, 0) + 1
    payload: dict[str, Any] = {
        "schema": "lumencore_master_whitepaper_source_ledger_v1",
        "status": STATUS,
        "generated_utc": generated_utc,
        "publication_authorized": False,
        "patent_archives_expanded": False,
        "source_count": len(ledger_records),
        "claim_class_counts": by_class,
        "records": ledger_records,
    }
    payload["ledger_sha256"] = stable_hash(payload)
    return payload


def render_source_ledger_markdown(ledger: dict[str, Any]) -> str:
    lines = [
        "# LumenCore Master Whitepaper Source Ledger",
        "",
        f"- Status: `{STATUS}`",
        f"- Sources: `{ledger['source_count']}`",
        "- Publication authorized: `false`",
        "- Patent archives expanded: `false`",
        f"- Ledger SHA-256: `{ledger['ledger_sha256']}`",
        "",
        "This ledger records source identity and disposition. It is not a patent, "
        "scientific, legal, authorship, or publication opinion.",
        "",
        "| Source | Class | Disposition | Use |",
        "|---|---|---|---|",
    ]
    for record in ledger["records"]:
        lines.append(
            "| `{source_id}` | `{claim_class}` | `{disposition}` | {use} |".format(
                **record
            )
        )
    return "\n".join(lines) + "\n"


def build_supersession_crosswalk(
    generated_utc: str,
    eia_v1: dict[str, Any],
    eia_v2: dict[str, Any],
    eia_v3: dict[str, Any],
) -> dict[str, Any]:
    crosswalk: dict[str, Any] = {
        "schema": "lumencore_eia_hourly_supersession_crosswalk.v1",
        "status": STATUS,
        "generated_utc": generated_utc,
        "publication_authorized": False,
        "version_succession_is_evidence_promotion": False,
        "automatic_promotion_allowed": False,
        "versions": [
            {
                "version": "v1",
                "protocol_role": "legacy_specialist_router_collection",
                "disposition": "PRESERVED_NONCONFIRMATORY_OPERATIONAL_CHAIN",
                "prediction_count": eia_v1.get("prediction_count"),
                "settlement_count": eia_v1.get("settlement_count"),
                "common_settled_hour_count": eia_v1.get(
                    "common_settled_hour_count"
                ),
                "sample_gates": eia_v1.get("sample_gates"),
                "promotion_evaluation_complete": eia_v1.get(
                    "promotion_evaluation_complete"
                ),
                "performance_claim_ready": False,
            },
            {
                "version": "v2",
                "protocol_role": "all_authority_direct_parent",
                "disposition": "ACTIVE_PRESERVED_PARENT",
                "prediction_panel_count": eia_v2.get("prediction_panel_count"),
                "settlement_panel_count": eia_v2.get("settlement_panel_count"),
                "sealed_authority_prediction_count": eia_v2.get(
                    "sealed_authority_prediction_count"
                ),
                "settled_authority_prediction_count": eia_v2.get(
                    "settled_authority_prediction_count"
                ),
                "common_settled_hour_count": eia_v2.get(
                    "common_settled_hour_count"
                ),
                "sample_gates": eia_v2.get("sample_gates"),
                "preliminary_sample_ready_only": bool(
                    (eia_v2.get("sample_gates") or {}).get("preliminary_ready")
                ),
                "promotion_evaluation_complete": eia_v2.get(
                    "promotion_evaluation_complete"
                ),
                "evidence_packet_release_ready": eia_v2.get(
                    "external_release_ready"
                ),
                "performance_claim_ready": eia_v2.get("performance_claim_ready"),
            },
            {
                "version": "v3",
                "protocol_role": "frozen_future_only_hybrid_confirmation",
                "disposition": "ACTIVE_FROZEN_PROSPECTIVE_CONFIRMATION",
                "prediction_panel_count": eia_v3.get("v3_prediction_panel_count"),
                "settlement_panel_count": eia_v3.get("v3_settlement_panel_count"),
                "sealed_authority_prediction_count": eia_v3.get(
                    "v3_sealed_authority_prediction_count"
                ),
                "settled_authority_prediction_count": eia_v3.get(
                    "v3_settled_authority_prediction_count"
                ),
                "complete_utc_day_count": eia_v3.get("complete_utc_day_count"),
                "sample_gates": eia_v3.get("sample_gates"),
                "scores_suppressed": (eia_v3.get("performance") or {}).get(
                    "scores_suppressed"
                ),
                "promotion_evaluation_complete": (
                    eia_v3.get("performance") or {}
                ).get("promotion_evaluation_complete"),
                "protocol_commit": eia_v3.get("protocol_commit"),
                "protocol_commit_bound": bool(eia_v3.get("protocol_commit")),
                "performance_claim_ready": False,
            },
            {
                "version": "v4",
                "protocol_role": "disjoint_temporal_replication",
                "disposition": "DEFERRED_UNTIL_DISJOINT_TEMPORAL_REPLICATION",
                "performance_claim_ready": False,
            },
            {
                "version": "v5",
                "protocol_role": "named_independent_evaluator_replication",
                "disposition": "DEFERRED_UNTIL_NAMED_INDEPENDENT_EVALUATOR",
                "performance_claim_ready": False,
            },
        ],
        "claim_boundary": (
            "V2 preliminary sample readiness is a sample-count fact only. V3 is a "
            "separate future-only chain with suppressed scores and no protocol-commit "
            "receipt. V4 and V5 do not exist as active evidence lanes."
        ),
    }
    crosswalk["crosswalk_sha256"] = stable_hash(crosswalk)
    return crosswalk


def render_crosswalk_markdown(crosswalk: dict[str, Any]) -> str:
    rows = []
    for version in crosswalk["versions"]:
        rows.append(
            "| `{version}` | `{protocol_role}` | `{disposition}` | `{ready}` |".format(
                version=version["version"],
                protocol_role=version["protocol_role"],
                disposition=version["disposition"],
                ready=str(version["performance_claim_ready"]).lower(),
            )
        )
    return "\n".join(
        [
            "# EIA Hourly Supersession Crosswalk",
            "",
            f"- Status: `{STATUS}`",
            "- Publication authorized: `false`",
            "- Version succession is evidence promotion: `false`",
            f"- Crosswalk SHA-256: `{crosswalk['crosswalk_sha256']}`",
            "",
            "| Version | Protocol role | Disposition | Performance claim ready |",
            "|---|---|---|---|",
            *rows,
            "",
            crosswalk["claim_boundary"],
            "",
        ]
    )


def build_payload(at: datetime | None = None) -> dict[str, Any]:
    at = at or datetime.now(timezone.utc)
    for path in (
        README_PATH,
        CURRENT_WHITEPAPER,
        SOURCE_NATIVE_LEDGER,
        SOURCE_NATIVE_V3_STATUS,
        EIA_V1_STATUS,
        EIA_V1_PROTOCOL,
        EIA_V1_RUNTIME,
        EIA_V2_STATUS,
        EIA_V2_PROTOCOL,
        EIA_V2_RUNTIME,
        EIA_V3_STATUS,
        EIA_V3_PROTOCOL,
        EIA_V3_RUNTIME,
        EIA_V3_WRAPPER,
        EIA_V3_TEST,
        EIA_V2_TO_V5_LADDER,
        REVIEWER_START,
        NOAHS_ARCHITECTURE,
        PUBLIC_EVIDENCE_LEDGER,
        PROOFLOCK_OFFER,
        PRODUCT_AUDIT,
        INVENTORY_PATH,
        LOGO_PATH,
        CONCEPT_OVERLAY,
    ):
        require_file(path)

    readme = README_PATH.read_text(encoding="utf-8")
    maturity_match = re.search(
        r"Current repository-wide supported maturity: Level\s+(\d+)", readme
    )
    if not maturity_match:
        raise ValueError("Repository maturity statement is missing")
    maturity = int(maturity_match.group(1))
    if maturity != 3:
        raise ValueError("Master paper requires explicit review for maturity changes")

    inventory, inventory_receipt = read_json_with_receipt(INVENTORY_PATH)
    if inventory.get("status") != STATUS:
        raise ValueError("Corpus inventory is not private-review gated")
    if inventory.get("publication_authorized") is not False:
        raise ValueError("Corpus publication gate must remain false")

    generated_utc = at.isoformat()
    source_ledger = build_source_ledger(inventory, generated_utc)
    source_native, source_native_receipt = read_json_with_receipt(
        SOURCE_NATIVE_LEDGER
    )
    source_summary = source_native.get("summary")
    if not isinstance(source_summary, dict):
        raise ValueError("Source-native summary is missing")
    if any(
        source_summary.get(field) is not False
        for field in (
            "field_validation_claim_allowed",
            "live_trading_or_autonomous_execution_allowed",
            "public_performance_claim_allowed",
            "real_dollar_savings_claim_allowed",
        )
    ):
        raise ValueError("Source-native claim gates unexpectedly opened")

    eia_v1, eia_v1_status_receipt = read_json_with_receipt(EIA_V1_STATUS)
    eia_v1_protocol, eia_v1_protocol_receipt = read_json_with_receipt(
        EIA_V1_PROTOCOL
    )
    eia_v2, eia_v2_status_receipt = read_json_with_receipt(EIA_V2_STATUS)
    eia_v2_protocol, eia_v2_protocol_receipt = read_json_with_receipt(
        EIA_V2_PROTOCOL
    )
    eia_v3, eia_v3_status_receipt = read_json_with_receipt(EIA_V3_STATUS)
    eia_v3_protocol, eia_v3_protocol_receipt = read_json_with_receipt(
        EIA_V3_PROTOCOL
    )

    if eia_v1.get("promotion_evaluation_complete") is not False:
        raise ValueError("EIA promotion state changed and requires human review")
    v1_sample_gates = eia_v1.get("sample_gates")
    if not isinstance(v1_sample_gates, dict) or any(
        v1_sample_gates.get(name) is not False
        for name in ("preliminary_ready", "confirmatory_ready", "durability_ready")
    ):
        raise ValueError("EIA V1 sample gates unexpectedly opened")

    v2_sample_gates = eia_v2.get("sample_gates")
    if not isinstance(v2_sample_gates, dict) or (
        v2_sample_gates.get("preliminary_ready") is not True
        or v2_sample_gates.get("confirmatory_ready") is not False
        or v2_sample_gates.get("durability_ready") is not False
    ):
        raise ValueError("EIA V2 sample-gate contract changed")
    if any(
        eia_v2.get(name) is not False
        for name in ("performance_claim_ready", "promotion_evaluation_complete")
    ):
        raise ValueError("EIA V2 performance gate unexpectedly opened")
    if not isinstance(eia_v2.get("external_release_ready"), bool):
        raise ValueError("EIA V2 evidence-packet release gate is not explicit")

    v3_sample_gates = eia_v3.get("sample_gates")
    v3_performance = eia_v3.get("performance")
    v3_succession = eia_v3.get("succession")
    if not isinstance(v3_sample_gates, dict) or any(
        v3_sample_gates.get(name) is not False
        for name in (
            "operational_shakeout_ready",
            "preliminary_sample_ready",
            "confirmatory_sample_ready",
        )
    ):
        raise ValueError("EIA V3 sample gates unexpectedly opened")
    if not isinstance(v3_performance, dict) or (
        v3_performance.get("scores_suppressed") is not True
        or v3_performance.get("promotion_evaluation_complete") is not False
        or v3_performance.get("automatic_promotion_allowed") is not False
    ):
        raise ValueError("EIA V3 performance boundary changed")
    expected_succession = {
        "v2": "ACTIVE_PRESERVED_PARENT",
        "v3": "ACTIVE_FROZEN_PROSPECTIVE_CONFIRMATION",
        "v4": "DEFERRED_UNTIL_DISJOINT_TEMPORAL_REPLICATION",
        "v5": "DEFERRED_UNTIL_NAMED_INDEPENDENT_EVALUATOR",
    }
    if v3_succession != expected_succession:
        raise ValueError("EIA V3 succession contract changed")

    for label, status, protocol_receipt in (
        ("V1", eia_v1, eia_v1_protocol_receipt),
        ("V2", eia_v2, eia_v2_protocol_receipt),
        ("V3", eia_v3, eia_v3_protocol_receipt),
    ):
        if status.get("protocol_sha256") != protocol_receipt["sha256"]:
            raise ValueError(f"EIA {label} protocol receipt mismatch")
    if not all(isinstance(item, dict) for item in (eia_v1_protocol, eia_v2_protocol, eia_v3_protocol)):
        raise ValueError("EIA protocol payload is not an object")
    eia_v3_runtime_receipt = file_receipt(EIA_V3_RUNTIME)
    if eia_v3.get("runtime_sha256") != eia_v3_runtime_receipt["sha256"]:
        raise ValueError("EIA V3 runtime receipt mismatch")

    prospective, prospective_receipt = read_json_with_receipt(
        SOURCE_NATIVE_V3_STATUS
    )
    if prospective.get("performance_claim_allowed") is not False:
        raise ValueError("Prospective performance gate unexpectedly opened")

    supersession_crosswalk = build_supersession_crosswalk(
        generated_utc, eia_v1, eia_v2, eia_v3
    )

    payload: dict[str, Any] = {
        "schema": "lumencore_private_master_whitepaper_v2",
        "status": STATUS,
        "generated_utc": generated_utc,
        "responsible_author": "Robert Ashworth",
        "affiliation": "LumenCore",
        "publication_authorized": False,
        "peer_reviewed": False,
        "independently_validated": False,
        "field_validated": False,
        "boundary": BOUNDARY,
        "repository_maturity_level": maturity,
        "source_native": {
            "registered_family_count": source_summary.get("registered_family_count"),
            "implementation_present_count": source_summary.get(
                "implementation_present_count"
            ),
            "implementation_required_count": source_summary.get(
                "implementation_required_count"
            ),
            "candidate_source_card_count": source_summary.get(
                "direct_candidate_source_card_count"
            ),
            "comparison_count": source_summary.get(
                "executed_direct_source_baseline_comparison_count"
            ),
            "global_holm_positive_count": source_summary.get(
                "individual_comparison_global_holm_positive_count"
            ),
            "promotion_gate_pass_count": source_summary.get(
                "internal_source_native_promotion_gate_pass_count"
            ),
            "market_signal_comparison_count": source_summary.get(
                "market_signal_comparison_count"
            ),
            "market_signal_inference_insufficient_count": source_summary.get(
                "market_signal_inference_insufficient_count"
            ),
        },
        "eia_hourly": {
            "state": eia_v1.get("state"),
            "generated_utc": eia_v1.get("generated_utc"),
            "prediction_count": eia_v1.get("prediction_count"),
            "settlement_count": eia_v1.get("settlement_count"),
            "common_settled_hour_count": eia_v1.get("common_settled_hour_count"),
            "first_common_settled_period": eia_v1.get("first_common_settled_period"),
            "latest_common_settled_period": eia_v1.get("latest_common_settled_period"),
            "protocol_commit": eia_v1.get("protocol_commit"),
            "sample_gates": v1_sample_gates,
        },
        "eia_authority_v2": {
            "state": eia_v2.get("state"),
            "prediction_panel_count": eia_v2.get("prediction_panel_count"),
            "settlement_panel_count": eia_v2.get("settlement_panel_count"),
            "sealed_authority_prediction_count": eia_v2.get(
                "sealed_authority_prediction_count"
            ),
            "settled_authority_prediction_count": eia_v2.get(
                "settled_authority_prediction_count"
            ),
            "common_settled_hour_count": eia_v2.get("common_settled_hour_count"),
            "protocol_commit": eia_v2.get("protocol_commit"),
            "sample_gates": v2_sample_gates,
            "evidence_packet_release_ready": eia_v2.get(
                "external_release_ready"
            ),
            "promotion_evaluation_complete": False,
            "performance_claim_ready": False,
        },
        "eia_hybrid_v3": {
            "state": eia_v3.get("state"),
            "prediction_panel_count": eia_v3.get("v3_prediction_panel_count"),
            "settlement_panel_count": eia_v3.get("v3_settlement_panel_count"),
            "sealed_authority_prediction_count": eia_v3.get(
                "v3_sealed_authority_prediction_count"
            ),
            "settled_authority_prediction_count": eia_v3.get(
                "v3_settled_authority_prediction_count"
            ),
            "complete_utc_day_count": eia_v3.get("complete_utc_day_count"),
            "protocol_commit": eia_v3.get("protocol_commit"),
            "protocol_commit_bound": bool(eia_v3.get("protocol_commit")),
            "sample_gates": v3_sample_gates,
            "scores_suppressed": True,
            "promotion_evaluation_complete": False,
            "automatic_promotion_allowed": False,
            "succession": v3_succession,
        },
        "source_native_v3": {
            "state": prospective.get("state"),
            "protocol_id": prospective.get("protocol_id"),
            "prediction_count": prospective.get("prediction_count"),
            "settlement_count": prospective.get("settlement_count"),
            "eligible_future_observation_count": prospective.get(
                "eligible_future_observation_count"
            ),
            "promotion_decision": prospective.get("promotion_decision"),
            "external_anchor_count": prospective.get("external_anchor_count"),
        },
        "corpus": {
            "source_count": inventory.get("source_count"),
            "extractable_count": inventory.get("extractable_count"),
            "inventory_sha256": inventory.get("inventory_sha256"),
            "source_ledger_sha256": source_ledger.get("ledger_sha256"),
        },
        "canonical_source_receipts": [
            file_receipt(path)
            for path in (
                README_PATH,
                CURRENT_WHITEPAPER,
                REVIEWER_START,
                NOAHS_ARCHITECTURE,
                PUBLIC_EVIDENCE_LEDGER,
                PROOFLOCK_OFFER,
                PRODUCT_AUDIT,
                EIA_V1_RUNTIME,
                EIA_V2_RUNTIME,
                EIA_V3_WRAPPER,
                EIA_V3_TEST,
                EIA_V2_TO_V5_LADDER,
            )
        ]
        + [
            source_native_receipt,
            prospective_receipt,
            eia_v1_status_receipt,
            eia_v1_protocol_receipt,
            eia_v2_status_receipt,
            eia_v2_protocol_receipt,
            eia_v3_status_receipt,
            eia_v3_protocol_receipt,
            eia_v3_runtime_receipt,
            inventory_receipt,
        ],
        "supersession_crosswalk": supersession_crosswalk,
        "source_ledger": source_ledger,
    }
    unsigned = dict(payload)
    unsigned.pop("source_ledger", None)
    payload["paper_payload_sha256"] = stable_hash(unsigned)
    return payload


def render_architecture_figure() -> None:
    FIGURE_ROOT.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(12, 7.2), dpi=180)
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 7.2)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    stages = [
        ("1", "Authorized\nsources", "Timestamped public or\nbuyer-approved inputs"),
        ("2", "Source-task\ncontract", "Series, cadence, horizon,\neligible population"),
        ("3", "Candidate +\nbaseline registry", "Named methods; complete\ncomparator roster"),
        ("4", "Frozen evaluation\nprotocol", "Cutoff, metrics, effect floor,\nsample gates"),
        ("5", "ProofLock\ncustody", "Hashes, manifests, ledgers,\nnegative results"),
        ("6", "Reviewer gate +\nHumanUnlock", "Independent receipt before\nclaim or action"),
    ]
    x_positions = [0.6, 4.15, 7.7]
    y_positions = [4.25, 1.55]
    for index, (number, title, detail) in enumerate(stages):
        row = 0 if index < 3 else 1
        col = index % 3
        x = x_positions[col]
        y = y_positions[row]
        face = "#E8F3F3" if index < 3 else "#EAF0F5"
        box = FancyBboxPatch(
            (x, y),
            3.0,
            1.25,
            boxstyle="round,pad=0.04,rounding_size=0.06",
            linewidth=1.4,
            edgecolor="#0A6B72",
            facecolor=face,
        )
        ax.add_patch(box)
        ax.text(
            x + 0.17,
            y + 0.9,
            number,
            fontsize=17,
            weight="bold",
            color="#B78A2E",
            va="center",
        )
        ax.text(
            x + 1.78,
            y + 0.91,
            title,
            fontsize=10.1,
            weight="bold",
            color="#17324D",
            ha="center",
            va="center",
            linespacing=0.95,
        )
        ax.text(
            x + 1.5,
            y + 0.36,
            detail,
            fontsize=8.2,
            color="#33424F",
            ha="center",
            va="center",
            linespacing=1.05,
        )

    arrow_pairs = [
        ((3.6, 4.88), (4.1, 4.88)),
        ((7.15, 4.88), (7.65, 4.88)),
        ((9.2, 4.18), (9.2, 2.85)),
        ((7.7, 2.18), (7.2, 2.18)),
        ((4.15, 2.18), (3.65, 2.18)),
    ]
    for start, end in arrow_pairs:
        ax.annotate(
            "",
            xy=end,
            xytext=start,
            arrowprops=dict(arrowstyle="->", lw=1.8, color="#5F6B76"),
        )

    hypothesis = FancyBboxPatch(
        (0.8, 6.0),
        10.2,
        0.72,
        boxstyle="round,pad=0.04,rounding_size=0.05",
        linewidth=1.2,
        linestyle="--",
        edgecolor="#B78A2E",
        facecolor="#F7F1E3",
    )
    ax.add_patch(hypothesis)
    ax.text(
        1.05,
        6.38,
        "Nature-inspired forms enter as registered hypotheses, not as evidence. Promotion requires the same source-native gate as every other candidate.",
        fontsize=10.1,
        color="#5A4616",
        va="center",
    )
    ax.annotate(
        "",
        xy=(6.0, 5.55),
        xytext=(6.0, 5.98),
        arrowprops=dict(arrowstyle="->", lw=1.5, color="#B78A2E"),
    )
    ax.text(
        6,
        0.55,
        "A missing, stale, inconsistent, unsealed, underpowered, or independently unverified link blocks claim promotion.",
        ha="center",
        fontsize=9.5,
        color="#9B2C2C",
        weight="bold",
    )
    fig.savefig(ARCHITECTURE_FIGURE, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def render_maturity_figure(current_level: int) -> None:
    FIGURE_ROOT.mkdir(parents=True, exist_ok=True)
    fig, ax = plt.subplots(figsize=(12, 4.2), dpi=180)
    ax.set_xlim(-0.2, 6.2)
    ax.set_ylim(0, 2.6)
    ax.axis("off")
    fig.patch.set_facecolor("white")
    labels = [
        ("0", "Concept", "Registered scope"),
        ("1", "Implemented", "Runnable + tested"),
        ("2", "Synthetic", "Frozen benchmark"),
        ("3", "Replay", "Hashed real sources"),
        ("4", "Prospective", "Post-freeze outcomes"),
        ("5", "Independent", "External evaluator"),
    ]
    for level, title, detail in labels:
        value = int(level)
        if value < current_level:
            face, edge = "#E8F3F3", "#0A6B72"
        elif value == current_level:
            face, edge = "#F7F1E3", "#B78A2E"
        else:
            face, edge = "#F2F4F6", "#8B949E"
        box = FancyBboxPatch(
            (value + 0.03, 0.75),
            0.94,
            1.15,
            boxstyle="round,pad=0.03,rounding_size=0.04",
            linewidth=2.0 if value == current_level else 1.2,
            edgecolor=edge,
            facecolor=face,
        )
        ax.add_patch(box)
        ax.text(value + 0.5, 1.62, level, ha="center", fontsize=17, weight="bold", color=edge)
        ax.text(value + 0.5, 1.29, title, ha="center", fontsize=9.7, weight="bold", color="#17324D")
        ax.text(value + 0.5, 0.98, detail, ha="center", fontsize=7.8, color="#5F6B76")
        if value < 5:
            ax.annotate(
                "",
                xy=(value + 1.02, 1.32),
                xytext=(value + 0.98, 1.32),
                arrowprops=dict(arrowstyle="->", lw=1.0, color="#8B949E"),
            )
    ax.text(
        3.5,
        2.25,
        "Current repository-wide supported maturity: Level 3",
        ha="center",
        fontsize=13,
        weight="bold",
        color="#17324D",
    )
    ax.text(
        4.95,
        0.38,
        "Levels 4 and 5 remain closed",
        ha="center",
        fontsize=9.5,
        color="#9B2C2C",
        weight="bold",
    )
    fig.savefig(MATURITY_FIGURE, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def materialize_markdown_layout(markdown: str) -> str:
    replacements = {
        "{PAGE_BREAK}": "<!-- PAGE_BREAK -->",
        "{FIGURE:maturity}": (
            "![Evidence maturity ladder with the current repository at Level 3]"
            "(figures/evidence_maturity_ladder.png)\n\n"
            "*Figure 1. Claim-specific evidence maturity. Levels 4 and 5 remain closed.*"
        ),
        "{FIGURE:architecture}": (
            "![Evidence-governed architecture from authorized sources through independent review]"
            "(figures/evidence_governed_architecture.png)\n\n"
            "*Figure 2. Evidence-governed architecture. Nature-inspired forms enter as hypotheses.*"
        ),
        "{FIGURE:concept_overlay}": "<!-- FIGURE:concept_overlay -->",
    }
    for marker, replacement in replacements.items():
        markdown = markdown.replace(marker, replacement)
    return markdown


def render_master_markdown(payload: dict[str, Any]) -> str:
    s = payload["source_native"]
    eia = payload["eia_hourly"]
    eia_v2 = payload["eia_authority_v2"]
    eia_v3 = payload["eia_hybrid_v3"]
    v3 = payload["source_native_v3"]
    crosswalk = payload["supersession_crosswalk"]
    source_rows = []
    for record in payload["source_ledger"]["records"]:
        source_rows.append(
            "| `{source_id}` | `{claim_class}` | `{disposition}` | {use} |".format(
                **record
            )
        )
    sources_table = "\n".join(
        [
            "| Source | Class | Disposition | Master-paper treatment |",
            "|---|---|---|---|",
            *source_rows,
        ]
    )
    receipts = "\n".join(
        f"- `{item['path']}` - `{item['bytes']}` bytes - SHA-256 `{item['sha256']}`"
        for item in payload["canonical_source_receipts"]
    )
    crosswalk_rows = "\n".join(
        "| `{version}` | `{protocol_role}` | `{disposition}` | `{ready}` |".format(
            version=row["version"],
            protocol_role=row["protocol_role"],
            disposition=row["disposition"],
            ready=str(row["performance_claim_ready"]).lower(),
        )
        for row in crosswalk["versions"]
    )
    markdown = f"""# LumenCore Master Whitepaper

## An Evidence-Governed Platform for Source-Native Benchmarking, Prospective Validation, and Nature-Inspired Systems Research

- Responsible author: **Robert Ashworth**
- Affiliation: **LumenCore**
- Edition: **Private review draft 0.1**
- Generated UTC: `{payload['generated_utc']}`
- Status: `{STATUS}`
- Peer reviewed: `false`
- Independently validated: `false`
- Publication authorized: `false`
- Paper payload SHA-256: `{payload['paper_payload_sha256']}`

> {BOUNDARY}

<!-- BODY -->

## Abstract

LumenCore began as a collection of nature-inspired computing, geometry, sensing, and control concepts. The current repository has evolved into a more defensible contribution: an evidence-governed research platform that registers candidate families, binds each candidate to an authorized source and task contract, compares it with named source-native baselines, preserves losses and inconclusive results, seals prospective protocols, and blocks consequential actions until a human or independent reviewer supplies the missing authority.

This paper reconciles those two histories. It reports what is implemented, what has been measured locally, what is prospectively collecting, what remains a hardware or algorithm hypothesis, and what is quarantined as unsupported speculation. The repository-wide supported maturity is Level {payload['repository_maturity_level']}: source-conditioned replay. The current source-native ledger contains {s['registered_family_count']} registered families, {s['implementation_present_count']} implementations, and {s['comparison_count']} executed direct candidate-source-baseline comparisons. No individual comparison is positive after the current global Holm correction and no candidate passes the complete promotion gate. The legacy EIA V1 lane has {eia['prediction_count']} sealed predictions and {eia['settlement_count']} settlements with all sample gates closed. The preserved all-authority V2 parent has {eia_v2['prediction_panel_count']} sealed panels and {eia_v2['settlement_panel_count']} settled panels; its preliminary sample threshold is met, but no promotion evaluation or performance claim is complete. The separate V3 hybrid confirmation has {eia_v3['prediction_panel_count']} sealed panels, {eia_v3['settlement_panel_count']} settled panels, and {eia_v3['complete_utc_day_count']} complete UTC days, with scores suppressed. These are collection and software-custody facts, not performance claims.

The central thesis is therefore narrow: nature-inspired forms can be useful generators of testable candidate families, but inspiration becomes evidence only through controlled comparison, prospective custody, falsification, and independent review. LumenCore's strongest present contribution is the machinery that makes that distinction inspectable.

**Keywords:** evidence governance; source-native benchmarking; prospective evaluation; negative results; reproducibility; time-series forecasting; multi-agent systems; biomimetic design; human-in-the-loop control.

{{PAGE_BREAK}}

## 1. Scientific Posture

### 1.1 Research question

For a named source, series, cadence, horizon, eligible population, and decision cutoff, can a predeclared candidate beat every accepted baseline by a meaningful effect while surviving leakage controls, dependence-aware inference, multiple-testing correction, sample gates, and independent reproduction?

This question is deliberately more demanding than asking whether one chart, one split, or one metric looks favorable. It defines the comparison unit before the outcome and treats an honest failure as useful information.

### 1.2 Claim taxonomy

| Class | Meaning | Allowed language | Prohibited promotion |
|---|---|---|---|
| Implemented | Runnable code, explicit input contract, deterministic tests | "Implemented for the named bounded workflow" | Production, safety, or field-readiness claims |
| Locally measured | Dated result from a named local source and protocol | Exact metric, split, baseline, and limitation | Universal superiority or independent validation |
| Protocol-sealed | Candidate, cutoff, metrics, and gates fixed before eligible outcomes | Collection state and exact gate status | Accuracy or economic claims before gates pass |
| Conceptual | Mechanism or design variable proposed for testing | "Hypothesis," "candidate," or "proposed experiment" | Benefit language stated as observed fact |
| Speculative quarantine | Claim lacks an accepted mechanism, measurable variables, or credible source chain | Historical disposition only | Scientific, funding, customer, or patent-support use |

No term in the older papers - including harmonic, spiral, resonance, coherence, FlowForm, LumenShell, or sacred geometry - bypasses this taxonomy. A compelling name does not lower the evidence threshold.

### 1.3 Evidence maturity

{{FIGURE:maturity}}

The repository's Level 3 status is claim-specific and does not imply product readiness, agency approval, patent validity, cybersecurity accreditation, or economic value. Level 4 requires post-freeze observations and complete preregistered gates. Level 5 requires a named outcome-independent evaluator and dated receipt. [I1, I2]

{{PAGE_BREAK}}

## 2. Platform Vocabulary and Architecture

LumenCore is the technical platform and research program. LumaTrader is the quantitative-market experimentation lane. NovaStack is the orchestration and context layer. ProofLock is the custody and evidence-control layer. HumanUnlock is the required human decision before submissions, certifications, spending, account changes, external sends, legal actions, or live orders. These names describe bounded roles rather than independent proof of capability. [I1]

{{FIGURE:architecture}}

### 2.1 Authorized source layer

Every measured claim begins with an authorized public source or buyer-approved source. The source receipt binds identity, retrieval time, byte hash, eligible population, exclusions, and version. A zero-row, stale, missing, or incompatible source produces abstention or a blocker, not a silent substitution.

### 2.2 Source-task contract

The contract fixes the target variable, cadence, horizon, prediction cutoff, available features, missing-data rule, seasonal structure, and scoring population. This prevents a candidate from being evaluated on one task while being described as useful for another.

### 2.3 Candidate and baseline registry

Nature-inspired candidate families enter the same registry as statistical, machine-learning, and rule-based candidates. Each source has a complete accepted baseline roster. A candidate is not promoted because it beats one weak baseline; it must clear every predeclared comparator and every guardrail.

### 2.4 Protocol and custody layer

Frozen protocols define splits, metrics, effect floors, sample floors, multiplicity correction, ablations, and failure states. ProofLock then binds code, input snapshots, predictions, settlements, manifests, and reviewer artifacts with hashes and append-only links. A hash demonstrates byte identity and chronology; it does not demonstrate scientific correctness.

### 2.5 Reviewer and action layer

NOAHS is the reviewer proof-chain architecture. It asks whether the required custody links are current, internally consistent, hash-verifiable, preregistered where applicable, reproducible for the exact release, and independently validated. It does not calculate model gains. HumanUnlock keeps consequential actions outside the autonomous evidence pipeline. [I6]

## 3. Source-Native Evaluation Method

### 3.1 Comparison unit

The authoritative unit is a candidate-by-source card evaluated against the baseline roster registered for that source, series, cadence, and horizon. Expanding history ends immediately before each forecast origin. Overlapping origins and horizons are not treated as independent observations.

### 3.2 Baseline discipline

The current source-native roster includes naive persistence, drift, moving average, exponential smoothing, linear trend, source-period seasonal naive, damped Holt/ETS, and autoregressive ridge baselines. Other bounded lanes add named gradient-boosting, neural, classical, and official-source comparators when those methods fit the source contract. [I2, I7]

### 3.3 Multiplicity and effect floors

The current protocol applies Holm step-down correction across the declared family of comparisons. Statistical significance alone is insufficient: the candidate must also satisfy a predeclared effect floor, cell-level guardrails, tail-error controls, coverage rules, and complete sample gates. [R1]

### 3.4 Negative-result retention

LumenCore preserves failed and inconclusive gates rather than relabeling them as wins. In the current EIA wave benchmark, geometry-inspired wave candidates lost to official baselines. A separate residual model produced a favorable aggregate result but failed a worst-authority guardrail. Both outcomes remain visible. This behavior is more scientifically useful than a selectively positive showcase. [I7]

### 3.5 Prospective custody

Predictions must be sealed before the relevant outcome is released, linked to the frozen protocol, and settled against the permitted source vintage. Backfill, post-outcome substitution, silent calendar compression, and sample-gate weakening are prohibited. Collection-state counts may be reported, but no performance conclusion exists until the complete decision rule is evaluable.

## 4. Current Evidence State as of 2026-08-02

### 4.1 Source-native ledger

| Measure | Current value | Interpretation |
|---|---:|---|
| Registered candidate families | {s['registered_family_count']} | Inventory breadth, not proof |
| Implementations present | {s['implementation_present_count']} | Runnable bounded families |
| Missing implementations | {s['implementation_required_count']} | Unimplemented hypotheses |
| Candidate-source cards | {s['candidate_source_card_count']} | Source-bound comparison units |
| Executed direct comparisons | {s['comparison_count']} | Candidate-source-baseline comparisons |
| Global Holm-positive comparisons | {s['global_holm_positive_count']} | No corrected individual lead |
| Complete promotion-gate passes | {s['promotion_gate_pass_count']} | No champion |
| Market-signal comparisons | {s['market_signal_comparison_count']} | Retrospective specialist screen |
| Inferentially insufficient market comparisons | {s['market_signal_inference_insufficient_count']} | No market promotion |

The current result is not "nothing worked." It is that the registered evidence does not support a promoted champion. The platform has learned which prior leads fail under complete baselines, clustering, or multiplicity. That is a legitimate research output.

### 4.2 EIA hourly V1 specialist-router lane

| Measure | Current value |
|---|---:|
| State | `{eia['state']}` |
| Sealed predictions | {eia['prediction_count']} |
| Settlements | {eia['settlement_count']} |
| Common settled hours | {eia['common_settled_hour_count']} |
| First common period | `{eia['first_common_settled_period']}` |
| Latest common period | `{eia['latest_common_settled_period']}` |
| Preliminary gate | `{str(eia['sample_gates']['preliminary_ready']).lower()}` |
| Confirmatory gate | `{str(eia['sample_gates']['confirmatory_ready']).lower()}` |
| Durability gate | `{str(eia['sample_gates']['durability_ready']).lower()}` |
| Promotion evaluation complete | `false` |

These counts belong to the frozen hourly specialist-router lane. They do not imply that the router is better than the current fixed candidate, improve grid reliability, reduce outage cost, or provide deployment authority.

### 4.3 EIA hourly V2 all-authority parent

| Measure | Current value |
|---|---:|
| State | `{eia_v2['state']}` |
| Prediction panels | {eia_v2['prediction_panel_count']} |
| Settlement panels | {eia_v2['settlement_panel_count']} |
| Sealed authority predictions | {eia_v2['sealed_authority_prediction_count']} |
| Settled authority predictions | {eia_v2['settled_authority_prediction_count']} |
| Common settled hours | {eia_v2['common_settled_hour_count']} |
| Preliminary sample ready | `{str(eia_v2['sample_gates']['preliminary_ready']).lower()}` |
| Confirmatory sample ready | `{str(eia_v2['sample_gates']['confirmatory_ready']).lower()}` |
| Durability sample ready | `{str(eia_v2['sample_gates']['durability_ready']).lower()}` |
| Evidence packet release ready | `{str(eia_v2['evidence_packet_release_ready']).lower()}` |
| Promotion evaluation complete | `{str(eia_v2['promotion_evaluation_complete']).lower()}` |
| Performance claim ready | `{str(eia_v2['performance_claim_ready']).lower()}` |

V2 is preserved as the active parent chain. Its 168-hour preliminary sample threshold has been crossed, and the custody packet meets its minimum completeness threshold for external review. Those are sample and packet-readiness facts only, not a favorable result, promotion decision, or independent validation.

### 4.4 EIA hourly V3 hybrid confirmation

| Measure | Current value |
|---|---:|
| State | `{eia_v3['state']}` |
| Prediction panels | {eia_v3['prediction_panel_count']} |
| Settlement panels | {eia_v3['settlement_panel_count']} |
| Sealed authority predictions | {eia_v3['sealed_authority_prediction_count']} |
| Settled authority predictions | {eia_v3['settled_authority_prediction_count']} |
| Complete UTC days | {eia_v3['complete_utc_day_count']} |
| Scores suppressed | `{str(eia_v3['scores_suppressed']).lower()}` |
| Promotion evaluation complete | `{str(eia_v3['promotion_evaluation_complete']).lower()}` |
| Automatic promotion allowed | `{str(eia_v3['automatic_promotion_allowed']).lower()}` |
| Protocol commit bound | `{str(eia_v3['protocol_commit_bound']).lower()}` |

V3 is a separate future-only chain; it does not replace V2 evidence or inherit a favorable conclusion. Its current protocol and runtime are byte-hash bound, but the status contains no protocol-commit receipt, so this lane is not yet commit-bound. V4 is deferred until a disjoint temporal replication, and V5 is deferred until a named independent evaluator controls timestamps and reproduction.

### EIA version-succession crosswalk

| Version | Protocol role | Disposition | Performance claim ready |
|---|---|---|---|
{crosswalk_rows}

Crosswalk SHA-256: `{crosswalk['crosswalk_sha256']}`. Version succession is not evidence promotion. Negative, invalid, and inconclusive results remain attached to their original lanes.

### 4.5 Source-native Version 3 prospective lane

Version 3 is a separate future-only protocol: `{v3['protocol_id']}`. It currently contains {v3['prediction_count']} sealed predictions, {v3['settlement_count']} settlements, and {v3['eligible_future_observation_count']} eligible future observations. Its state is `{v3['state']}` and its decision is `{v3['promotion_decision']}`. The external-anchor count is {v3['external_anchor_count']}. These facts block, rather than support, a prospective accuracy claim.

### 4.6 Commercially useful work before model promotion

The strongest near-term product is not a performance promise. ProofLock Opportunity Operations can support a bounded buyer workflow: source-bound triage, ownership routing, reviewer-ready drafting, preflight checks, provenance, abstention, and action receipts. A separate fixed-scope evidence-protocol sprint can help one buyer freeze an authorized source, incumbent baselines, metrics, exclusions, and replay bundle. Pricing, recipient, data terms, and any external send remain separately approved. [I8, I9]

## 5. Nature-Inspired Research Lineage

### 5.1 What survives scientific translation

The early papers repeatedly proposed curved paths, spiral routing, honeycomb packing, oscillator coherence, distributed sensing, adaptive control, and modular agent networks. Several of those ideas map to established engineering questions when stripped of benefit language:

- Does a conformal or curved layout change thermal resistance, pressure drop, signal path length, impedance, EMI, mechanical strain, or manufacturability relative to a matched flat control?
- Does honeycomb packing improve surface-area-to-volume ratio or coolant-channel access after cell count, chemistry, mass, enclosure volume, and safety constraints are held fixed?
- Can a coupled-oscillator controller minimize a defined phase or energy-loss cost more robustly than conventional control under a fixed disturbance set?
- Can distributed agents improve fault isolation or graceful degradation under a preregistered network-failure protocol without creating unsafe emergent actions?
- Can a biometric interface classify a bounded user state with consent, calibration, latency, false-trigger, and privacy controls?

These are testable questions. None carries an assumed positive answer.

### 5.2 FlowForm

FlowForm is retained as a family of geometry and packaging hypotheses, not as evidence that energy "prefers" curved paths. A defensible first experiment would compare flat and curved/conformal test articles with matched materials, copper area, component population, electrical load, airflow boundary conditions, and sensor placement.

Required endpoints include junction and surface temperature, thermal resistance, pressure drop, insertion loss, return loss, impedance discontinuity, radiated and conducted emissions, mechanical strain, defect rate, repairability, and manufacturing yield. Test articles, simulation meshes, and analysis scripts must be frozen before outcome inspection. Patent-sensitive dimensions remain outside this paper.

### 5.3 Harmonic backpropagation

The one-page harmonic-backprop document provides a useful conceptual mapping between neural-network optimization and instrumented resonator control: model parameters correspond to frequency, phase, coupling, gain, damping, and quality factor; loss corresponds to a measurable coherence or leakage cost. That mapping is not yet a demonstrated physical learning system.

The lowest-risk route is software-first. Implement the oscillator state equations, freeze a cost function and disturbance family, compare gradient, derivative-free, and conventional controller baselines, and require stability and energy constraints. Only then should a physical resonator be considered with calibrated phase sensing, lock-in detection, and independent instrumentation. Health, consciousness, or special-frequency claims are outside scope.

### 5.4 LumenShell and AetherFrame

The useful portion of the LumenShell specification is a modular VR testbed: sensor bridge, logic router, bounded scene actions, latency logs, false-trigger rates, uptime, and a user study. Any biometric work requires informed consent, data minimization, access control, and a study protocol appropriate to the risk. The source contains third-party boilerplate and therefore cannot establish clean authorship or implementation.

AetherFrame and LumenKing are best understood as orchestration roadmaps. The current repository's actual implemented control vocabulary is NovaStack, ProofLock, and HumanUnlock. Roadmap statements are not deployment receipts.

### 5.5 Historical concept graphic

{{FIGURE:concept_overlay}}

The geometric overlay is retained as a record of design inspiration. It is not a circuit model, field solution, optimization proof, or measured performance result. A future geometry study must translate each visual feature into explicit parameters and matched controls.

### 5.6 Quarantined claims

The legacy corpus includes BioGeometry benefit claims, passive thermal and EMI improvements, 30-50 percent cooling savings, cognition effects, special-frequency healing, nuclear-emission harmonization, scalar fields, zero-point energy, weather control, consciousness-field interaction, and wormhole-adjacent effects. The current evidence stack does not support these claims. They are excluded from the scientific body, funding packets, customer materials, and publication surfaces unless independently rederived into accepted, measurable physics and passed through the full evidence gate.

{{PAGE_BREAK}}

## 6. Proposed Research Program

### 6.1 Track A - Benchmark and evidence platform

1. Finish implementations only for families with a source-task rationale and a predeclared falsification test.
2. Expand direct-source lanes while preserving source-specific baselines and timing contracts.
3. Reconcile every aggregate count to one canonical generator and current source receipt.
4. Publish bounded replay capsules only from clean, commit-bound worktrees after privacy and license review.

### 6.2 Track B - Independent prospective evaluation

1. Complete the current sample gates without weakening them.
2. Name an outcome-independent evaluator before confirmatory scoring.
3. Transfer only the minimum frozen packet needed to reproduce arithmetic and custody.
4. Require a signed result receipt that permits pass, fail, invalid, or inconclusive.
5. Keep economic conversion closed until an operational owner accepts both the technical metric and the assumptions.

### 6.3 Track C - FlowForm hardware coupon study

1. Freeze one geometry parameter family and one matched conventional control.
2. Obtain design review from electrical, thermal, mechanical, safety, and manufacturing specialists.
3. Simulate both designs with identical boundary conditions, then fabricate blinded coupons.
4. Use calibrated thermal, RF, electrical, and mechanical instrumentation.
5. Report all endpoints, corrections, failures, and manufacturing defects.
6. Replicate at an independent laboratory before any efficiency, EMI, durability, or savings claim.

### 6.4 Track D - Human-interface study

1. Limit the first study to observable interface behavior: classification, latency, false triggers, uptime, and user-reported usability.
2. Obtain informed consent and minimize biometric retention.
3. Exclude diagnostic, therapeutic, consciousness, or healing endpoints.
4. Predefine participant exclusions, calibration, stopping rules, and analysis.

### 6.5 Track E - High-energy and nuclear contexts

Any nuclear or radiation-adjacent work must be partner-led by appropriately licensed facilities and domain experts. The first acceptable scope is simulation, instrumentation, shielding, or control-system evidence using conventional quantities and approved sources. This paper does not propose radiation conversion, healing, field harmonization, or operator-intention control.

## 7. Governance, Security, and Reproducibility

### 7.1 ProofLock controls

- Content hashes bind exact artifacts.
- Manifests bind expected artifact sets, byte counts, and versions.
- Append-only prediction and settlement chains expose deletion or reordering.
- Source receipts bind retrieval time and source identity.
- Negative results and abstentions remain first-class outcomes.
- Claim gates are machine-readable and fail closed.

These controls support auditability, not truth by themselves.

The current WhiteHole audit preserves the historical read-only market-diagnostic
and custody work while excluding its heuristic ranks, stale scheduled tasks, and
legacy website from current performance or deployment evidence. [I10] The pitch
deck governance register separately identifies one current review-required deck
among 31 registered PPTX files and blocks release of every stale, application-
specific, source-template, or high-risk presentation. [I11]

### 7.2 HumanUnlock

HumanUnlock is mandatory before external sends, uploads, submissions, certifications, legal representations, spending, account changes, production deployment, or live orders. The system may prepare a bounded draft and list missing facts; it may not invent the facts or certify them.

### 7.3 AI risk and software security

The governance roadmap should align internal controls to NIST AI RMF functions - Govern, Map, Measure, and Manage - and to the NIST Secure Software Development Framework. Alignment is a work program, not a certification claim. [R6, R7]

### 7.4 Reproducibility packet

Every reviewer-facing result should identify the repository commit, environment lock, source hashes, protocol hash, candidate and baseline versions, eligible population, exact command, output hashes, known failures, and claim boundary. A clean external runner and a named evaluator are still required for independent validation.

## 8. Limitations

- {s['implementation_required_count']} of {s['registered_family_count']} registered families lack implementations.
- Only a subset of lanes has executable direct-source adapters.
- No current source-native candidate passes the complete promotion gate.
- The EIA hourly collection has not reached any preregistered sample gate.
- Source-native Version 3 has zero eligible future observations and no verified external anchor receipt.
- The worktree and public release chain require a clean, selective rebuild before publication.
- No named independent evaluator has completed the current prospective reproduction.
- No current artifact establishes field performance, realized savings, regulatory approval, patent scope, or profitable live trading.
- The legacy corpus contains duplicate, mislabeled, placeholder, private, patent-sensitive, and unsupported material.
- This synthesis is AI-assisted and must receive responsible-author and domain-specialist review before any external use.

## 9. Conclusion

LumenCore's defensible identity is not a claim that geometry or harmonic models already outperform every incumbent. It is a disciplined research program for turning unconventional inspiration into auditable hypotheses. The platform registers candidates, binds them to real source contracts, tests them against named baselines, preserves losses, seals future-only protocols, and refuses to promote a claim when evidence is incomplete.

That posture gives the older ideas a credible path forward. Curved hardware, oscillator control, distributed sensing, and biometric interfaces can be investigated through matched experiments. Unsupported nuclear, healing, scalar-field, zero-point, weather-control, and wormhole claims remain quarantined. The next scientific milestone is not a stronger adjective. It is a completed prospective gate and an independent receipt.

## 10. Authorship and Research Integrity

Robert Ashworth is responsible for the research question, concept ownership, interpretation, release decisions, and scientific claims.

**AI assistance disclosure.** Luma (OpenAI Codex) assisted with corpus extraction, source reconciliation, software implementation, test scaffolding, literature verification, quality assurance, and document production. AI assistance is not evidence, is not listed as authorship, and does not assume responsibility for the work.

**Data and code availability.** Canonical code, protocols, ledgers, and local source receipts are identified by path, byte count, and SHA-256. Public availability is not claimed. Patent archives were not expanded. Private identifiers and patent-sensitive details are omitted.

**Declarations.** Funding, competing-interest, contributor, ownership, and legal declarations require responsible-author confirmation before external release.

{{PAGE_BREAK}}

## Appendix A. Source Disposition Ledger

The corpus contains {payload['corpus']['source_count']} inventoried sources, of which {payload['corpus']['extractable_count']} yielded usable text. Source disposition determines whether a document contributes current evidence, hypothesis context, historical context, or only a sealed receipt.

{sources_table}

## Appendix B. Canonical Internal Receipts

{receipts}

## References

- [I1] LumenCore repository `README.md`, evidence maturity and platform vocabulary, current local receipt above.
- [I2] `docs/LUMENCORE_SOURCE_NATIVE_BENCHMARK_WHITEPAPER_CURRENT.md`, current source-native method, results, and limitations.
- [I3] `out/ops/source_native_family_baseline_ledger_latest.json`, authoritative candidate-source-baseline ledger.
- [I4] `out/eia_grid_prospective_hourly_router/prospective_status_latest.json`, dated hourly collection state.
- [I5] `out/time_series_source_native_prospective_v3/prospective_status_latest.json`, Version 3 future-only state.
- [I6] `docs/NOAHS_REVIEWER_ARCHITECTURE_2026-07-25.md`, conjunctive reviewer proof chain.
- [I7] `docs/PUBLIC_SAFE_MODEL_AND_GEOMETRY_EVIDENCE_LEDGER_2026-07-13.md`, bounded model and geometry evidence.
- [I8] `docs/LUMENCORE_EVIDENCE_PROTOCOL_REVIEW_FIXED_SCOPE_OFFER_2026-07-30.md`, bounded professional-services scope.
- [I9] `docs/PRODUCT_LANE_EVIDENCE_AUDIT_2026-07-29.md`, product evidence and pilot gates.
- [I10] `docs/WHITEHOLE_WHITEHOLELAB_AUDIT_2026-08-02.md`, historical research custody, implementation defects, and current claim boundary.
- [I11] `docs/PITCH_DECK_GOVERNANCE_2026-07-29.md`, current presentation registry and release controls.
- [R1] Holm, S. (1979). A simple sequentially rejective multiple test procedure. *Scandinavian Journal of Statistics*, 6(2), 65-70. https://doi.org/10.2307/4615733
- [R2] Hyndman, R. J., and Koehler, A. B. (2006). Another look at measures of forecast accuracy. *International Journal of Forecasting*, 22(4), 679-688. https://doi.org/10.1016/j.ijforecast.2006.03.001
- [R3] Kunsch, H. R. (1989). The jackknife and the bootstrap for general stationary observations. *The Annals of Statistics*, 17(3), 1217-1241. https://doi.org/10.1214/aos/1176347265
- [R4] White, H. (2000). A reality check for data snooping. *Econometrica*, 68(5), 1097-1126. https://doi.org/10.1111/1468-0262.00152
- [R5] Puente-Baliarda, C., Romeu, J., Pous, R., and Cardama, A. (1998). On the behavior of the Sierpinski multiband fractal antenna. *IEEE Transactions on Antennas and Propagation*, 46(4), 517-524. https://doi.org/10.1109/8.664115
- [R6] Tabassi, E. (2023). *Artificial Intelligence Risk Management Framework (AI RMF 1.0)*. NIST AI 100-1. https://doi.org/10.6028/NIST.AI.100-1
- [R7] Souppaya, M., Scarfone, K., and Dodson, D. (2022). *Secure Software Development Framework (SSDF) Version 1.1*. NIST SP 800-218. https://doi.org/10.6028/NIST.SP.800-218
"""
    return materialize_markdown_layout(markdown)


def audit_private_output(text: str) -> None:
    failures = [name for name, pattern in PRIVACY_PATTERNS.items() if pattern.search(text)]
    if failures:
        raise ValueError(f"Private-output safety audit failed: {', '.join(failures)}")


def audit_layout_markers(text: str) -> None:
    raw_markers = (
        "{PAGE_BREAK}",
        "{{PAGE_BREAK}}",
        "{FIGURE:",
        "{{FIGURE:",
    )
    if any(marker in text for marker in raw_markers):
        raise ValueError("Private-output layout audit failed: raw layout marker remains")


def set_run_font(
    run: Any,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_table_row_together(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell: Any, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table: Any, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    total_dxa = sum(int(width * 1440) for width in widths)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_field(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text, end):
        run._r.append(node)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, NAVY, 16, 8),
        "Heading 2": (13, TEAL, 12, 6),
        "Heading 3": (11.5, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    caption = document.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = True


def add_inline_markdown(paragraph: Any, text: str, default_size: float | None = None) -> None:
    token_pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
    cursor = 0
    for match in token_pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=default_size)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=(default_size or 10.5) - 0.5, color=NAVY)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=default_size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=default_size, italic=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=default_size)


def add_callout(document: Document, text: str, fill: str = PALE_GOLD, border: str = GOLD) -> None:
    table = document.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    add_inline_markdown(paragraph, text, 10.5)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "8")
        node.set(qn("w:color"), border)
        tc_borders.append(node)
    tc_pr.append(tc_borders)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_picture_with_alt(document: Document, path: Path, width: float, alt_text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(path), width=Inches(width))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
            continue
        rows.append(cells)
    return rows


def table_widths(column_count: int) -> list[float]:
    return {
        2: [2.1, 4.4],
        3: [1.7, 1.6, 3.2],
        4: [1.35, 1.35, 1.7, 2.1],
    }.get(column_count, [6.5 / column_count] * column_count)


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = len(rows[0])
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    set_table_geometry(table, table_widths(column_count))
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        keep_table_row_together(table.rows[row_index])
        if len(values) != column_count:
            raise ValueError("Markdown table has inconsistent column count")
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            if row_index == 0:
                set_cell_shading(cell, NAVY)
            elif row_index % 2 == 0:
                set_cell_shading(cell, PALE_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            add_inline_markdown(paragraph, value, 8.6 if column_count >= 4 else 9.2)
            for run in paragraph.runs:
                if row_index == 0:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.bold = True
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_header_footer(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_after = Pt(0)
        run = hp.add_run("LumenCore Master Whitepaper | Private Review")
        set_run_font(run, size=8.5, color=MUTED, bold=True)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        fp.paragraph_format.space_after = Pt(0)
        run = fp.add_run(f"{STATUS}  |  ")
        set_run_font(run, size=8, color=MUTED)
        add_page_field(fp)


def build_docx(payload: dict[str, Any], markdown: str) -> None:
    document = Document()
    configure_styles(document)
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    document.core_properties.title = "LumenCore Master Whitepaper"
    document.core_properties.subject = "Private evidence-governed technical synthesis"
    document.core_properties.author = "Robert Ashworth"
    document.core_properties.keywords = (
        "LumenCore, source-native benchmarking, prospective validation, evidence governance"
    )
    document.core_properties.comments = STATUS

    cover_spacer = document.add_paragraph()
    cover_spacer.paragraph_format.space_after = Pt(8)
    add_picture_with_alt(document, LOGO_PATH, 2.15, "LumaArc geometric eclipse logo")
    seal_caption = document.add_paragraph()
    seal_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    seal_caption.paragraph_format.space_before = Pt(1)
    seal_caption.paragraph_format.space_after = Pt(2)
    run = seal_caption.add_run("LumaArc evidence-governance seal")
    set_run_font(run, size=8.2, color=MUTED)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(5)
    run = title.add_run("LumenCore Master Whitepaper")
    set_run_font(run, size=27, color=NAVY, bold=True)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run(
        "An Evidence-Governed Platform for Source-Native Benchmarking,\n"
        "Prospective Validation, and Nature-Inspired Systems Research"
    )
    set_run_font(run, size=14, color=TEAL, bold=True)

    metadata = document.add_table(rows=4, cols=2)
    set_table_geometry(metadata, [1.75, 4.75])
    metadata_rows = [
        ("Responsible author", "Robert Ashworth"),
        ("Affiliation", "LumenCore"),
        ("Edition", "Private review draft 0.1 | 2026-08-02"),
        ("Evidence posture", f"Repository Level {payload['repository_maturity_level']}; no independent validation"),
    ]
    for index, (label, value) in enumerate(metadata_rows):
        left, right = metadata.rows[index].cells
        set_cell_shading(left, PALE_BLUE)
        for cell in (left, right):
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        lr = left.paragraphs[0].add_run(label)
        set_run_font(lr, size=9.4, color=NAVY, bold=True)
        rr = right.paragraphs[0].add_run(value)
        set_run_font(rr, size=9.4, color=INK)
    document.add_paragraph()
    add_callout(document, f"**{STATUS}.** {BOUNDARY}", PALE_GOLD, GOLD)
    receipt = document.add_paragraph()
    receipt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    receipt.paragraph_format.space_before = Pt(10)
    run = receipt.add_run(f"Paper payload SHA-256: {payload['paper_payload_sha256']}")
    set_run_font(run, name="Consolas", size=7.6, color=MUTED)
    document.add_page_break()

    body = markdown.split("<!-- BODY -->", 1)[1]
    lines = body.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped in {"{{PAGE_BREAK}}", "{PAGE_BREAK}", "<!-- PAGE_BREAK -->"}:
            document.add_page_break()
            index += 1
            continue
        if stripped in {"{{FIGURE:maturity}}", "{FIGURE:maturity}"}:
            add_picture_with_alt(
                document,
                MATURITY_FIGURE,
                6.3,
                "Evidence maturity ladder with current repository status at Level 3",
            )
            caption = document.add_paragraph(style="Caption")
            caption.add_run(
                "Figure 1. Claim-specific evidence maturity. Levels 4 and 5 remain closed."
            )
            index += 1
            continue
        if stripped in {"{{FIGURE:architecture}}", "{FIGURE:architecture}"}:
            add_picture_with_alt(
                document,
                ARCHITECTURE_FIGURE,
                6.35,
                "Evidence-governed architecture from authorized sources through independent review",
            )
            caption = document.add_paragraph(style="Caption")
            caption.add_run(
                "Figure 2. Evidence-governed architecture. Nature-inspired forms enter as hypotheses."
            )
            index += 1
            continue
        if stripped == (
            "![Evidence maturity ladder with the current repository at Level 3]"
            "(figures/evidence_maturity_ladder.png)"
        ):
            add_picture_with_alt(
                document,
                MATURITY_FIGURE,
                6.3,
                "Evidence maturity ladder with current repository status at Level 3",
            )
            index += 1
            continue
        if stripped == (
            "![Evidence-governed architecture from authorized sources through independent review]"
            "(figures/evidence_governed_architecture.png)"
        ):
            add_picture_with_alt(
                document,
                ARCHITECTURE_FIGURE,
                6.35,
                "Evidence-governed architecture from authorized sources through independent review",
            )
            index += 1
            continue
        if stripped in {
            "{{FIGURE:concept_overlay}}",
            "{FIGURE:concept_overlay}",
            "<!-- FIGURE:concept_overlay -->",
        }:
            add_picture_with_alt(
                document,
                CONCEPT_OVERLAY,
                4.45,
                "Historical geometric overlay combining circular construction and golden-ratio rectangles",
            )
            caption = document.add_paragraph(style="Caption")
            caption.add_run(
                "Figure 3. Historical design-inspiration graphic. It is not performance evidence."
            )
            index += 1
            continue
        caption_match = re.fullmatch(r"\*(Figure \d+\..+)\*", stripped)
        if caption_match:
            caption = document.add_paragraph(style="Caption")
            caption.add_run(caption_match.group(1))
            index += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(document, parse_table(table_lines))
            continue
        if stripped.startswith("### "):
            paragraph = document.add_paragraph(style="Heading 2")
            add_inline_markdown(paragraph, stripped[4:])
            index += 1
            continue
        if stripped.startswith("## "):
            paragraph = document.add_paragraph(style="Heading 1")
            add_inline_markdown(paragraph, stripped[3:])
            index += 1
            continue
        if stripped.startswith("# "):
            paragraph = document.add_paragraph(style="Heading 1")
            add_inline_markdown(paragraph, stripped[2:])
            index += 1
            continue
        if stripped.startswith("> "):
            add_callout(document, stripped[2:], PALE_GOLD, GOLD)
            index += 1
            continue
        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(" ")
            is_reference = bool(re.match(r"^- \[(?:I|R)\d+\]", stripped))
            if is_reference:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.08
            add_inline_markdown(
                paragraph,
                stripped[2:],
                default_size=9.5 if is_reference else None,
            )
            index += 1
            continue
        number_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if number_match:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(" ")
            add_inline_markdown(paragraph, number_match.group(1))
            index += 1
            continue
        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate:
                index += 1
                break
            if (
                candidate.startswith(
                    ("#", "- ", "> ", "|", "{{", "{", "<!--", "![", "*Figure ")
                )
                or re.match(r"^\d+\.\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph = document.add_paragraph()
        add_inline_markdown(paragraph, " ".join(paragraph_lines))

    add_header_footer(document)
    document.save(OUTPUT_DOCX)


def extract_docx_text(path: Path) -> str:
    document = Document(path)
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def find_soffice() -> Path:
    candidates = [
        Path("C:/Program Files/LibreOffice/program/soffice.com"),
        ROOT.parent / ".tools" / "LibreOffice" / "program" / "soffice.com",
        Path("C:/Program Files/LibreOffice/program/soffice.exe"),
    ]
    for command in ("soffice.com", "soffice.exe", "soffice"):
        discovered = shutil.which(command)
        if discovered:
            candidates.append(Path(discovered))
    for candidate in candidates:
        if candidate.is_file():
            return candidate.resolve()
    raise ValueError("LibreOffice is required for the canonical private-review PDF")


def convert_docx_to_pdf() -> None:
    OUTPUT_PDF.unlink(missing_ok=True)
    profile = Path(tempfile.mkdtemp(prefix="lumencore-master-whitepaper-lo-"))
    try:
        result = subprocess.run(
            [
                str(find_soffice()),
                "--headless",
                f"-env:UserInstallation={profile.resolve().as_uri()}",
                "--convert-to",
                "pdf",
                "--outdir",
                str(OUTPUT_ROOT),
                str(OUTPUT_DOCX),
            ],
            cwd=ROOT,
            check=False,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=120,
        )
        if result.returncode == 0:
            deadline = time.monotonic() + 30
            while not OUTPUT_PDF.is_file() and time.monotonic() < deadline:
                time.sleep(0.25)
        if result.returncode != 0 or not OUTPUT_PDF.is_file():
            raise ValueError(
                "Canonical private-review PDF conversion failed: "
                + (result.stderr or result.stdout or "unknown LibreOffice error").strip()
            )
    finally:
        shutil.rmtree(profile, ignore_errors=True)


def extract_pdf_text(path: Path) -> tuple[str, int, bool]:
    reader = PdfReader(path)
    if reader.is_encrypted:
        return "", len(reader.pages), True
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    return text, len(reader.pages), False


def inspect_canonical_pdf(payload_sha256: str) -> dict[str, Any]:
    text, pages, encrypted = extract_pdf_text(OUTPUT_PDF)
    result = {
        "pages": pages,
        "encrypted": encrypted,
        "searchable": len(text.strip()) >= 5000,
        "payload_sha256_present": payload_sha256 in text,
        "private_review_banner_present": STATUS in text,
    }
    result["all_checks_pass"] = bool(
        pages > 0
        and not encrypted
        and result["searchable"]
        and result["payload_sha256_present"]
        and result["private_review_banner_present"]
    )
    return result


def validate_canonical_source_receipts(payload: dict[str, Any]) -> None:
    receipts = payload.get("canonical_source_receipts")
    if not isinstance(receipts, list) or not receipts:
        raise ValueError("Canonical source receipt contract is missing")
    for receipt in receipts:
        if not isinstance(receipt, dict) or not receipt.get("exists"):
            raise ValueError("Canonical source receipt is invalid")
        path_value = receipt.get("path")
        if not isinstance(path_value, str):
            raise ValueError("Canonical source receipt path is invalid")
        path = Path(path_value)
        if not path.is_absolute():
            path = ROOT / path
        current = file_receipt(path)
        if current["bytes"] != receipt.get("bytes") or current["sha256"] != receipt.get(
            "sha256"
        ):
            raise ValueError(f"Canonical source changed during reseal: {path}")


def write_outputs(payload: dict[str, Any]) -> None:
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    render_architecture_figure()
    render_maturity_figure(payload["repository_maturity_level"])
    source_ledger = payload["source_ledger"]
    SOURCE_LEDGER_JSON.write_text(
        json.dumps(source_ledger, indent=2, ensure_ascii=True) + "\n",
        encoding="utf-8",
    )
    SOURCE_LEDGER_MD.write_text(
        render_source_ledger_markdown(source_ledger), encoding="utf-8"
    )
    crosswalk = payload["supersession_crosswalk"]
    CROSSWALK_JSON.write_text(
        json.dumps(crosswalk, indent=2, ensure_ascii=True) + "\n",
        encoding="utf-8",
    )
    CROSSWALK_MD.write_text(
        render_crosswalk_markdown(crosswalk), encoding="utf-8"
    )
    markdown = render_master_markdown(payload)
    audit_private_output(markdown)
    audit_layout_markers(markdown)
    OUTPUT_MD.write_text(markdown, encoding="utf-8")
    build_docx(payload, markdown)
    output_docx_text = extract_docx_text(OUTPUT_DOCX)
    audit_private_output(output_docx_text)
    audit_layout_markers(output_docx_text)
    convert_docx_to_pdf()
    pdf_text, _, _ = extract_pdf_text(OUTPUT_PDF)
    audit_private_output(pdf_text)
    audit_layout_markers(pdf_text)


def validate_generated_outputs(payload: dict[str, Any]) -> None:
    for path in GENERATED_ARTIFACTS:
        require_file(path)

    expected_markdown = render_master_markdown(payload)
    actual_markdown = OUTPUT_MD.read_text(encoding="utf-8")
    if actual_markdown != expected_markdown:
        raise ValueError(
            "Generated master whitepaper is stale relative to the current source payload"
        )

    payload_sha256 = payload.get("paper_payload_sha256")
    if not isinstance(payload_sha256, str) or len(payload_sha256) != 64:
        raise ValueError("Master-whitepaper payload hash is invalid")
    if payload_sha256 not in extract_docx_text(OUTPUT_DOCX):
        raise ValueError(
            "Generated master-whitepaper DOCX is stale relative to the current source payload"
        )

    source_ledger = payload.get("source_ledger")
    if not isinstance(source_ledger, dict):
        raise ValueError("Master-whitepaper source ledger is missing")
    if read_json(SOURCE_LEDGER_JSON) != source_ledger:
        raise ValueError("Generated source ledger JSON is stale")
    if SOURCE_LEDGER_MD.read_text(encoding="utf-8") != render_source_ledger_markdown(
        source_ledger
    ):
        raise ValueError("Generated source ledger Markdown is stale")

    crosswalk = payload.get("supersession_crosswalk")
    if not isinstance(crosswalk, dict):
        raise ValueError("Master-whitepaper supersession crosswalk is missing")
    if read_json(CROSSWALK_JSON) != crosswalk:
        raise ValueError("Generated supersession crosswalk JSON is stale")
    if CROSSWALK_MD.read_text(encoding="utf-8") != render_crosswalk_markdown(
        crosswalk
    ):
        raise ValueError("Generated supersession crosswalk Markdown is stale")

    pdf_inspection = inspect_canonical_pdf(payload_sha256)
    if not pdf_inspection["all_checks_pass"]:
        raise ValueError(f"Canonical private-review PDF failed QA: {pdf_inspection}")

    output_docx_text = extract_docx_text(OUTPUT_DOCX)
    audit_private_output(actual_markdown)
    audit_layout_markers(actual_markdown)
    audit_private_output(output_docx_text)
    audit_layout_markers(output_docx_text)
    pdf_text, _, _ = extract_pdf_text(OUTPUT_PDF)
    audit_private_output(pdf_text)
    audit_layout_markers(pdf_text)


def write_manifest(payload: dict[str, Any]) -> dict[str, Any]:
    validate_canonical_source_receipts(payload)
    validate_generated_outputs(payload)
    outputs = [file_receipt(path) for path in GENERATED_ARTIFACTS]
    pdf_inspection = inspect_canonical_pdf(payload["paper_payload_sha256"])
    manifest: dict[str, Any] = {
        "schema": "lumencore_private_master_whitepaper_manifest_v2",
        "status": STATUS,
        "generated_utc": payload["generated_utc"],
        "paper_payload_sha256": payload["paper_payload_sha256"],
        "source_ledger_sha256": payload["source_ledger"]["ledger_sha256"],
        "publication_authorized": False,
        "peer_reviewed": False,
        "independently_validated": False,
        "canonical_source_receipts": payload["canonical_source_receipts"],
        "supersession_crosswalk_sha256": payload["supersession_crosswalk"][
            "crosswalk_sha256"
        ],
        "canonical_pdf": {
            "receipt": file_receipt(OUTPUT_PDF),
            "inspection": pdf_inspection,
        },
        "generated_artifacts": outputs,
    }
    manifest["manifest_sha256"] = stable_hash(manifest)
    OUTPUT_MANIFEST.write_text(
        json.dumps(manifest, indent=2, ensure_ascii=True) + "\n",
        encoding="utf-8",
    )
    return manifest


def verify_manifest() -> dict[str, Any]:
    manifest = read_json(OUTPUT_MANIFEST)
    if manifest.get("schema") != "lumencore_private_master_whitepaper_manifest_v2":
        raise ValueError("Master-whitepaper manifest schema is stale")
    expected_hash = manifest.get("manifest_sha256")
    unsigned = dict(manifest)
    unsigned.pop("manifest_sha256", None)
    if expected_hash != stable_hash(unsigned):
        raise ValueError("Master-whitepaper manifest hash mismatch")
    if manifest.get("status") != STATUS or manifest.get("publication_authorized") is not False:
        raise ValueError("Master-whitepaper publication gate is invalid")
    generated_artifacts = manifest.get("generated_artifacts")
    if not isinstance(generated_artifacts, list):
        raise ValueError("Master-whitepaper generated artifact list is invalid")
    expected_generated_paths = [file_receipt(path)["path"] for path in GENERATED_ARTIFACTS]
    actual_generated_paths = [
        receipt.get("path") if isinstance(receipt, dict) else None
        for receipt in generated_artifacts
    ]
    if actual_generated_paths != expected_generated_paths:
        raise ValueError("Master-whitepaper generated artifact contract mismatch")
    for receipt in (
        manifest.get("canonical_source_receipts", [])
        + generated_artifacts
    ):
        if not isinstance(receipt, dict) or not receipt.get("exists"):
            raise ValueError("Manifest contains a missing artifact")
        path_value = receipt.get("path")
        if not isinstance(path_value, str):
            raise ValueError("Manifest artifact path is invalid")
        path = Path(path_value)
        if not path.is_absolute():
            path = ROOT / path
        current = file_receipt(path)
        if current["bytes"] != receipt.get("bytes") or current["sha256"] != receipt.get(
            "sha256"
        ):
            raise ValueError(f"Artifact receipt mismatch: {path}")
    output_markdown = OUTPUT_MD.read_text(encoding="utf-8")
    output_docx_text = extract_docx_text(OUTPUT_DOCX)
    payload_sha256 = manifest.get("paper_payload_sha256")
    if not isinstance(payload_sha256, str) or (
        payload_sha256 not in output_markdown or payload_sha256 not in output_docx_text
    ):
        raise ValueError("Master-whitepaper artifacts are not bound to the sealed payload")
    source_ledger = read_json(SOURCE_LEDGER_JSON)
    if source_ledger.get("ledger_sha256") != manifest.get("source_ledger_sha256"):
        raise ValueError("Master-whitepaper source ledger is not bound to the manifest")
    if SOURCE_LEDGER_MD.read_text(encoding="utf-8") != render_source_ledger_markdown(
        source_ledger
    ):
        raise ValueError("Master-whitepaper source ledger Markdown is stale")
    crosswalk = read_json(CROSSWALK_JSON)
    if crosswalk.get("crosswalk_sha256") != manifest.get(
        "supersession_crosswalk_sha256"
    ):
        raise ValueError("Master-whitepaper supersession crosswalk is not bound")
    if CROSSWALK_MD.read_text(encoding="utf-8") != render_crosswalk_markdown(
        crosswalk
    ):
        raise ValueError("Master-whitepaper supersession crosswalk Markdown is stale")
    canonical_pdf = manifest.get("canonical_pdf")
    if not isinstance(canonical_pdf, dict):
        raise ValueError("Canonical private-review PDF receipt is missing")
    if canonical_pdf.get("receipt") != file_receipt(OUTPUT_PDF):
        raise ValueError("Canonical private-review PDF receipt mismatch")
    if canonical_pdf.get("inspection") != inspect_canonical_pdf(str(payload_sha256)):
        raise ValueError("Canonical private-review PDF inspection is stale")
    if not canonical_pdf["inspection"].get("all_checks_pass"):
        raise ValueError("Canonical private-review PDF QA is not passing")
    audit_private_output(output_markdown)
    audit_layout_markers(output_markdown)
    audit_private_output(output_docx_text)
    audit_layout_markers(output_docx_text)
    pdf_text, _, _ = extract_pdf_text(OUTPUT_PDF)
    audit_private_output(pdf_text)
    audit_layout_markers(pdf_text)
    return manifest


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--as-of-utc")
    parser.add_argument("--seal", action="store_true")
    parser.add_argument("--reseal", action="store_true")
    parser.add_argument("--check", action="store_true")
    args = parser.parse_args()
    if args.check:
        result = verify_manifest()
    else:
        at = (
            datetime.fromisoformat(args.as_of_utc.replace("Z", "+00:00"))
            if args.as_of_utc
            else datetime.now(timezone.utc)
        )
        payload = build_payload(at)
        if args.seal or args.reseal:
            write_outputs(payload)
            validate_canonical_source_receipts(payload)
            result = write_manifest(payload)
        else:
            write_outputs(payload)
            result = {
                "schema": payload["schema"],
                "status": payload["status"],
                "repository_maturity_level": payload["repository_maturity_level"],
                "paper_payload_sha256": payload["paper_payload_sha256"],
                "output_docx": str(OUTPUT_DOCX),
                "publication_authorized": payload["publication_authorized"],
            }
    print(json.dumps(result, indent=2, ensure_ascii=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
