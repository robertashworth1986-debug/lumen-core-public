"""Build the public-safe CDC AI acquisition support capability statement."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = RGBColor(31, 78, 121)
INK = RGBColor(28, 36, 44)
MUTED = RGBColor(88, 96, 105)
LIGHT = "F2F4F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_bottom_border(paragraph, color="D9DEE5", size="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, before, after in (
        ("Heading 1", 15, 12, 6),
        ("Heading 2", 12.5, 9, 4),
        ("Heading 3", 11.5, 6, 3),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("Capability Bullet", "Capability Number"):
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.base_style = normal
        style.font.name = "Calibri"
        style.font.size = Pt(10.25)
        style.paragraph_format.left_indent = Inches(0.43)
        style.paragraph_format.first_line_indent = Inches(-0.2)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.05


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Capability Bullet")
    paragraph.style = doc.styles["List Bullet"]
    paragraph.paragraph_format.left_indent = Inches(0.43)
    paragraph.paragraph_format.first_line_indent = Inches(-0.2)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.05
    paragraph.add_run(text)


def add_numbered(doc: Document, label: str, text: str) -> None:
    paragraph = doc.add_paragraph(style="Capability Number")
    run = paragraph.add_run(f"{label}. ")
    run.bold = True
    run.font.color.rgb = BLUE
    paragraph.add_run(text)


def add_metadata_table(doc: Document) -> None:
    rows = [
        ("Respondent", "LumenCore", "Notice", "75D301-26-RFI-73483"),
        ("Contact", "Robert Ashworth", "Status", "Prototype-stage / in development"),
        ("Location", "Nashville, Tennessee", "Response", "AI acquisition support capability"),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    set_table_geometry(table, [1200, 3480, 1150, 3530])
    for row_index, row_data in enumerate(rows):
        for col_index, value in enumerate(row_data):
            cell = table.cell(row_index, col_index)
            cell.text = value
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.runs[0]
            run.font.size = Pt(9.25)
            if col_index % 2 == 0:
                run.bold = True
                run.font.color.rgb = BLUE
                set_cell_shading(cell, LIGHT)
            else:
                set_cell_shading(cell, WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "LumenCore | CDC AI Acquisition Support RFI"
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header.paragraph_format.space_after = Pt(0)
        for run in header.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(8.5)
            run.font.color.rgb = MUTED

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("75D301-26-RFI-73483 | Market research response | Public-safe")
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = MUTED


def build(output: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    configure_styles(doc)
    add_header_footer(doc)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(3)
    run = kicker.add_run("CAPABILITY STATEMENT")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Evidence-Bounded AI Support for the Federal Acquisition Lifecycle")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = INK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(9)
    run = subtitle.add_run("CDC Reverse Industry Day | Notice 75D301-26-RFI-73483")
    run.font.size = Pt(10.5)
    run.font.color.rgb = MUTED
    add_bottom_border(subtitle)

    add_metadata_table(doc)

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "LumenCore is a founder-built, prototype-stage AI evidence and evaluation platform. "
        "It is designed to help acquisition and technical review teams determine what an "
        "AI-assisted workflow measured, which sources and assumptions it used, how it compared "
        "with a locked incumbent baseline, and which conclusions remain unsupported. The product "
        "is still in development and is not represented as federally deployed, commercially "
        "validated, certified, or authorized for autonomous acquisition decisions."
    )
    doc.add_paragraph(
        "For CDC, LumenCore could be demonstrated as a bounded evidence layer around acquisition "
        "planning, market research, requirements development, proposal evaluation support, "
        "contract administration, and closeout. Human officials retain all decision and approval "
        "authority."
    )

    doc.add_heading("Acquisition Lifecycle Alignment", level=1)
    add_bullet(doc, "Planning and market research: register the acquisition question, source inventory, update times, assumptions, and evidence gaps before analysis begins.")
    add_bullet(doc, "Requirements development: trace each requirement to source material, acceptance criteria, exclusions, and reviewer ownership.")
    add_bullet(doc, "AI tool and workflow evaluation: compare candidate workflows against a preselected manual or non-AI baseline using locked metrics and holdout rules.")
    add_bullet(doc, "Proposal evaluation support: preserve evidence citations, scoring rationale, uncertainty, conflicts, abstentions, and human reviewer decisions without automating source selection.")
    add_bullet(doc, "Administration and closeout: assemble milestone, deliverable, exception, and disposition records into machine-readable manifests and concise reviewer packets.")

    doc.add_page_break()
    doc.add_heading("Technical Approach", level=1)
    add_numbered(doc, "1", "Source gate - record origin, owner, timestamp, access boundary, hash, transformation, and measured-versus-synthetic status.")
    add_numbered(doc, "2", "Evaluation-plan gate - freeze the baseline, candidate, metric, threshold, holdout policy, and prohibited claims before scoring.")
    add_numbered(doc, "3", "Replay gate - execute the approved comparison, preserve configuration and failures, and block results when source quality or required controls fail.")
    add_numbered(doc, "4", "Review gate - generate a machine-readable manifest plus a human-readable proof card showing improvements, regressions, uncertainty, missing evidence, and claim boundaries.")
    add_numbered(doc, "5", "Decision gate - require an authorized human to approve any recommendation, pilot transition, external transmission, or high-impact action.")

    doc.add_heading("Current Product and Evidence Status", level=1)
    add_bullet(doc, "Working prototype components include source inventory, provenance records, locked evaluation plans, baseline-versus-candidate replay, evidence manifests, reviewer summaries, and fail-closed workflow controls.")
    add_bullet(doc, "Public reviewer surface: https://lumen-core.ai/proof_to_pilot.html")
    add_bullet(doc, "Evidence is internally reproducible and hash-addressed; it is not represented as CDC validation, field performance, realized savings, ATO, FedRAMP authorization, or CMMC certification.")
    add_bullet(doc, "LumenCore is available for a controlled, non-production demonstration using synthetic, public, or CDC-authorized data and pre-agreed evaluation criteria.")

    doc.add_heading("Proposed Reverse Industry Day Demonstration", level=1)
    doc.add_paragraph(
        "LumenCore would demonstrate a small acquisition workflow in which a requirement and "
        "evaluation plan are registered before scoring; two AI-assisted approaches are compared "
        "with a locked baseline; provenance, configuration, exceptions, and uncertainty are "
        "captured; and a reviewer packet distinguishes measured evidence from assumptions. The "
        "demonstration would emphasize auditability, abstention when evidence is insufficient, and "
        "human control rather than autonomous acquisition authority."
    )

    doc.add_heading("Contact", level=1)
    contact = doc.add_paragraph()
    contact.paragraph_format.keep_with_next = True
    run = contact.add_run("Robert Ashworth | LumenCore | Nashville, Tennessee")
    run.bold = True
    contact.add_run("\nrobertashworth4444@gmail.com | https://lumen-core.ai")

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("grant_submissions/funding_sprint_20260709/LumenCore_CDC_AI_Acquisition_RFI_75D301-26-RFI-73483_2026-07-15.docx"),
    )
    args = parser.parse_args()
    build(args.output.resolve())


if __name__ == "__main__":
    main()
